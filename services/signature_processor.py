import re
from typing import Optional, Tuple, Pattern, List, Union
from docx import Document
from docx.text.paragraph import Paragraph
from docx.document import Document as DocumentType


def create_name_regex(name_str: str) -> Pattern[str]:
    """
    Создает регулярное выражение для поиска ФИО, учитывающее различные порядки
    (Фамилия И.О. или И.О. Фамилия) и наличие пробелов/точек.
    """
    parts: List[str] = re.findall(r'[А-ЯЁа-яёA-Za-z]+', name_str)
    if not parts:
        return re.compile(re.escape(name_str), re.IGNORECASE)

    surname: str = parts[0] if len(parts[0]) > 2 else parts[-1]
    initials: List[str] = [p[0] for p in parts if p != surname]

    if not initials:
        return re.compile(re.escape(surname), re.IGNORECASE)

    dots_spaces: str = r"\.?\s*"
    init_pattern: str = dots_spaces.join(initials) + r"\.?"

    pattern1: str = rf"{init_pattern}\s*{surname}"
    pattern2: str = rf"{surname}\s*{init_pattern}"

    return re.compile(f"({pattern1}|{pattern2})", re.IGNORECASE)


def create_title_regex(title_str: Optional[str]) -> Optional[Pattern[str]]:
    """
    Создает гибкое регулярное выражение для должности, учитывающее
    возможные сокращения (например, 'зав.' вместо 'заведующий').
    """
    if not title_str:
        return None
    pattern: str = re.escape(title_str)
    pattern = pattern.replace(r'заведующий', r'зав(едующий)?\.?')
    pattern = pattern.replace(r'Заведующий', r'[Зз]ав(едующий)?\.?')
    pattern = pattern.replace(r'\ ', r'\s+')
    return re.compile(pattern, re.IGNORECASE)


def is_signature_zone(paragraph: Paragraph, cell_context: bool = False) -> bool:
    """
    Определяет, является ли параграф зоной подписи на основе ключевых слов,
    длины текста и визуальных индикаторов (подчеркивания, даты).
    """
    text: str = paragraph.text.strip()
    if not text:
        return False

    if len(text) > 200:
        return False

    indicators: List[str] = ["_", "20", "г.", "____________"]
    has_indicator: bool = any(ind in text for ind in indicators)

    keywords: List[str] = [
        "зав", "кафедр", "руковод", "программ", "декан",
        "председател", "составител", "разработчик"
    ]

    para_content: str = text.lower()
    has_keyword: bool = any(kw in para_content for kw in keywords)

    if cell_context:
        return True

    return has_indicator or has_keyword


def process_docx_signatures(
    file_path: str,
    old_name: str,
    new_name: str,
    old_title: Optional[str] = None,
    new_title: Optional[str] = None
) -> Tuple[bool, str]:
    """
    Ищет и заменяет ФИО и должность в документе DOCX, ограничивая область поиска
    только зонами подписей (в тексте и таблицах).
    """
    try:
        doc: DocumentType = Document(file_path)
        name_regex: Pattern[str] = create_name_regex(old_name)
        title_regex: Optional[Pattern[str]] = create_title_regex(old_title) if old_title else None
        is_changed: bool = False

        def replace_in_paragraph(p: Paragraph, is_cell: bool = False) -> None:
            nonlocal is_changed
            has_name: bool = bool(name_regex.search(p.text))
            has_title: bool = bool(title_regex.search(p.text)) if title_regex else False

            if has_name or has_title:
                if is_signature_zone(p, is_cell):
                    full_text: str = p.text

                    if has_title and title_regex and new_title:
                        full_text = title_regex.sub(new_title, full_text)

                    if has_name:
                        full_text = name_regex.sub(new_name, full_text)

                    if full_text != p.text:
                        for i in range(len(p.runs)):
                            p.runs[i].text = ""
                        if p.runs:
                            p.runs[0].text = full_text
                        else:
                            p.add_run(full_text)
                        is_changed = True

        for para in doc.paragraphs:
            replace_in_paragraph(para)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        replace_in_paragraph(para, is_cell=True)

        if is_changed:
            doc.save(file_path)
            return True, "Успешно обновлено"
        else:
            return False, "Данные не найдены в подходящих блоках"

    except Exception as e:
        return False, str(e)