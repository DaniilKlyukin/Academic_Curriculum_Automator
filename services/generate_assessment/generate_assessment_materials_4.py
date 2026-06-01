import os
import re
import copy
import logging
import time
import json
from pathlib import Path
from typing import List, Dict, Any, Optional
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Попытка импорта официальной библиотеки Google AI
try:
    import google.generativeai as genai

    HAS_GEMINI = True
except ImportError:
    HAS_GEMINI = False

logger = logging.getLogger(__name__)


class RateLimiter:
    """Класс для ограничения частоты запросов к API (Rate Limiting)."""

    def __init__(self, rpm: int = 15):
        self.delay = 60.0 / rpm if rpm > 0 else 0.0
        self.last_call = 0.0

    def wait(self):
        if self.delay <= 0:
            return
        now = time.time()
        elapsed = now - self.last_call
        if elapsed < self.delay:
            time.sleep(self.delay - elapsed)
        self.last_call = time.time()


def set_font(run, font_name="Times New Roman", size_pt=14, bold=False, italic=False, color_rgb=(0, 0, 0)):
    """Устанавливает параметры шрифта для текстового прогона (run)."""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(*color_rgb)

    try:
        rPr = run._r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:cs'), font_name)
        rFonts.set(qn('w:eastAsia'), font_name)
        rPr.append(rFonts)
    except Exception as e:
        logger.debug(f"Не удалось применить XML-стили шрифтов: {e}")


def format_cell_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT):
    """Сбрасывает интервалы абзаца внутри ячеек таблицы."""
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0


def resolve_style_names(doc: Document):
    """Определяет доступные имена системных стилей в документе."""
    existing_styles = {s.name for s in doc.styles}
    style_normal = "Normal"
    for name in ["Обычный", "Normal"]:
        if name in existing_styles:
            style_normal = name
            break
    return style_normal


def extract_direction_from_doc(doc: Document) -> str:
    """Извлекает название направления подготовки с титульного листа документа.

    Считывает данные, перенесенные из листа Excel 'Титул' на первом шаге генерации.
    """
    paragraphs = list(doc.paragraphs)
    for i, p in enumerate(paragraphs[:60]):  # Сканируем первые 60 абзацев (титульный лист)
        txt = p.text.strip()
        # Поиск по ключевой фразе поля ввода
        if "направление подготовки" in txt.lower():
            # Направление находится в абзацах сразу под меткой поля
            for offset in range(1, 4):
                if i + offset < len(paragraphs):
                    val = paragraphs[i + offset].text.strip()
                    # Исключаем пустые строки, символы подчеркивания и подстрочные подсказки
                    if val and not val.startswith("_") and "код и наименование" not in val.lower():
                        return val
    return ""


def clean_name_for_match(name: str) -> str:
    """Очищает строку для точного сопоставления имен дисциплин."""
    s = name.lower()
    s = "".join(s.split())
    s = s.replace("«", "").replace("»", "").replace('"', '').replace("'", "")
    return s


def get_rp_title(doc: Document) -> str:
    """Извлекает официальное название дисциплины с титульного листа РП."""
    found_marker = False
    for p in doc.paragraphs:
        txt = " ".join(p.text.split()).strip().lower()
        if "рабочая программа" in txt or "программа дисциплины" in txt:
            found_marker = True
            continue
        if found_marker:
            title = " ".join(p.text.split()).strip()
            if title:
                return title
    for p in doc.paragraphs[:12]:
        txt = " ".join(p.text.split()).strip()
        if txt and len(txt) > 5 and not txt.isupper() and "министерство" not in txt.lower():
            return txt
    return ""


def find_test_elements(doc: Document) -> tuple[list[Any], bool]:
    """Находит XML-элементы теста в документе РП."""
    body_elements = list(doc.element.body)
    start_idx = -1

    for i, elem in enumerate(body_elements):
        if elem.tag.endswith('p'):
            text = "".join(elem.itertext()).strip().lower()
            if "примерныйвариант" in text or "вариантитоговоготеста" in text or "итоговыйтест" in text or "варианттеста" in text:
                start_idx = i
                break

    if start_idx == -1:
        for i, elem in enumerate(body_elements):
            text = "".join(elem.itertext()).strip().lower()
            if "тест" in text and len(text) < 100:
                start_idx = i
                break

    if start_idx == -1:
        return [], False

    test_elements = []
    has_keys_table = False

    for i in range(start_idx, len(body_elements)):
        elem = body_elements[i]
        text = "".join(elem.itertext()).strip().lower()

        if i > start_idx:
            if "критерии" in text or "наименование:" in text or "представлениевфос:" in text:
                break

            if "ключи" in text or "ответы" in text:
                test_elements.append(elem)
                for j in range(i + 1, min(i + 5, len(body_elements))):
                    next_elem = body_elements[j]
                    if next_elem.tag.endswith('tbl'):
                        test_elements.append(next_elem)
                        has_keys_table = True
                        break
                break

        test_elements.append(elem)

    if not has_keys_table:
        return [], False

    return test_elements, True


def generate_test_via_ai(subject_name: str, competency_text: str, api_key: str, rate_limiter: RateLimiter, direction: str = "", retries: int = 3) -> Optional[dict]:
    """Запрашивает генерацию вопросов теста через API Gemini с поддержкой повторных попыток (retries)."""
    if not HAS_GEMINI:
        logger.error("Библиотека google-generativeai не установлена. Выполните: pip install google-generativeai")
        return None

    # Жесткое ограничение на упоминание метаданных компетенций в тестах
    avoid_metadata_instruction = (
        "КАТЕГОРИЧЕСКИ ЗАПРЕЩЕНО упоминать код компетенции (например, УК-1, ОПК-2), само слово 'компетенция', "
        "а также методические термины (например, 'индикатор', 'знать', 'уметь', 'владеть') в текстах вопросов и вариантов ответов."
    )

    # Требование к грамотному, но не избыточно официальному стилю
    style_instruction = (
        "Стиль написания вопросов: используйте четкий, грамотный и профессиональный язык. Избегайте избыточного академизма, "
        "канцеляризмов, слишком длинных и сложных предложений. Формулировки должны быть лаконичными, ясными и понятными "
        "для студентов, но при этом академически корректными. Не используйте пассивный залог там, где можно обойтись активным "
        "Стремитесь к естественности и точности формулировок."
    )

    # Формирование промпта в зависимости от контекста (Раздел 2 или Раздел 3)
    if not subject_name:
        # Для Раздела 3 (варианты диагностической работы без привязки к конкретному предмету)
        subject_context = f"общего междисциплинарного оценочного средства по направлению подготовки «{direction}»" if direction else "общего междисциплинарного оценочного средства"
        prompt = f"""Вы — опытный преподаватель вуза, способный четко и грамотно формулировать задания. 
Сгенерируйте качественный академический тест из 5 вопросов для студентов в рамках {subject_context}.

Компетенция (используйте ее исключительно как тематический ориентир для подбора вопросов, но не упоминайте в тексте): {competency_text}

Требования к тесту:
1. Ровно 5 вопросов множественного выбора. Вопросы должны оценивать общее междисциплинарное понимание и практическое применение данной компетенции во всей профессиональной области {f'«{direction}»' if direction else ''}.
2. {avoid_metadata_instruction}
3. {style_instruction}
4. Для каждого вопроса должно быть ровно 4 варианта ответа (маркированные русскими строчными буквами: а, б, в, г).
5. Вопросы должны быть содержательными, концептуальными и высокоуровневыми.
6. Ответ должен быть СТРОГО в формате JSON, без какого-либо окружающего текста, комментариев или разметки. Схема JSON:
{{
  "test_questions": [
    {{
      "number": 1,
      "question": "Текст вопроса...",
      "options": {{
        "а": "Текст варианта а",
        "б": "Текст варианта б",
        "в": "Текст варианта в",
        "г": "Текст варианта г"
      }},
      "correct_option": "а"
    }}
  ]
}}
Верните исключительно валидный JSON.
"""
    else:
        # Для Раздела 2 (конкретная дисциплина или практика)
        prompt = f"""Вы — опытный преподаватель вуза, способный четко и грамотно формулировать задания. 
Сгенерируйте качественный академический тест из 5 вопросов для студентов по дисциплине.

Дисциплина/Практика: {subject_name}
Компетенция (используйте ее исключительно как тематический ориентир для подбора вопросов, но не упоминайте в тексте): {competency_text}

Требования к тесту:
1. Ровно 5 вопросов множественного выбора.
2. {avoid_metadata_instruction}
3. {style_instruction}
4. Для каждого вопроса должно быть ровно 4 варианта ответа (маркированные русскими строчными буквами: а, б, в, г).
5. Вопросы должны быть содержательными и напрямую соотноситься со спецификой указанной компетенции и предмета.
6. Ответ должен быть СТРОГО в формате JSON, без какого-либо окружающего текста, комментариев или разметки. Схема JSON:
{{
  "test_questions": [
    {{
      "number": 1,
      "question": "Текст вопроса...",
      "options": {{
        "а": "Текст варианта а",
        "б": "Текст варианта б",
        "в": "Текст варианта в",
        "г": "Текст варианта г"
      }},
      "correct_option": "а"
    }}
  ]
}}
Верните исключительно валидный JSON.
"""

    for attempt in range(1, retries + 1):
        rate_limiter.wait()
        try:
            genai.configure(api_key=api_key)
            model_name = "gemini-3.1-flash-lite-preview" #"gemini-3.1-flash-lite-preview"
            model = genai.GenerativeModel(
                model_name,
                generation_config={"response_mime_type": "application/json"}
            )
            response = model.generate_content(prompt)
            text_response = response.text.strip()

            if text_response.startswith("```json"):
                text_response = text_response.split("```json", 1)[1]
            elif text_response.startswith("```"):
                text_response = text_response.split("```", 1)[1]
            if text_response.endswith("```"):
                text_response = text_response.rsplit("```", 1)[0]
            text_response = text_response.strip()

            data = json.loads(text_response)
            return data
        except Exception as e:
            logger.warning(f"Попытка {attempt}/{retries} генерации теста завершилась ошибкой для '{subject_name or 'Раздел 3'}' (Компетенция: {competency_text[:40]}): {e}")
            if attempt < retries:
                time.sleep(3)  # Небольшая пауза перед следующей попыткой
            else:
                logger.error(f"Не удалось сгенерировать тест после {retries} попыток.")
    return None


def build_test_elements(test_data: dict, style_normal: str) -> list:
    """Строит структуру XML-элементов (абзацы и таблицу) на основе сгенерированного ИИ JSON."""
    temp_doc = Document()
    questions = test_data.get("test_questions", [])

    for q in questions:
        num = q.get("number", 1)
        q_text = q.get("question", "")
        options = q.get("options", {})

        # Вопрос (шрифт 14pt, жирный)
        p_q = temp_doc.add_paragraph()
        try:
            p_q.style = style_normal
        except Exception:
            pass
        p_q.paragraph_format.space_before = Pt(6)
        p_q.paragraph_format.space_after = Pt(2)
        p_q.paragraph_format.line_spacing = 1.0

        run_num = p_q.add_run(f"{num}. ")
        set_font(run_num, size_pt=14, bold=True)
        run_text = p_q.add_run(q_text)
        set_font(run_text, size_pt=14, bold=True)

        # Варианты ответов (буквами а, б, в, г, шрифт 14pt, не жирный)
        for opt_key in ["а", "б", "в", "г"]:
            opt_val = options.get(opt_key, "")
            p_opt = temp_doc.add_paragraph()
            try:
                p_opt.style = style_normal
            except Exception:
                pass
            p_opt.paragraph_format.space_before = Pt(0)
            p_opt.paragraph_format.space_after = Pt(2)
            p_opt.paragraph_format.line_spacing = 1.0

            run_opt_lbl = p_opt.add_run(f"{opt_key}) ")
            set_font(run_opt_lbl, size_pt=14, bold=False)
            run_opt_val = p_opt.add_run(opt_val)
            set_font(run_opt_val, size_pt=14, bold=False)

        # Пустой абзац-разделитель между вопросами
        p_space = temp_doc.add_paragraph()
        try:
            p_space.style = style_normal
        except Exception:
            pass
        p_space.paragraph_format.space_after = Pt(6)

    # Метка таблицы ключей
    p_keys_lbl = temp_doc.add_paragraph()
    try:
        p_keys_lbl.style = style_normal
    except Exception:
        pass
    p_keys_lbl.paragraph_format.space_before = Pt(12)
    p_keys_lbl.paragraph_format.space_after = Pt(6)
    run_keys_lbl = p_keys_lbl.add_run("Ключи теста")
    set_font(run_keys_lbl, size_pt=14, bold=False)

    # Таблица ключей ответов
    table = temp_doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT

    col_widths = [Inches(1.5), Inches(4.5)]
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width

    # Шапка таблицы
    headers = ["Вопрос", "Ответ"]
    for idx, text in enumerate(headers):
        cell = table.cell(0, idx)
        cell.text = ""
        p = cell.paragraphs[0]
        try:
            p.style = style_normal
        except Exception:
            pass
        format_cell_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER)
        run = p.add_run(text)
        set_font(run, size_pt=12, bold=True)

    # Строки ответов
    for idx, q in enumerate(questions, 1):
        cell_num = table.cell(idx, 0)
        cell_num.text = ""
        p_num = cell_num.paragraphs[0]
        try:
            p_num.style = style_normal
        except Exception:
            pass
        format_cell_paragraph(p_num, align=WD_ALIGN_PARAGRAPH.CENTER)
        run_num = p_num.add_run(str(idx))
        set_font(run_num, size_pt=12, bold=False)

        correct = q.get("correct_option", "")
        cell_ans = table.cell(idx, 1)
        cell_ans.text = ""
        p_ans = cell_ans.paragraphs[0]
        try:
            p_ans.style = style_normal
        except Exception:
            pass
        format_cell_paragraph(p_ans, align=WD_ALIGN_PARAGRAPH.LEFT)
        run_ans = p_ans.add_run(str(correct))
        set_font(run_ans, size_pt=12, bold=False)

    p_tail = temp_doc.add_paragraph()
    try:
        p_tail.style = style_normal
    except Exception:
        pass

    elements_to_insert = [el for el in temp_doc.element.body if el.tag.endswith('p') or el.tag.endswith('tbl')]
    return elements_to_insert


def insert_elements_after_paragraph(anchor_paragraph, elements: list):
    """Вставляет копии XML-элементов (параграфы и таблицы) после указанного абзаца."""
    current_element = anchor_paragraph._element
    for elem in elements:
        new_elem = copy.deepcopy(elem)
        current_element.addnext(new_elem)
        current_element = new_elem


class CompetencyReportGenerator:
    """Класс управления интеграцией тестов из Рабочих Программ и ИИ."""

    def __init__(self, word_path: str, rp_folder_path: str, ai_mode: int = 3, api_key: str = "", rpm_limit: int = 15):
        self.word_path = word_path
        self.rp_folder_path = rp_folder_path
        self.ai_mode = ai_mode
        self.api_key = api_key
        self.rpm_limit = rpm_limit
        self.rate_limiter = RateLimiter(rpm=rpm_limit)

    def generate(self):
        logger.info(f"Загрузка итогового документа: {self.word_path}")
        if not os.path.exists(self.word_path):
            logger.error(f"Файл не найден: {self.word_path}")
            return

        try:
            doc = Document(self.word_path)
        except Exception as e:
            logger.error(f"Не удалось открыть документ: {e}")
            return

        rp_folder = Path(self.rp_folder_path)
        rp_map: Dict[str, Path] = {}

        # Сканирование файлов РП производится только при использовании смешанного или оригинального режимов
        if self.ai_mode in [2, 3]:
            if rp_folder.exists():
                logger.info("Сканирование папки с Рабочими Программами...")
                for docx_path in rp_folder.glob("*.docx"):
                    if docx_path.name.startswith("~$"):
                        continue
                    try:
                        rp_doc = Document(docx_path)
                        title = get_rp_title(rp_doc)
                        if title:
                            cleaned_title = clean_name_for_match(title)
                            rp_map[cleaned_title] = docx_path
                            logger.info(f"  -> Найдена РП: '{title}' ({docx_path.name})")
                    except Exception as e:
                        logger.warning(f"Не удалось прочитать файл РП {docx_path.name}: {e}")
                logger.info(f"Успешно проиндексировано РП: {len(rp_map)}")
            else:
                logger.warning(f"Папка с РП не найдена по пути: {self.rp_folder_path}")

        logger.info("Интеграция тестов в структуру документа...")
        subject_pattern = re.compile(r'^(Дисциплина|Практика)\s+«(.*?)»')

        # Автоматически определяем направление подготовки для Раздела 3
        direction = extract_direction_from_doc(doc)
        if direction:
            logger.info(f"Определено направление подготовки для междисциплинарных тестов: '{direction}'")
        else:
            logger.warning(
                "Направление подготовки не обнаружено на титульном листе. Будут сгенерированы общие вопросы.")

        paragraphs = list(doc.paragraphs)
        updated_count = 0
        current_competency = "Не определена"

        for i, p in enumerate(paragraphs):
            p_text = p.text.strip()

            # Обновление названия активной компетенции
            if i > 0 and paragraphs[i - 1].text.strip().lower() == "компетенция":
                current_competency = p_text

            is_subject = False
            is_variant = False
            subj_type = ""
            subj_name = ""

            # Распознаем Дисциплину/Практику (Раздел 2) или Вариант (Раздел 3)
            match = subject_pattern.match(p_text)
            if match:
                is_subject = True
                subj_type = match.group(1)
                subj_name = match.group(2)
                cleaned_subj_name = clean_name_for_match(subj_name)
            elif p_text.lower() == "вариант":
                is_variant = True

            if not (is_subject or is_variant):
                continue

            use_ai_for_this = False
            test_elements = None
            has_keys_table = False

            # === ОБРАБОТКА ДИСЦИПЛИН И ПРАКТИК (РАЗДЕЛ 2) ===
            if is_subject:
                if self.ai_mode == 1:
                    use_ai_for_this = True
                elif self.ai_mode == 2:
                    rp_path = rp_map.get(cleaned_subj_name)
                    if rp_path:
                        try:
                            rp_doc = Document(rp_path)
                            test_elements, has_keys_table = find_test_elements(rp_doc)
                        except Exception as e:
                            logger.error(f"Ошибка парсинга РП {rp_path.name}: {e}")

                    if not (test_elements and has_keys_table):
                        use_ai_for_this = True
                else:
                    rp_path = rp_map.get(cleaned_subj_name)
                    if rp_path:
                        try:
                            rp_doc = Document(rp_path)
                            test_elements, has_keys_table = find_test_elements(rp_doc)
                        except Exception as e:
                            logger.error(f"Ошибка парсинга РП {rp_path.name}: {e}")

                if use_ai_for_this:
                    if not self.api_key:
                        logger.warning(
                            f"  [-] Режим ИИ активен, но API-ключ не задан. Пропускаем дисциплину «{subj_name}».")
                        continue

                    print(
                        f"  [ИИ] Генерация теста для {subj_type.lower()} «{subj_name}» (Компетенция: {current_competency})...")
                    ai_data = generate_test_via_ai(subj_name, current_competency, self.api_key, self.rate_limiter)
                    if ai_data:
                        style_normal = resolve_style_names(doc)
                        test_elements = build_test_elements(ai_data, style_normal)
                        has_keys_table = True
                    else:
                        logger.warning(f"  [-] Не удалось получить данные теста через ИИ для «{subj_name}».")
                        continue

                if test_elements and has_keys_table:
                    p_conduct = None
                    for k in range(i + 1, min(i + 5, len(paragraphs))):
                        if "проведение работы" in paragraphs[k].text.lower():
                            p_conduct = paragraphs[k]
                            break

                    if not p_conduct:
                        continue

                    main_elements = list(doc.element.body)
                    p_idx = main_elements.index(p_conduct._element)

                    elements_to_delete = []
                    for k in range(p_idx + 1, len(main_elements)):
                        elem = main_elements[k]
                        if elem.tag.endswith('tbl'):
                            elements_to_delete.append(elem)
                            break
                        else:
                            elements_to_delete.append(elem)

                    for elem in elements_to_delete:
                        doc.element.body.remove(elem)

                    insert_elements_after_paragraph(p_conduct, test_elements)

                    if use_ai_for_this:
                        logger.info(f"  [+] Добавлен ИИ-тест для {subj_type.lower()} «{subj_name}»")
                    else:
                        logger.info(f"  [+] Добавлен тест из РП для {subj_type.lower()} «{subj_name}»")
                    updated_count += 1

            # === ОБРАБОТКА ВАРИАНТОВ ДИАГНОСТИЧЕСКОЙ РАБОТЫ (РАЗДЕЛ 3) ===
            elif is_variant:
                # В Разделе 3 тесты создаются только если ИИ-режимы включены (1 или 2)
                if self.ai_mode in [1, 2]:
                    if not self.api_key:
                        logger.warning(
                            f"  [-] Режим ИИ активен, но API-ключ не задан. Пропускаем междисциплинарный тест для «{current_competency}».")
                        continue

                    print(f"  [ИИ] Генерация общего междисциплинарного теста для {current_competency}...")
                    ai_data = generate_test_via_ai("", current_competency, self.api_key, self.rate_limiter,
                                                   direction=direction)
                    if ai_data:
                        style_normal = resolve_style_names(doc)
                        test_elements = build_test_elements(ai_data, style_normal)
                        has_keys_table = True
                    else:
                        logger.warning(
                            f"  [-] Не удалось получить данные теста через ИИ для Варианта по компетенции {current_competency}.")
                        continue

                    if test_elements and has_keys_table:
                        main_elements = list(doc.element.body)
                        p_idx = main_elements.index(p._element)

                        # Очищаем заглушки непосредственно после слова "Вариант" до пустой таблицы
                        elements_to_delete = []
                        for k in range(p_idx + 1, len(main_elements)):
                            elem = main_elements[k]
                            if elem.tag.endswith('tbl'):
                                elements_to_delete.append(elem)
                                break
                            else:
                                elements_to_delete.append(elem)

                        for elem in elements_to_delete:
                            doc.element.body.remove(elem)

                        insert_elements_after_paragraph(p, test_elements)
                        logger.info(f"  [+] Добавлен междисциплинарный ИИ-тест для Варианта ({current_competency})")
                        updated_count += 1

        try:
            doc.save(self.word_path)
            print(f"\n[Успешно] Обработка тестов завершена. Интегрировано тестов: {updated_count}")
        except Exception as e:
            logger.error(f"Не удалось сохранить обновленный файл Word: {e}")


def main():
    print("=== Панель управления интеграцией тестов из РП и ИИ ===")

    user_word_path: str = input(
        "Шаг 1. Введите путь к итоговому файлу Word (например, Оценочные материалы.docx): "
    ).strip()
    user_rp_folder: str = input("Шаг 2. Введите путь к папке с Рабочими Программами (РП): ").strip()

    if not user_word_path:
        user_word_path = "Оценочные материалы.docx"

    if not user_rp_folder:
        user_rp_folder = "."

    print("\nРежимы работы:")
    print("  1 - Полный ИИ (все тесты генерируются ИИ, файлы РП игнорируются)")
    print("  2 - Смешанный (тесты берутся из РП, при отсутствии — генерируются ИИ)")
    print("  3 - Без ИИ (только из РП, при отсутствии — остаются заглушки)")

    ai_mode_input = input("Выберите режим [По умолчанию: 3]: ").strip()
    ai_mode = int(ai_mode_input) if ai_mode_input in ["1", "2", "3"] else 3

    api_key = ""
    rpm_limit = 15
    if ai_mode in [1, 2]:
        api_key = os.environ.get("GEMINI_API_KEY", "").strip()
        if not api_key:
            api_key = input(
                "Введите ваш API-ключ Gemini (или оставьте пустым при наличии системной переменной): ").strip()
        rpm_input = input("Лимит запросов в минуту (RPM) [По умолчанию: 15]: ").strip()
        if rpm_input.isdigit():
            rpm_limit = int(rpm_input)

    print("\nЗапуск процесса интеграции...")
    generator = CompetencyReportGenerator(
        word_path=user_word_path,
        rp_folder_path=user_rp_folder,
        ai_mode=ai_mode,
        api_key=api_key,
        rpm_limit=rpm_limit
    )
    generator.generate()


if __name__ == "__main__":
    main()