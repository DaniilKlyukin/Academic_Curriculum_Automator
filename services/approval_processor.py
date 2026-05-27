import os
import re
from typing import List, Tuple
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.table import _Cell


def generate_years(start_year: int, end_year: int) -> List[str]:
    """
    Генерирует список строк учебных годов в формате 'YYYY – YYYY+1'.
    """
    return [f"{y} – {y + 1}" for y in range(start_year, end_year)]


def set_cell_format(cell: _Cell, text: str, align_center: bool = False) -> None:
    """
    Устанавливает текст, шрифт и выравнивание для указанной ячейки таблицы.
    """
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if align_center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)


def process_docx(file_path: str, years_list: List[str]) -> Tuple[bool, str]:
    """
    Ищет в документе таблицу согласования РПД по ключевым словам и обновляет в ней список годов.
    """
    try:
        doc = Document(file_path)
        keywords: List[str] = ["учебн", "год", "согласов", "рпд", "лист"]

        for table in doc.tables:
            if not table.rows:
                continue

            check_limit: int = min(len(table.rows), 3)
            rows_data: List[str] = []
            for i in range(check_limit):
                text: str = " ".join(cell.text for cell in table.rows[i].cells).lower()
                rows_data.append(re.sub(r'\s+', ' ', text))

            full_header_text: str = " ".join(rows_data)
            matches: List[str] = [kw for kw in keywords if kw in full_header_text]

            if len(matches) >= 2:
                target_row_index: int = 0
                for idx, text in enumerate(rows_data):
                    if "учебн" in text and "год" in text:
                        target_row_index = idx
                    elif "согласов" in text and target_row_index == 0:
                        target_row_index = idx

                for i in range(len(table.rows) - 1, target_row_index, -1):
                    row_el = table.rows[i]._element
                    row_el.getparent().remove(row_el)

                for year_str in years_list:
                    new_row = table.add_row()
                    set_cell_format(new_row.cells[0], year_str, align_center=True)
                    if len(new_row.cells) > 1:
                        set_cell_format(new_row.cells[1], "\n\n", align_center=False)

                doc.save(file_path)
                return True, "Успешно"

        return False, "Таблица не найдена"
    except Exception as e:
        return False, str(e)


def main():
    print("=== Обновление учебных лет в Листах согласования ===")
    folder_path = input("Введите путь до папки с .docx: ").strip().strip('"')

    if not os.path.isdir(folder_path):
        print("Ошибка: Путь не найден.")
        return

    try:
        start_y = int(input("Введите начальный год: "))
        end_y = int(input("Введите конечный год: "))
    except ValueError:
        print("Ошибка: Введите числа.")
        return

    years_list = generate_years(start_y, end_y)

    files = [f for f in os.listdir(folder_path) if f.lower().endswith('.docx') and not f.startswith('~$')]
    total = len(files)

    print(f"\n{'№':<9} | {'Статус':<8} | {'Файл'}")
    print("-" * 80)

    for i, filename in enumerate(files, 1):
        file_path = os.path.join(folder_path, filename)
        success, message = process_docx(file_path, years_list)

        status = "OK" if success else "SKIP"

        print(f"[{i:03}/{total:03}] | {status:<8} | {os.path.basename(file_path)[:60]}")

if __name__ == "__main__":
    main()