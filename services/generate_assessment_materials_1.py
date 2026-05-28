import re
import logging
from pathlib import Path
from typing import List, Dict, Set, Tuple, Optional, Union, Any
from openpyxl import load_workbook
from openpyxl.worksheet.worksheet import Worksheet
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


class ExcelConfig:
    """Конфигурация столбцов и имен листов Excel по умолчанию."""
    COMPETENCY_SHEET_NAME: str = "Компетенции"
    PLAN_SHEET_NAME: str = "ПланСвод"  # Приоритетный лист по умолчанию

    # Дефолтная колонка для текста содержания (если не найдена автопоиском)
    COMP_TXT_COL: int = 5

    PLAN_IDX_COL: int = 2
    PLAN_NAME_COL: int = 3


def set_font(run, font_name="Times New Roman", size_pt=12, bold=False, italic=False, color_rgb=(0, 0, 0)):
    """Устанавливает шрифт, размер, цвет и начертание для текстового прогона (run),

    обеспечивая поддержку кириллицы в MS Word.
    """
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(*color_rgb)

    try:
        rPr = run._r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:cs'), font_name)
        rFonts.set(qn('w:eastAsia'), font_name)
        rPr.append(rFonts)
    except Exception as e:
        logger.debug(f"Не удалось применить XML-стили шрифтов: {e}")


def resolve_style_names(doc: Document):
    """Определяет доступные системные имена стилей в открытом документе."""
    existing_styles = {s.name for s in doc.styles}

    style_h1 = "Heading 1"
    for name in ["Заголовок 1", "Heading 1", "Heading1"]:
        if name in existing_styles:
            style_h1 = name
            break

    style_h2 = "Heading 2"
    for name in ["Заголовок 2", "Heading 2", "Heading2"]:
        if name in existing_styles:
            style_h2 = name
            break

    style_normal = "Normal"
    for name in ["Обычный", "Normal"]:
        if name in existing_styles:
            style_normal = name
            break

    return style_h1, style_h2, style_normal


def configure_document_styles(doc: Document):
    """Программно настраивает глобальные стили документа согласно заданным параметрам."""
    style_h1, style_h2, style_normal = resolve_style_names(doc)
    black_color = (0, 0, 0)

    def setup_style(style_obj, font_name="Times New Roman", size_pt=14, bold=False, italic=False, color_rgb=black_color,
                    align=None):
        font = style_obj.font
        font.name = font_name
        font.size = Pt(size_pt)
        font.bold = bold
        font.italic = italic
        font.color.rgb = RGBColor(*color_rgb)

        # Настройка абзацных интервалов и выравнивания
        pf = style_obj.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0
        if align is not None:
            pf.alignment = align

    # Настройка стиля «Обычный» (Обычный текст по ширине)
    if style_normal in doc.styles:
        setup_style(doc.styles[style_normal], size_pt=14, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # Настройка стиля «Заголовок 1» (По центру, Жирный)
    if style_h1 in doc.styles:
        setup_style(doc.styles[style_h1], size_pt=14, bold=True, italic=False, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Настройка стиля «Заголовок 2» (По ширине, Курсив)
    if style_h2 in doc.styles:
        setup_style(doc.styles[style_h2], size_pt=14, bold=False, italic=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY)


def set_repeat_table_header(row):
    """Включает повторение строки таблицы на каждой новой странице (tblHeader)."""
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    trPr.append(tblHeader)


def set_row_cant_split(row):
    """Запрещает разрыв строки таблицы при переходе на новую страницу (cantSplit)."""
    trPr = row._tr.get_or_add_trPr()
    cantSplit = OxmlElement('w:cantSplit')
    trPr.append(cantSplit)


def set_cell_text(cell, text: str, style_name="Normal", bold=False, italic=False, size_pt=12,
                  align=WD_ALIGN_PARAGRAPH.LEFT, keep_with_next=False):
    """Записывает текст в ячейку таблицы с очисткой и применением форматирования."""
    cell.text = ""  # Сброс стандартного содержимого ячейки
    p = cell.paragraphs[0]
    try:
        p.style = style_name
    except Exception:
        pass

    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0

    if keep_with_next:
        p.paragraph_format.keep_with_next = True

    if text:
        run = p.add_run(text)
        set_font(run, size_pt=size_pt, bold=bold, italic=italic)


class CompetencyReportGenerator:
    """Класс для сопоставления компетенций и дисциплин из учебного плана в документ Word."""

    def __init__(self, excel_path: Union[str, Path], word_folder_path: Union[str, Path]) -> None:
        self.excel_path: Path = Path(excel_path)
        self.word_folder_path: Path = Path(word_folder_path)

    def _find_semester_columns(self, sheet: Worksheet) -> Dict[int, int]:
        """
        Сканирует первые строки листа для автоматического определения столбцов семестров (18-25).
        """
        semester_cols: Dict[int, int] = {}
        for col in range(1, sheet.max_column + 1):
            for row in range(1, 6):
                cell_val: str = str(sheet.cell(row=row, column=col).value or "").lower()
                if "сем" in cell_val:
                    match = re.search(r"сем.*?(\d+)", cell_val)
                    if match:
                        sem_num: int = int(match.group(1))
                        semester_cols[sem_num] = col
                        break
        return semester_cols

    def _find_assessment_columns(self, sheet: Worksheet) -> List[int]:
        """
        Находит индексы колонок промежуточной аттестации (Экзамен, Зачет, Зачет с оц., КР, Реферат).
        """
        cols: List[int] = []
        for col in range(1, sheet.max_column + 1):
            for row in range(1, 6):
                val: str = str(sheet.cell(row=row, column=col).value or "").lower()
                if any(x in val for x in ["экзамен", "зачет", "кр", "реферат", "кп"]):
                    cols.append(col)
                    break
        return list(set(cols))

    def _find_competency_text_column(self, sheet: Worksheet) -> int:
        """
        Ищет столбец 'Содержание' на листе компетенций.
        """
        txt_col: int = ExcelConfig.COMP_TXT_COL
        for col in range(1, sheet.max_column + 1):
            for row in range(1, 11):
                val: str = str(sheet.cell(row=row, column=col).value or "").lower().strip()
                if "содержание" in val:
                    return col
        return txt_col

    @staticmethod
    def _extract_semesters_from_assessment(val: Any) -> List[int]:
        """
        Извлекает номера семестров из ячеек промежуточной аттестации (например, '123' -> [1, 2, 3]).
        """
        if val is None:
            return []
        val_str: str = str(val).strip()
        semesters: List[int] = []

        if val_str.isdigit():
            for char in val_str:
                sem_num = int(char)
                if 1 <= sem_num <= 12:
                    semesters.append(sem_num)
        else:
            parts = re.split(r'[,\s;]+', val_str)
            for part in parts:
                if part.isdigit():
                    sem_num = int(part)
                    if 1 <= sem_num <= 12:
                        semesters.append(sem_num)
        return list(set(semesters))

    def _parse_plan_sheet(self, sheet: Worksheet) -> Dict[str, Dict[str, Any]]:
        """
        Парсит учебный план, собирая названия предметов и семестры их проведения.
        """
        plan_db: Dict[str, Dict[str, Any]] = {}
        semester_cols: Dict[int, int] = self._find_semester_columns(sheet)
        assessment_cols: List[int] = self._find_assessment_columns(sheet)

        logger.info(f"Найдены столбцы семестров: {list(semester_cols.keys())}")
        logger.info(f"Найдены столбцы аттестации (индексы): {assessment_cols}")

        for row in range(1, sheet.max_row + 1):
            idx_val: str = str(sheet.cell(row=row, column=ExcelConfig.PLAN_IDX_COL).value or "").strip()
            if re.match(r"^(Б\d|ФТД)", idx_val):
                name_val: str = str(sheet.cell(row=row, column=ExcelConfig.PLAN_NAME_COL).value or "").strip()

                active_semesters: List[int] = []

                # 1. Считываем семестры из колонок нагрузки
                for sem_num, col_idx in semester_cols.items():
                    val: Optional[Any] = sheet.cell(row=row, column=col_idx).value
                    if val is not None:
                        try:
                            if float(val) > 0:
                                active_semesters.append(sem_num)
                        except ValueError:
                            pass

                # 2. Дополнительно считываем семестры из колонок аттестации
                for col_idx in assessment_cols:
                    val_assess: Optional[Any] = sheet.cell(row=row, column=col_idx).value
                    if val_assess is not None:
                        active_semesters.extend(self._extract_semesters_from_assessment(val_assess))

                plan_db[idx_val] = {
                    "name": name_val,
                    "semesters": sorted(list(set(active_semesters)))
                }
        return plan_db

    def _parse_competencies(self, sheet: Worksheet) -> List[Dict[str, Any]]:
        """
        Парсит структуру компетенций, динамически определяя смещение структуры (для ПК).
        """
        competencies: List[Dict[str, Any]] = []
        current_comp: Optional[Dict[str, Any]] = None

        txt_col = self._find_competency_text_column(sheet)
        logger.info(f"Определен столбец Содержания компетенций -> {txt_col}")

        # Помехоустойчивые паттерны к разным типам тире/дефисов и пробелов
        comp_pattern = re.compile(r"^([A-Za-zА-Яа-я]+)\s*[-–—]\s*\d+$")
        indicator_pattern = re.compile(r"^([A-Za-zА-Яа-я]+)\s*[-–—]\s*\d+\.\d+$")
        item_pattern = re.compile(r"^(Б\d|ФТД)")

        print("\n" + "=" * 50)
        print("ДЕТАЛЬНЫЙ ЛОГ ИМПОРТА ИЗЛИСТА 'КОМПЕТЕНЦИИ' (АДАПТИВНЫЙ ПАРСЕР):")
        print("=" * 50)

        # Вычисляем глубину сканирования с запасом
        max_scan_rows: int = max(sheet.max_row, 1000)
        consecutive_empty_rows: int = 0

        for row in range(1, max_scan_rows + 1):
            # Считываем 4 колонки из-за возможного смещения ПК вправо
            val_col1: str = str(sheet.cell(row=row, column=1).value or "").strip()
            val_col2: str = str(sheet.cell(row=row, column=2).value or "").strip()
            val_col3: str = str(sheet.cell(row=row, column=3).value or "").strip()
            val_col4: str = str(sheet.cell(row=row, column=4).value or "").strip()
            text_val: str = str(sheet.cell(row=row, column=txt_col).value or "").strip()

            # Предотвращение преждевременного завершения
            if not val_col1 and not val_col2 and not val_col3 and not val_col4:
                consecutive_empty_rows += 1
                if consecutive_empty_rows > 50:  # Прекращаем парсинг после 50 пустых строк подряд
                    break
                continue

            consecutive_empty_rows = 0  # Сбрасываем счетчик, если строка не пуста

            # 1. Поиск КОМПЕТЕНЦИИ (может быть в колонке 1 или 2)
            comp_code: Optional[str] = None
            if val_col1 and comp_pattern.match(val_col1):
                comp_code = val_col1
            elif val_col2 and comp_pattern.match(val_col2):
                comp_code = val_col2

            if comp_code:
                current_comp = {
                    "comp_code": comp_code,
                    "comp_name": text_val,
                    "indicators": [],
                    "mapped_codes": set()
                }
                competencies.append(current_comp)
                print(f"Строка {row:03} | [Компетенция] {comp_code} -> {text_val[:50]}...")
                continue

            # 2. Поиск ИНДИКАТОРА (может быть в колонке 2 или 3)
            indicator_code: Optional[str] = None
            if val_col2 and indicator_pattern.match(val_col2):
                indicator_code = val_col2
            elif val_col3 and indicator_pattern.match(val_col3):
                indicator_code = val_col3

            if indicator_code:
                if current_comp:
                    current_comp["indicators"].append(f"{indicator_code} {text_val}")
                    print(f"Строка {row:03} |   [Индикатор]   {indicator_code} -> {text_val[:50]}...")
                else:
                    print(f"Строка {row:03} |   [ВНИМАНИЕ] Пропущен индикатор без компетенции: {indicator_code}")
                continue

            # 3. Поиск ДИСЦИПЛИНЫ/ПРАКТИКИ (может быть в колонке 3 или 4)
            item_code: Optional[str] = None
            if val_col3 and item_pattern.match(val_col3):
                item_code = val_col3
            elif val_col4 and item_pattern.match(val_col4):
                item_code = val_col4

            if item_code:
                if current_comp:
                    current_comp["mapped_codes"].add(item_code)
                    print(f"Строка {row:03} |     [Привязка]    {item_code} -> {text_val[:50]}...")
                else:
                    print(f"Строка {row:03} |     [ВНИМАНИЕ] Пропущена дисциплина без компетенции: {item_code}")

        print("=" * 50 + "\n")
        return competencies

    @staticmethod
    def _format_semesters(sem_list: List[int]) -> str:
        """Форматирует список семестров в строку."""
        if not sem_list:
            return ""
        sorted_sems: List[int] = sorted(list(set(sem_list)))
        sems_str: str = ", ".join(map(str, sorted_sems))
        if len(sorted_sems) == 1:
            return f"{sems_str} семестр"
        return f"{sems_str} семестры"

    def generate(self) -> None:
        """
        Запускает процесс обработки и сохранения результатов.
        """
        if not self.excel_path.exists():
            print(f"Ошибка: Исходный файл {self.excel_path.name} не найден по указанному пути.")
            return

        try:
            wb = load_workbook(str(self.excel_path.absolute()), data_only=True)
        except Exception as e:
            logger.error(f"Не удалось открыть файл Excel: {e}")
            return

        # Интеллектуальный поиск листа учебного плана
        plan_sheet_name: str = ExcelConfig.PLAN_SHEET_NAME
        if plan_sheet_name not in wb.sheetnames:
            if "План" in wb.sheetnames:
                plan_sheet_name = "План"
            else:
                for name in wb.sheetnames:
                    if "план" in name.lower():
                        plan_sheet_name = name
                        break

        if ExcelConfig.COMPETENCY_SHEET_NAME not in wb.sheetnames or plan_sheet_name not in wb.sheetnames:
            logger.error(f"В книге отсутствуют необходимые листы. Доступные листы: {wb.sheetnames}")
            return

        print(f"Анализ структуры учебного плана на листе '{plan_sheet_name}'...")
        plan_db: Dict[str, Dict[str, Any]] = self._parse_plan_sheet(wb[plan_sheet_name])

        sheet_comp = wb[ExcelConfig.COMPETENCY_SHEET_NAME]
        competencies: List[Dict[str, Any]] = self._parse_competencies(sheet_comp)

        if not competencies:
            print("\n[ВНИМАНИЕ] Данные по компетенциям не были импортированы.")
            return

        total: int = len(competencies)
        print(f"\n{'№':<9} | {'Статус':<8} | {'Код':<8} | {'Компетенция'}")
        print("-" * 100)

        # Инициализация Word
        doc = Document()

        # Настройка и глобальное переопределение параметров стилей
        configure_document_styles(doc)
        style_h1, style_h2, style_normal = resolve_style_names(doc)

        # Настройка альбомной ориентации для широкой таблицы
        section = doc.sections[-1]
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height

        # 1. Добавление Раздела 1 перед таблицей (Стиль "Заголовок 1")
        p_header = doc.add_paragraph()
        try:
            p_header.style = style_h1
        except Exception:
            pass

        p_header.paragraph_format.space_before = Pt(12)
        p_header.paragraph_format.space_after = Pt(18)
        p_header.paragraph_format.keep_with_next = True  # Привязываем заголовок к таблице

        run_header = p_header.add_run(
            "Раздел 1. Матрица соответствия между компетенциями и дисциплинами и практиками их формирующими"
        )
        set_font(run_header, font_name="Times New Roman", size_pt=14, bold=True, italic=False)

        # 2. Создание таблицы
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"

        # Настройка шапки таблицы (дублирование на каждой странице + запрет разрывов)
        hdr_row = table.rows[0]
        set_repeat_table_header(hdr_row)
        set_row_cant_split(hdr_row)

        headers_text = [
            "Код и наименование компетенции",
            "Код и наименование индикатора достижения компетенции",
            "Дисциплины (модули)",
            "Практики",
            "Семестр формирования"
        ]

        hdr_cells = hdr_row.cells
        for idx, text in enumerate(headers_text):
            set_cell_text(
                hdr_cells[idx],
                text,
                style_name=style_normal,
                bold=True,
                italic=True,
                size_pt=12,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                keep_with_next=True  # Привязываем шапку к первой строке контента
            )

        # Заполнение строк таблицы данными
        for i, comp in enumerate(competencies, 1):
            comp_code: str = comp["comp_code"]
            status: str = "OK"

            try:
                new_row = table.add_row()
                row_cells = new_row.cells

                # Использование списков с последующим исключением дубликатов по названию предмета
                disciplines_names: List[str] = []
                practices_names: List[str] = []
                disc_semesters: Set[int] = set()
                prac_semesters: Set[int] = set()

                for item_code in sorted(list(comp["mapped_codes"])):
                    item_info: Optional[Dict[str, Any]] = plan_db.get(item_code)
                    if not item_info:
                        name: str = f"{item_code} (название не найдено)"
                        sems: List[int] = []
                    else:
                        name = item_info["name"]
                        sems = item_info["semesters"]

                    # Распределение на дисциплины и практики
                    if item_code.startswith("Б1") or item_code.startswith("ФТД"):
                        if name not in disciplines_names:
                            disciplines_names.append(name)
                        disc_semesters.update(sems)
                    elif item_code.startswith("Б2") or item_code.startswith("Б3"):
                        if name not in practices_names:
                            practices_names.append(name)
                        prac_semesters.update(sems)

                # Подготовка текстов для ячеек
                comp_text = f"{comp_code} {comp['comp_name']}"
                indicators_text = "\n".join(comp["indicators"])
                disc_text = "\n".join(sorted(disciplines_names))
                prac_text = "\n".join(sorted(practices_names))

                sem_info: List[str] = []
                if disc_semesters:
                    sem_info.append(f"Дисциплины — {self._format_semesters(list(disc_semesters))}")
                if prac_semesters:
                    sem_info.append(f"Практики — {self._format_semesters(list(prac_semesters))}")
                sem_text = "\n".join(sem_info)

                # Запись данных в ячейки с применением стиля "Обычный" (Normal) и переопределением на 12pt для компактности
                set_cell_text(row_cells[0], comp_text, style_name=style_normal, size_pt=12,
                              align=WD_ALIGN_PARAGRAPH.LEFT)
                set_cell_text(row_cells[1], indicators_text, style_name=style_normal, size_pt=12,
                              align=WD_ALIGN_PARAGRAPH.LEFT)
                set_cell_text(row_cells[2], disc_text, style_name=style_normal, size_pt=12,
                              align=WD_ALIGN_PARAGRAPH.LEFT)
                set_cell_text(row_cells[3], prac_text, style_name=style_normal, size_pt=12,
                              align=WD_ALIGN_PARAGRAPH.LEFT)
                set_cell_text(row_cells[4], sem_text, style_name=style_normal, size_pt=12,
                              align=WD_ALIGN_PARAGRAPH.LEFT)

            except Exception as e:
                status = "ERR"
                logger.error(f"Непредвиденная ошибка при обработке {comp_code}: {e}")

            comp_desc: str = comp['comp_name'][:50]
            print(f"[{i:03}/{total:03}] | {status:<8} | {comp_code:<8} | {comp_desc}...")

        try:
            # Создаем папку, если она не существует
            self.word_folder_path.mkdir(parents=True, exist_ok=True)

            # Конструируем итоговый путь к файлу
            final_docx_path = self.word_folder_path / "ИИ Оценочные материалы.docx"

            doc.save(str(final_docx_path.absolute()))
            print(
                f"\nПроцесс завершен. Таблица успешно сохранена в файл:\n'{final_docx_path.name}'\nПо адресу: {final_docx_path.parent}")
        except Exception as e:
            logger.error(f"Не удалось сохранить итоговый документ Word: {e}")


def main():
    print("=== Панель управления генерацией отчетов ===")

    user_excel_path: str = input("Шаг 1. Введите путь к исходному файлу Excel (например, plan.xlsx): ").strip()
    user_folder_path: str = input(
        "Шаг 2. Введите путь к папке для сохранения документа Word (например, C:\\Reports): ").strip()

    if not user_excel_path:
        user_excel_path = "plan.xlsx"
        print(f"Используется путь по умолчанию для Excel: {user_excel_path}")

    if not user_folder_path:
        user_folder_path = "."
        print(f"Используется текущая рабочая папка по умолчанию: {Path(user_folder_path).absolute()}")

    print("\nЗапуск процесса генерации...")
    generator = CompetencyReportGenerator(
        excel_path=user_excel_path,
        word_folder_path=user_folder_path
    )
    generator.generate()


if __name__ == "__main__":
    main()