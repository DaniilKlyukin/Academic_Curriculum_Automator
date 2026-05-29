import re
import logging
from pathlib import Path
from typing import List, Dict, Set, Tuple, Optional, Union, Any
from openpyxl import load_workbook
from openpyxl.worksheet.worksheet import Worksheet
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


class ExcelConfig:
    """Конфигурация столбцов и имен листов Excel по умолчанию."""
    COMPETENCY_SHEET_NAME: str = "Компетенции"
    PLAN_SHEET_NAME: str = "ПланСвод"  # Приоритетный лист по умолчанию

    # Дефолтная колонка для текста содержания (если не найдена автопоиском)
    COMP_TXT_COL: int = 5

    PLAN_IDX_COL: int = 2
    PLAN_NAME_COL: int = 3


def _parse_title_info(wb) -> Dict[str, str]:
    """Интеллектуальный поиск метаданных на листе 'Титул'."""
    info = {
        "code": "",
        "direction": "",
        "profile": "",
        "qualification": "",
        "form": "",
        "year": ""
    }
    if "Титул" not in wb.sheetnames:
        return info

    sheet = wb["Титул"]
    for r in range(1, min(sheet.max_row + 1, 100)):
        for c in range(1, min(sheet.max_column + 1, 20)):
            cell_val = sheet.cell(row=r, column=c).value
            if cell_val is None:
                continue
            val = clean_cell_value(cell_val)
            if not val:
                continue

            # 1. Поиск кода и названия направления (01.03.04 и т.д.)
            if re.match(r"^\d{2}\.\d{2}\.\d{2}$", val):
                info["code"] = val
                for offset in range(1, 4):
                    potential_dir = clean_cell_value(sheet.cell(row=r + offset, column=c).value)
                    if potential_dir and not potential_dir.lower().startswith(("профиль", "кафедра")):
                        info["direction"] = potential_dir
                        break

            # 2. Поиск профиля
            if val.lower().startswith("профиль:"):
                parts = val.split(":", 1)
                if len(parts) > 1 and parts[1].strip():
                    info["profile"] = parts[1].strip()
                else:
                    next_val = sheet.cell(row=r, column=c + 1).value
                    if next_val:
                        info["profile"] = str(next_val).strip()
            elif val.lower() == "профиль":
                next_val = sheet.cell(row=r, column=c + 1).value
                if next_val:
                    info["profile"] = str(next_val).strip()

            # 3. Поиск квалификации
            if val.lower().startswith("квалификация:"):
                parts = val.split(":", 1)
                if len(parts) > 1 and parts[1].strip():
                    info["qualification"] = parts[1].strip()
            elif val.lower().startswith("квалификация"):
                next_val = sheet.cell(row=r, column=c + 1).value
                if next_val:
                    info["qualification"] = str(next_val).strip()

            # 4. Поиск формы обучения
            if val.lower().startswith("форма обучения:"):
                parts = val.split(":", 1)
                if len(parts) > 1 and parts[1].strip():
                    info["form"] = parts[1].strip()
            elif val.lower().startswith("форма обучения"):
                next_val = sheet.cell(row=r, column=c + 1).value
                if next_val:
                    info["form"] = str(next_val).strip()

            # 5. Поиск года начала подготовки
            if val.lower().startswith("год начала подготовки"):
                for col_offset in range(1, 10):
                    temp_val = str(sheet.cell(row=r, column=c + col_offset).value or "").strip()
                    if temp_val.isdigit() and len(temp_val) == 4:
                        info["year"] = temp_val
                        break
    return info


def generate_title_page(doc: Document, info: Dict[str, str], style_normal: str):
    """Генерирует титульный лист в портретной ориентации."""
    # Приложение
    p_app = doc.add_paragraph()
    try:
        p_app.style = style_normal
    except Exception:
        pass
    p_app.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_app = p_app.add_run("Приложение")
    set_font(run_app, size_pt=14, bold=True)

    for _ in range(4):
        doc.add_paragraph()

    # Вспомогательная функция для добавления подчеркнутых полей
    def add_field(label: str, value: str, subtext: str = ""):
        p_lbl = doc.add_paragraph()
        try:
            p_lbl.style = style_normal
        except Exception:
            pass
        p_lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_lbl = p_lbl.add_run(label)
        set_font(run_lbl, size_pt=14, bold=True)

        p_val = doc.add_paragraph()
        try:
            p_val.style = style_normal
        except Exception:
            pass
        p_val.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_val = p_val.add_run(value if value else "__________________________________________________")
        set_font(run_val, size_pt=14, bold=False)
        run_val.underline = True

        if subtext:
            p_sub = doc.add_paragraph()
            try:
                p_sub.style = style_normal
            except Exception:
                pass
            p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_sub.paragraph_format.space_before = Pt(2)
            run_sub = p_sub.add_run(subtext)
            set_font(run_sub, size_pt=9, bold=False)

        doc.add_paragraph()

    # Заголовок
    p_title = doc.add_paragraph()
    try:
        p_title.style = style_normal
    except Exception:
        pass
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(12)
    run_title = p_title.add_run("Оценочные материалы образовательной программы\nвысшего образования")
    set_font(run_title, size_pt=14, bold=True)
    doc.add_paragraph()

    # Определение уровня образования по квалификации
    qual = info.get("qualification", "").lower()
    level_of_edu = "Бакалавриат"
    if "магистр" in qual:
        level_of_edu = "Магистратура"
    elif "специал" in qual:
        level_of_edu = "Специалитет"

    add_field("Уровень высшего образования", level_of_edu)

    dir_text = ""
    if info.get("code") or info.get("direction"):
        dir_text = f"{info.get('code', '')} {info.get('direction', '')}".strip()
    add_field("Направление подготовки (специальность)", dir_text,
              "код и наименование направления подготовки (специальности)")

    add_field("Направленность (профиль/программа/специализация)", info.get("profile", ""),
              "наименование направленности (профиля) подготовки (специализации)")
    add_field("Квалификация", info.get("qualification", ""))
    add_field("Форма обучения", info.get("form", ""), "очная, очно-заочная, заочная")
    add_field("Год приема", info.get("year", ""))


def generate_toc_page(doc: Document, style_normal: str):
    """Генерирует страницу содержания с динамическим полем оглавления MS Word."""
    p_title = doc.add_paragraph()
    try:
        p_title.style = style_normal
    except Exception:
        pass
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(12)
    run = p_title.add_run("Содержание")
    set_font(run, size_pt=14, bold=False)

    doc.add_paragraph()

    # Вставка поля оглавления Word (TOC)
    p_toc = doc.add_paragraph()
    try:
        p_toc.style = style_normal
    except Exception:
        pass
    run_toc = p_toc.add_run()

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    run_toc._r.append(fldChar1)
    run_toc._r.append(instrText)
    run_toc._r.append(fldChar2)
    run_toc._r.append(fldChar3)


def set_font(run, font_name="Times New Roman", size_pt=12, bold=False, italic=False, color_rgb=(0, 0, 0)):
    """Устанавливает шрифт, размер, цвет и начертание для текстового прогона (run),

    обеспечивая поддержку кириллицы в MS Word.
    """
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(*color_rgb)

    try:
        rPr = run._r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:cs'), font_name)
        rFonts.set(qn('w:eastAsia'), font_name)
        rPr.append(rFonts)
    except Exception as e:
        logger.debug(f"Не удалось применить XML-стили шрифтов: {e}")


def resolve_style_names(doc: Document):
    """Определяет доступные системные имена стилей в открытом документе."""
    existing_styles = {s.name for s in doc.styles}

    style_h1 = "Heading 1"
    for name in ["Заголовок 1", "Heading 1", "Heading1"]:
        if name in existing_styles:
            style_h1 = name
            break

    style_h2 = "Heading 2"
    for name in ["Заголовок 2", "Heading 2", "Heading2"]:
        if name in existing_styles:
            style_h2 = name
            break

    style_normal = "Normal"
    for name in ["Обычный", "Normal"]:
        if name in existing_styles:
            style_normal = name
            break

    return style_h1, style_h2, style_normal


def clean_cell_value(val: Any) -> str:
    """Удаляет из значения ячейки Excel скрытые артефакты разметки и переносов строк."""
    if val is None:
        return ""
    s = str(val)
    # Удаляем артефакты Excel _x000D_ и системные переносы каретки \r
    s = s.replace("_x000D_", "").replace("_x000D", "").replace("\r", "")
    return s.strip()


def force_xml_update_on_open(doc: Document):
    """Добавляет в настройки документа команду на автоматическое обновление полей (включая оглавление) при открытии."""
    try:
        settings_xml = doc.settings._element
        updateFields = OxmlElement('w:updateFields')
        updateFields.set(qn('w:val'), 'true')
        settings_xml.append(updateFields)
    except Exception as e:
        logger.debug(f"Не удалось включить автообновление полей на старте: {e}")


def set_style_font_robust(style_obj, font_name="Times New Roman", size_pt=14, bold=False, italic=False,
                          color_rgb=(0, 0, 0)):
    """Устанавливает шрифт для стиля с гарантированной поддержкой кириллицы

    и полным отключением переопределения системных тем Word (Calibri/Arial).
    """
    font = style_obj.font
    font.name = font_name
    font.size = Pt(size_pt)
    font.bold = bold
    font.italic = italic
    font.color.rgb = RGBColor(*color_rgb)

    try:
        # Переходим к XML-свойствам символов стиля (rPr)
        rPr = style_obj.element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()

        # КРИТИЧЕСКИ ВАЖНО: Удаляем привязку к теме оформления Word (+minor / +major).
        # Если этого не сделать, Word принудительно вернет Calibri для кириллицы.
        theme_attrs = ['asciiTheme', 'hAnsiTheme', 'eastAsiaTheme', 'cstheme']
        for attr in theme_attrs:
            attrib_key = qn(f'w:{attr}')
            if attrib_key in rFonts.attrib:
                del rFonts.attrib[attrib_key]

        # Явно закрепляем Times New Roman для всех наборов символов
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:cs'), font_name)
        rFonts.set(qn('w:eastAsia'), font_name)
    except Exception as e:
        logger.debug(f"Не удалось применить глубокую очистку стилей шрифтов: {e}")


def get_or_create_style(doc: Document, name: str, style_type=WD_STYLE_TYPE.PARAGRAPH):
    """Возвращает существующий стиль или принудительно создает его в XML-структуре, если он латентный."""
    try:
        return doc.styles[name]
    except KeyError:
        try:
            return doc.styles.add_style(name, style_type)
        except Exception:
            # Резервный поиск по имени без учета регистра
            for s in doc.styles:
                if s.name.lower() == name.lower():
                    return s
    return None


def get_or_create_toc_style(doc: Document, level: int):
    """Находит существующий или создает и корректно регистрирует встроенный стиль оглавления (TOC)

    для того, чтобы MS Word гарантированно применял его при обновлении полей.
    """
    english_name = f"toc {level}"
    russian_name = f"Оглавление {level}"
    style_name_to_add = f"TOC {level}"

    # Сначала пытаемся найти стиль по любому из возможных названий
    for s in doc.styles:
        if s.name.lower() in [english_name.lower(), russian_name.lower(), style_name_to_add.lower()]:
            return s

    # Если стиль отсутствует в документе, создаем его временную заготовку
    try:
        style = doc.styles.add_style(style_name_to_add, WD_STYLE_TYPE.PARAGRAPH)
    except Exception:
        try:
            return doc.styles[style_name_to_add]
        except Exception:
            return doc.styles[russian_name]

    # Корректируем XML-структуру стиля, превращая его из пользовательского во встроенный
    try:
        style_el = style.element
        # Задаем системный ID, который ожидает генератор Word ('toc 1', 'toc 2' и т.д.)
        style_el.set(qn('w:styleId'), english_name)
        # Снимаем флаг пользовательского стиля (0 означает "встроенный")
        style_el.set(qn('w:customStyle'), '0')

        # Задаем внутреннее имя стиля
        name_el = style_el.find(qn('w:name'))
        if name_el is not None:
            name_el.set(qn('w:val'), english_name)

        # Добавляем локализованный псевдоним (alias) для поддержки русскоязычной версии Word
        aliases_el = style_el.find(qn('w:aliases'))
        if aliases_el is None:
            aliases_el = OxmlElement('w:aliases')
            aliases_el.set(qn('w:val'), russian_name)
            if name_el is not None:
                idx = style_el.index(name_el)
                style_el.insert(idx + 1, aliases_el)
            else:
                style_el.insert(0, aliases_el)
        else:
            aliases_el.set(qn('w:val'), russian_name)
    except Exception as e:
        logger.debug(f"Не удалось применить XML-патч для стиля оглавления: {e}")

    return style


def configure_toc_styles(doc: Document):
    """Настраивает форматирование строк оглавления в точном соответствии с макетом."""
    black_color = (0, 0, 0)

    # Получаем/регистрируем системные стили оглавления
    style_toc1 = get_or_create_toc_style(doc, 1)
    style_toc2 = get_or_create_toc_style(doc, 2)

    # 1. Настройка Оглавление 1 (Разделы: Жирный, 14pt, выравнивание по ширине, интервалы 6/6 пт)
    if style_toc1:
        set_style_font_robust(
            style_toc1,
            font_name="Times New Roman",
            size_pt=14,
            bold=True,
            italic=False,  # Отключено (только жирный, как на первом скриншоте)
            color_rgb=black_color
        )
        pf1 = style_toc1.paragraph_format
        pf1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Выравнивание по ширине
        pf1.space_before = Pt(6)                    # Интервал перед: 6 пт
        pf1.space_after = Pt(6)                     # Интервал после: 6 пт
        pf1.line_spacing = 1.0                      # Одинарный межстрочный интервал
        pf1.left_indent = Cm(0)                     # Без отступа слева

    # 2. Настройка Оглавление 2 (Подразделы: Курсив, 13pt, выравнивание по ширине, отступ 0.39 см, интервалы 3/3 пт)
    if style_toc2:
        set_style_font_robust(
            style_toc2,
            font_name="Times New Roman",
            size_pt=13,
            bold=False,                                 # Не жирный
            italic=True,                                # Курсив (как на втором скриншоте)
            color_rgb=black_color
        )
        pf2 = style_toc2.paragraph_format
        pf2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Выравнивание по ширине
        pf2.left_indent = Cm(0.39)                  # Отступ слева: 0.39 см
        pf2.space_before = Pt(3)                    # Интервал перед: 3 пт
        pf2.space_after = Pt(3)                     # Интервал после: 3 пт
        pf2.line_spacing = 1.0                      # Одинарный межстрочный интервал


def configure_document_styles(doc: Document):
    """Программно настраивает глобальные стили документа согласно заданным параметрам."""
    style_h1, style_h2, style_normal = resolve_style_names(doc)
    black_color = (0, 0, 0)

    def setup_style(style_obj, size_pt=14, bold=False, italic=False, align=None):
        # Применяем надежную настройку шрифта стиля (с поддержкой кириллицы)
        set_style_font_robust(style_obj, font_name="Times New Roman", size_pt=size_pt, bold=bold, italic=italic,
                              color_rgb=black_color)

        # Настройка абзацных интервалов и выравнивания
        pf = style_obj.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0
        if align is not None:
            pf.alignment = align

    # Настройка стиля «Обычный» (Обычный текст по ширине)
    if style_normal in doc.styles:
        setup_style(doc.styles[style_normal], size_pt=14, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # Настройка стиля «Заголовок 1» (По центру, Жирный)
    if style_h1 in doc.styles:
        setup_style(doc.styles[style_h1], size_pt=14, bold=True, italic=False, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Настройка стиля «Заголовок 2» (По ширине, Курсив)
    if style_h2 in doc.styles:
        setup_style(doc.styles[style_h2], size_pt=14, bold=False, italic=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # Настройка стилей отображения элементов автооглавления
    configure_toc_styles(doc)

def set_repeat_table_header(row):
    """Включает повторение строки таблицы на каждой новой странице (tblHeader)."""
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    trPr.append(tblHeader)


def set_row_cant_split(row):
    """Запрещает разрыв строки таблицы при переходе на новую страницу (cantSplit)."""
    trPr = row._tr.get_or_add_trPr()
    cantSplit = OxmlElement('w:cantSplit')
    trPr.append(cantSplit)


def set_cell_text(cell, text: str, style_name="Normal", bold=False, italic=False, size_pt=12,
                  align=WD_ALIGN_PARAGRAPH.LEFT, keep_with_next=False):
    """Записывает текст в ячейку таблицы с очисткой и применением форматирования."""
    cell.text = ""  # Сброс стандартного содержимого ячейки
    p = cell.paragraphs[0]
    try:
        p.style = style_name
    except Exception:
        pass

    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0

    if keep_with_next:
        p.paragraph_format.keep_with_next = True

    if text:
        run = p.add_run(text)
        set_font(run, size_pt=size_pt, bold=bold, italic=italic)


class CompetencyReportGenerator:
    """Класс для сопоставления компетенций и дисциплин из учебного плана в документ Word."""

    def __init__(self, excel_path: Union[str, Path], word_file_path: Union[str, Path]) -> None:
        self.excel_path: Path = Path(excel_path)
        self.word_file_path: Path = Path(word_file_path)

    def _find_semester_columns(self, sheet: Worksheet) -> Dict[int, int]:
        """
        Сканирует первые строки листа для автоматического определения столбцов семестров (18-25).
        """
        semester_cols: Dict[int, int] = {}
        for col in range(1, sheet.max_column + 1):
            for row in range(1, 6):
                cell_val: str = str(sheet.cell(row=row, column=col).value or "").lower()
                if "сем" in cell_val:
                    match = re.search(r"сем.*?(\d+)", cell_val)
                    if match:
                        sem_num: int = int(match.group(1))
                        semester_cols[sem_num] = col
                        break
        return semester_cols

    def _find_assessment_columns(self, sheet: Worksheet) -> List[int]:
        """
        Находит индексы колонок промежуточной аттестации (Экзамен, Зачет, Зачет с оц., КР, Реферат).
        """
        cols: List[int] = []
        for col in range(1, sheet.max_column + 1):
            for row in range(1, 6):
                val: str = str(sheet.cell(row=row, column=col).value or "").lower()
                if any(x in val for x in ["экзамен", "зачет", "кр", "реферат", "кп"]):
                    cols.append(col)
                    break
        return list(set(cols))

    def _find_competency_text_column(self, sheet: Worksheet) -> int:
        """
        Ищет столбец 'Содержание' на листе компетенций.
        """
        txt_col: int = ExcelConfig.COMP_TXT_COL
        for col in range(1, sheet.max_column + 1):
            for row in range(1, 11):
                val: str = str(sheet.cell(row=row, column=col).value or "").lower().strip()
                if "содержание" in val:
                    return col
        return txt_col

    @staticmethod
    def _extract_semesters_from_assessment(val: Any) -> List[int]:
        """
        Извлекает номера семестров из ячеек промежуточной аттестации (например, '123' -> [1, 2, 3]).
        """
        if val is None:
            return []
        val_str: str = str(val).strip()
        semesters: List[int] = []

        if val_str.isdigit():
            for char in val_str:
                sem_num = int(char)
                if 1 <= sem_num <= 12:
                    semesters.append(sem_num)
        else:
            parts = re.split(r'[,\s;]+', val_str)
            for part in parts:
                if part.isdigit():
                    sem_num = int(part)
                    if 1 <= sem_num <= 12:
                        semesters.append(sem_num)
        return list(set(semesters))

    def _parse_plan_sheet(self, sheet: Worksheet) -> Dict[str, Dict[str, Any]]:
        """
        Парсит учебный план, собирая названия предметов и семестры их проведения.
        """
        plan_db: Dict[str, Dict[str, Any]] = {}
        semester_cols: Dict[int, int] = self._find_semester_columns(sheet)
        assessment_cols: List[int] = self._find_assessment_columns(sheet)

        logger.info(f"Найдены столбцы семестров: {list(semester_cols.keys())}")
        logger.info(f"Найдены столбцы аттестации (индексы): {assessment_cols}")

        for row in range(1, sheet.max_row + 1):
            idx_val: str = str(sheet.cell(row=row, column=ExcelConfig.PLAN_IDX_COL).value or "").strip()
            if re.match(r"^(Б\d|ФТД)", idx_val):
                name_val: str = str(sheet.cell(row=row, column=ExcelConfig.PLAN_NAME_COL).value or "").strip()

                active_semesters: List[int] = []

                # 1. Считываем семестры из колонок нагрузки
                for sem_num, col_idx in semester_cols.items():
                    val: Optional[Any] = sheet.cell(row=row, column=col_idx).value
                    if val is not None:
                        try:
                            if float(val) > 0:
                                active_semesters.append(sem_num)
                        except ValueError:
                            pass

                # 2. Дополнительно считываем семестры из колонок аттестации
                for col_idx in assessment_cols:
                    val_assess: Optional[Any] = sheet.cell(row=row, column=col_idx).value
                    if val_assess is not None:
                        active_semesters.extend(self._extract_semesters_from_assessment(val_assess))

                plan_db[idx_val] = {
                    "name": name_val,
                    "semesters": sorted(list(set(active_semesters)))
                }
        return plan_db

    def _parse_competencies(self, sheet: Worksheet) -> List[Dict[str, Any]]:
        """
        Парсит структуру компетенций, динамически определяя смещение структуры (для ПК).
        """
        competencies: List[Dict[str, Any]] = []
        current_comp: Optional[Dict[str, Any]] = None

        txt_col = self._find_competency_text_column(sheet)
        logger.info(f"Определен столбец Содержания компетенций -> {txt_col}")

        # Помехоустойчивые паттерны к разным типам тире/дефисов и пробелов
        comp_pattern = re.compile(r"^([A-Za-zА-Яа-я]+)\s*[-–—]\s*\d+$")
        indicator_pattern = re.compile(r"^([A-Za-zА-Яа-я]+)\s*[-–—]\s*\d+\.\d+$")
        item_pattern = re.compile(r"^(Б\d|ФТД)")

        print("\n" + "=" * 50)
        print("ДЕТАЛЬНЫЙ ЛОГ ИМПОРТА ИЗЛИСТА 'КОМПЕТЕНЦИИ' (АДАПТИВНЫЙ ПАРСЕР):")
        print("=" * 50)

        # Вычисляем глубину сканирования с запасом
        max_scan_rows: int = max(sheet.max_row, 1000)
        consecutive_empty_rows: int = 0

        for row in range(1, max_scan_rows + 1):
            # Считываем 4 колонки из-за возможного смещения ПК вправо
            val_col1: str = str(sheet.cell(row=row, column=1).value or "").strip()
            val_col2: str = str(sheet.cell(row=row, column=2).value or "").strip()
            val_col3: str = str(sheet.cell(row=row, column=3).value or "").strip()
            val_col4: str = str(sheet.cell(row=row, column=4).value or "").strip()
            text_val: str = str(sheet.cell(row=row, column=txt_col).value or "").strip()

            # Предотвращение преждевременного завершения
            if not val_col1 and not val_col2 and not val_col3 and not val_col4:
                consecutive_empty_rows += 1
                if consecutive_empty_rows > 50:  # Прекращаем парсинг после 50 пустых строк подряд
                    break
                continue

            consecutive_empty_rows = 0  # Сбрасываем счетчик, если строка не пуста

            # 1. Поиск КОМПЕТЕНЦИИ (может быть в колонке 1 или 2)
            comp_code: Optional[str] = None
            if val_col1 and comp_pattern.match(val_col1):
                comp_code = val_col1
            elif val_col2 and comp_pattern.match(val_col2):
                comp_code = val_col2

            if comp_code:
                current_comp = {
                    "comp_code": comp_code,
                    "comp_name": text_val,
                    "indicators": [],
                    "mapped_codes": set()
                }
                competencies.append(current_comp)
                print(f"Строка {row:03} | [Компетенция] {comp_code} -> {text_val[:50]}...")
                continue

            # 2. Поиск ИНДИКАТОРА (может быть в колонке 2 или 3)
            indicator_code: Optional[str] = None
            if val_col2 and indicator_pattern.match(val_col2):
                indicator_code = val_col2
            elif val_col3 and indicator_pattern.match(val_col3):
                indicator_code = val_col3

            if indicator_code:
                if current_comp:
                    current_comp["indicators"].append(f"{indicator_code} {text_val}")
                    print(f"Строка {row:03} |   [Индикатор]   {indicator_code} -> {text_val[:50]}...")
                else:
                    print(f"Строка {row:03} |   [ВНИМАНИЕ] Пропущен индикатор без компетенции: {indicator_code}")
                continue

            # 3. Поиск ДИСЦИПЛИНЫ/ПРАКТИКИ (может быть в колонке 3 или 4)
            item_code: Optional[str] = None
            if val_col3 and item_pattern.match(val_col3):
                item_code = val_col3
            elif val_col4 and item_pattern.match(val_col4):
                item_code = val_col4

            if item_code:
                if current_comp:
                    current_comp["mapped_codes"].add(item_code)
                    print(f"Строка {row:03} |     [Привязка]    {item_code} -> {text_val[:50]}...")
                else:
                    print(f"Строка {row:03} |     [ВНИМАНИЕ] Пропущена дисциплина без компетенции: {item_code}")

        print("=" * 50 + "\n")
        return competencies

    @staticmethod
    def _format_semesters(sem_list: List[int]) -> str:
        """Форматирует список семестров в строку."""
        if not sem_list:
            return ""
        sorted_sems: List[int] = sorted(list(set(sem_list)))
        sems_str: str = ", ".join(map(str, sorted_sems))
        if len(sorted_sems) == 1:
            return f"{sems_str} семестр"
        return f"{sems_str} семестры"

    def generate(self) -> None:
        """
        Запускает процесс обработки и сохранения результатов.
        """
        if not self.excel_path.exists():
            print(f"Ошибка: Исходный файл {self.excel_path.name} не найден по указанному пути.")
            return

        try:
            wb = load_workbook(str(self.excel_path.absolute()), data_only=True)
        except Exception as e:
            logger.error(f"Не удалось открыть файл Excel: {e}")
            return

        # Интеллектуальный поиск листа учебного плана
        plan_sheet_name: str = ExcelConfig.PLAN_SHEET_NAME
        if plan_sheet_name not in wb.sheetnames:
            if "План" in wb.sheetnames:
                plan_sheet_name = "План"
            else:
                for name in wb.sheetnames:
                    if "план" in name.lower():
                        plan_sheet_name = name
                        break

        if ExcelConfig.COMPETENCY_SHEET_NAME not in wb.sheetnames or plan_sheet_name not in wb.sheetnames:
            logger.error(f"В книге отсутствуют необходимые листы. Доступные листы: {wb.sheetnames}")
            return

        print(f"Анализ структуры учебного плана на листе '{plan_sheet_name}'...")
        plan_db: Dict[str, Dict[str, Any]] = self._parse_plan_sheet(wb[plan_sheet_name])

        # Парсинг информации для титульного листа
        title_info: Dict[str, str] = _parse_title_info(wb)

        sheet_comp = wb[ExcelConfig.COMPETENCY_SHEET_NAME]
        competencies: List[Dict[str, Any]] = self._parse_competencies(sheet_comp)

        if not competencies:
            print("\n[ВНИМАНИЕ] Данные по компетенциям не были импортированы.")
            return

        total: int = len(competencies)
        print(f"\n{'№':<9} | {'Статус':<8} | {'Код':<8} | {'Компетенция'}")
        print("-" * 100)

        # Инициализация Word
        doc = Document()

        force_xml_update_on_open(doc)

        sect0 = doc.sections[0]
        sect0.top_margin = Cm(2.0)
        sect0.bottom_margin = Cm(2.0)
        sect0.left_margin = Cm(3.0)
        sect0.right_margin = Cm(1.5)

        # Настройка и глобальное переопределение параметров стилей
        configure_document_styles(doc)
        style_h1, style_h2, style_normal = resolve_style_names(doc)

        # 1. Первая страница: Титульный лист (ориентация Книжная по умолчанию)
        generate_title_page(doc, title_info, style_normal)

        # 2. Вторая страница: Содержание (ориентация Книжная)
        doc.add_page_break()
        generate_toc_page(doc, style_normal)

        # 3. Третья страница: Создание новой секции с Альбомной ориентацией для таблицы
        table_section = doc.add_section()
        table_section.orientation = WD_ORIENT.LANDSCAPE
        new_width, new_height = table_section.page_height, table_section.page_width
        table_section.page_width = new_width
        table_section.page_height = new_height

        table_section.top_margin = Cm(3.0)
        table_section.bottom_margin = Cm(1.5)
        table_section.left_margin = Cm(2.0)
        table_section.right_margin = Cm(2.0)

        # Добавление Раздела 1 перед таблицей (Стиль "Заголовок 1")
        p_header = doc.add_paragraph()
        try:
            p_header.style = style_h1
        except Exception:
            pass

        p_header.paragraph_format.space_before = Pt(12)
        p_header.paragraph_format.space_after = Pt(18)
        p_header.paragraph_format.keep_with_next = True  # Привязываем заголовок к таблице

        p_header.add_run(
            "Раздел 1. Матрица соответствия между компетенциями и дисциплинами и практиками их формирующими"
        )

        # 4. Создание таблицы
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"

        # Настройка шапки таблицы (дублирование на каждой странице + запрет разрывов)
        hdr_row = table.rows[0]
        set_repeat_table_header(hdr_row)
        set_row_cant_split(hdr_row)

        headers_text = [
            "Код и наименование компетенции",
            "Код и наименование индикатора достижения компетенции",
            "Дисциплины (модули)",
            "Практики",
            "Семестр формирования"
        ]

        hdr_cells = hdr_row.cells
        for idx, text in enumerate(headers_text):
            set_cell_text(
                hdr_cells[idx],
                text,
                style_name=style_normal,
                bold=True,
                italic=True,
                size_pt=12,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                keep_with_next=True  # Привязываем шапку к первой строке контента
            )

        # Заполнение строк таблицы данными
        for i, comp in enumerate(competencies, 1):
            comp_code: str = comp["comp_code"]
            status: str = "OK"

            try:
                new_row = table.add_row()
                row_cells = new_row.cells

                # Использование списков с последующим исключением дубликатов по названию предмета
                disciplines_names: List[str] = []
                practices_names: List[str] = []
                disc_semesters: Set[int] = set()
                prac_semesters: Set[int] = set()

                for item_code in sorted(list(comp["mapped_codes"])):
                    item_info: Optional[Dict[str, Any]] = plan_db.get(item_code)
                    if not item_info:
                        name: str = f"{item_code} (название не найдено)"
                        sems: List[int] = []
                    else:
                        name = item_info["name"]
                        sems = item_info["semesters"]

                    # Распределение на дисциплины и практики
                    if item_code.startswith("Б1") or item_code.startswith("ФТД"):
                        if name not in disciplines_names:
                            disciplines_names.append(name)
                        disc_semesters.update(sems)
                    elif item_code.startswith("Б2") or item_code.startswith("Б3"):
                        if name not in practices_names:
                            practices_names.append(name)
                        prac_semesters.update(sems)

                # Подготовка текстов для ячеек
                comp_text = f"{comp_code} {comp['comp_name']}"
                indicators_text = "\n".join(comp["indicators"])
                disc_text = "\n".join(sorted(disciplines_names))
                prac_text = "\n".join(sorted(practices_names))

                sem_info: List[str] = []
                if disc_semesters:
                    sem_info.append(f"Дисциплины — {self._format_semesters(list(disc_semesters))}")
                if prac_semesters:
                    sem_info.append(f"Практики — {self._format_semesters(list(prac_semesters))}")
                sem_text = "\n".join(sem_info)

                # Запись данных в ячейки с применением стиля "Обычный" (Normal) и переопределением на 12pt для компактности
                set_cell_text(row_cells[0], comp_text, style_name=style_normal, size_pt=12,
                              align=WD_ALIGN_PARAGRAPH.LEFT)
                set_cell_text(row_cells[1], indicators_text, style_name=style_normal, size_pt=12,
                              align=WD_ALIGN_PARAGRAPH.LEFT)
                set_cell_text(row_cells[2], disc_text, style_name=style_normal, size_pt=12,
                              align=WD_ALIGN_PARAGRAPH.LEFT)
                set_cell_text(row_cells[3], prac_text, style_name=style_normal, size_pt=12,
                              align=WD_ALIGN_PARAGRAPH.LEFT)
                set_cell_text(row_cells[4], sem_text, style_name=style_normal, size_pt=12,
                              align=WD_ALIGN_PARAGRAPH.LEFT)

            except Exception as e:
                status = "ERR"
                logger.error(f"Непредвиденная ошибка при обработке {comp_code}: {e}")

            comp_desc: str = comp['comp_name'][:50]
            print(f"[{i:03}/{total:03}] | {status:<8} | {comp_code:<8} | {comp_desc}...")

        try:
            # Создаем папку, если она не существует
            self.word_file_path.parent.mkdir(parents=True, exist_ok=True)

            doc.save(str(self.word_file_path.absolute()))
            print(
                f"\nПроцесс завершен. Таблица успешно сохранена в файл:\n'{self.word_file_path.name}'\nПо адресу: {self.word_file_path.parent}")
        except Exception as e:
            logger.error(f"Не удалось сохранить итоговый документ Word: {e}")


def main():
    print("=== Панель управления генерацией отчетов ===")

    user_excel_path: str = input("Шаг 1. Введите путь к исходному файлу Excel (например, plan.xlsx): ").strip()
    user_folder_path: str = input(
        "Шаг 2. Введите путь к папке для сохранения документа Word (например, C:\\Reports): ").strip()

    if not user_excel_path:
        user_excel_path = "plan.xlsx"
        print(f"Используется путь по умолчанию для Excel: {user_excel_path}")

    if not user_folder_path:
        user_folder_path = ".."
        print(f"Используется текущая рабочая папка по умолчанию: {Path(user_folder_path).absolute()}")

    print("\nЗапуск процесса генерации...")
    file_path = Path(user_folder_path) / "Оценочные материалы.docx"

    generator = CompetencyReportGenerator(
        excel_path=Path(user_excel_path),
        word_file_path=file_path
    )
    generator.generate()


if __name__ == "__main__":
    main()