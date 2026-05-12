import os
import tempfile
import logging
from typing import List, Optional, Generator, Union, Any, Type

import fitz
from docx import Document
from docx.shared import Mm
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, nsmap
from docx.enum.text import WD_BREAK
from docx.document import Document as DocumentType
from docx.text.paragraph import Paragraph
from docx.shape import InlineShape

EXT_NAMESPACES: dict[str, str] = {
    'wp14': 'http://schemas.microsoft.com/office/word/2010/wordml',
    'w14': 'http://schemas.microsoft.com/office/word/2010/wordml',
    'w15': 'http://schemas.microsoft.com/office/word/2012/wordml',
}
nsmap.update(EXT_NAMESPACES)

logger = logging.getLogger(__name__)


class DocxEditor:
    """Класс для редактирования документов DOCX: вставка сканов, управление позиционированием."""

    def __init__(self, file_path: str) -> None:
        self.file_path: str = file_path
        self.doc: Optional[DocumentType] = None
        self._temp_images: List[str] = []

    def __enter__(self) -> 'DocxEditor':
        self.doc = Document(self.file_path)
        return self

    def __exit__(self, exc_type: Optional[Type[BaseException]],
                 exc_val: Optional[BaseException],
                 exc_tb: Any) -> None:
        if self.doc and exc_type is None:
            self.doc.save(self.file_path)
        self._cleanup_temp_files()

    def _cleanup_temp_files(self) -> None:
        """Удаляет временные файлы изображений, созданные при конвертации PDF."""
        for tmp in self._temp_images:
            try:
                if os.path.exists(tmp):
                    os.remove(tmp)
            except Exception as e:
                logger.debug(f"Ошибка при очистке временного файла {tmp}: {e}")

    def _prepare_image(self, path: str) -> str:
        """
        Если файл PDF — конвертирует первую страницу в PNG.
        Если нет — возвращает исходный путь.
        """
        if not path.lower().endswith('.pdf'):
            return path
        try:
            pdf_doc = fitz.open(path)
            page = pdf_doc.load_page(0)
            pix = page.get_pixmap(dpi=300)

            fd, tmp_path = tempfile.mkstemp(suffix=".png")
            os.close(fd)
            pix.save(tmp_path)
            pdf_doc.close()
            self._temp_images.append(tmp_path)
            return tmp_path
        except Exception as e:
            logger.error(f"Ошибка конвертации PDF в изображение: {e}")
            return path

    def _make_floating(self, shape: InlineShape, width_mm: float, height_mm: float) -> None:
        """
        Превращает inline-картинку в плавающую (Anchor) с позиционированием 'Над текстом'.
        """
        w_emu = int(width_mm * 36000)
        h_emu = int(height_mm * 36000)

        inline = shape._inline

        anchor_xml = (
            f'<wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" '
            f'relativeHeight="251658240" behindDoc="0" locked="0" layoutInCell="1" '
            f'allowOverlap="1" {nsdecls("wp", "wp14", "pic", "r")}>'
            f'<wp:simplePos x="0" y="0"/>'
            f'<wp:positionH relativeFrom="page"><wp:posOffset>0</wp:posOffset></wp:positionH>'
            f'<wp:positionV relativeFrom="page"><wp:posOffset>0</wp:posOffset></wp:positionV>'
            f'<wp:extent cx="{w_emu}" cy="{h_emu}"/>'
            f'<wp:effectExtent l="0" t="0" r="0" b="0"/>'
            f'<wp:wrapNone/><wp:docPr id="1" name="Scan"/><wp:cNvGraphicFramePr/>'
            f'{inline.graphic.xml}'
            f'</wp:anchor>'
        )
        inline.getparent().replace(inline, parse_xml(anchor_xml))

    def _get_paragraph_by_page(self, page_num: int) -> Paragraph:
        """
        Находит (или создает) параграф на указанной странице.
        В DOCX нет жесткого понятия 'страница', поэтому ищем разрывы или создаем их.
        """
        if not self.doc:
            raise RuntimeError("Документ не загружен. Используйте 'with DocxEditor(...)'.")

        if not self.doc.paragraphs:
            self.doc.add_paragraph()

        if page_num <= 1:
            return self.doc.paragraphs[0]

        current_page = 1
        for i, p in enumerate(self.doc.paragraphs):
            p_xml = p._element.xml
            if 'w:br w:type="page"' in p_xml or 'lastRenderedPageBreak' in p_xml:
                current_page += 1
                if current_page == page_num:
                    if i + 1 < len(self.doc.paragraphs):
                        return self.doc.paragraphs[i + 1]
                    else:
                        return self.doc.add_paragraph()

        last_p = self.doc.paragraphs[-1]
        while current_page < page_num:
            last_p.add_run().add_break(WD_BREAK.PAGE)
            last_p = self.doc.add_paragraph()
            current_page += 1
        return last_p

    def add_scan_to_page(self, page_num: int, image_path: str, floating: bool = True) -> None:
        """
        Вставляет изображение на конкретную страницу.

        Args:
            page_num: Номер страницы (1-based).
            image_path: Путь к файлу (картинка или PDF).
            floating: Если True, картинка будет размещена поверх текста на всю страницу.
        """
        p = self._get_paragraph_by_page(page_num)
        img_src = self._prepare_image(image_path)
        run = p.add_run()
        shape: InlineShape = run.add_picture(img_src, width=Mm(210))

        if floating:
            self._make_floating(shape, width_mm=210, height_mm=297)

    def insert_image_after_text(self, text: str, image_path: str) -> bool:
        """
        Ищет текст в документе и вставляет скан ПОВЕРХ страницы, на которой найден текст.

        Returns:
            bool: True если текст найден и изображение вставлено, иначе False.
        """
        if not self.doc:
            return False

        target = text.lower().strip()

        def all_paras() -> Generator[Paragraph, None, None]:
            """Генератор для обхода всех параграфов, включая таблицы."""
            if not self.doc:
                return
            for p in self.doc.paragraphs:
                yield p
            for t in self.doc.tables:
                for r in t.rows:
                    for c in r.cells:
                        for p in c.paragraphs:
                            yield p

        for p in all_paras():
            if target in p.text.lower():
                img_src = self._prepare_image(image_path)
                run = p.add_run()
                shape: InlineShape = run.add_picture(img_src, width=Mm(210))
                self._make_floating(shape, width_mm=210, height_mm=297)
                return True
        return False