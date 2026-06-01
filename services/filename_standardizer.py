import os
import re
import logging
from pathlib import Path
from typing import Dict, Tuple, Any, List
from openpyxl import load_workbook
from docx import Document

# Настройка логирования
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


def normalize_homoglyphs(text: str) -> str:
    """Заменяет английские буквы-омоглифы на их русские аналоги."""
    eng_to_rus = {
        'a': 'а', 'b': 'в', 'c': 'с', 'e': 'е', 'h': 'н', 'k': 'к', 'm': 'м',
        'o': 'о', 'p': 'р', 't': 'т', 'x': 'х', 'y': 'у', 'ё': 'е'
    }
    return "".join(eng_to_rus.get(char, char) for char in text.lower())


def clean_name_for_match(name: str) -> str:
    """Очищает строку от кавычек, пробелов и регистра с устранением омоглифов."""
    if not name:
        return ""
    s = name.lower()
    s = normalize_homoglyphs(s)
    s = "".join(s.split())
    s = s.replace("«", "").replace("»", "").replace('"', '').replace("'", "")
    return s


def abbreviate_word(word: str) -> str:
    """Сокращает слово до первой согласной после первой гласной с поддержкой спецсимволов."""
    vowels = set("аеёиоуыэюя")
    word_clean = "".join([c for c in word if c.isalnum() or c in ('+', '#')])
    if not word_clean:
        return ""

    first_vowel_idx = -1
    for idx, char in enumerate(word_clean):
        if char.lower() in vowels:
            first_vowel_idx = idx
            break

    if first_vowel_idx == -1:
        return word_clean

    first_consonant_after_vowel_idx = -1
    for idx in range(first_vowel_idx + 1, len(word_clean)):
        if word_clean[idx].lower() not in vowels:
            first_consonant_after_vowel_idx = idx
            break

    if first_consonant_after_vowel_idx == -1:
        return word_clean

    return word_clean[:first_consonant_after_vowel_idx + 1]


def abbreviate_discipline(name: str) -> str:
    """Сокращает название дисциплины (идентично для сканов и DOCX)."""
    normalized_name = name.replace("-", " ")
    words = normalized_name.split()

    stop_words = {"на", "по", "в", "с", "для", "под", "о", "об", "за", "из", "от", "до", "без"}

    abbr_words = []
    for w in words:
        if w.lower() in stop_words:
            continue

        abbr = abbreviate_word(w)
        if not abbr:
            continue

        if w.lower() == "и":
            abbr_words.append("и")
        else:
            abbr_words.append(abbr.capitalize())

    return "".join(abbr_words)


def _parse_title_info_from_sheet(sheet) -> Dict[str, str]:
    info = {"code": ""}
    for r in range(1, min(sheet.max_row + 1, 100)):
        for c in range(1, min(sheet.max_column + 1, 20)):
            cell_val = sheet.cell(row=r, column=c).value
            if cell_val is None:
                continue
            val = str(cell_val).strip()
            if re.match(r"^\d{2}\.\d{2}\.\d{2}$", val):
                info["code"] = val
                return info
    return info


def load_excel_mapping(excel_path: Path) -> Tuple[Dict[str, str], str]:
    wb = load_workbook(str(excel_path.absolute()), data_only=True)

    specialty_code = ""
    if "Титул" in wb.sheetnames:
        title_info = _parse_title_info_from_sheet(wb["Титул"])
        specialty_code = title_info.get("code", "")

    plan_map = {}
    plan_sheet_name = "План"
    if plan_sheet_name not in wb.sheetnames:
        for name in wb.sheetnames:
            if "план" in name.lower() and "свод" not in name.lower():
                plan_sheet_name = name
                break

    if plan_sheet_name in wb.sheetnames:
        sheet = wb[plan_sheet_name]
        for row in range(1, sheet.max_row + 1):
            code_val = str(sheet.cell(row=row, column=2).value or "").strip()
            name_val = str(sheet.cell(row=row, column=3).value or "").strip()

            if code_val and name_val:
                if re.match(r"^[БФТД]\d+", code_val):
                    cleaned_name = clean_name_for_match(name_val)
                    plan_map[cleaned_name] = code_val

    return plan_map, specialty_code


def get_rp_title(doc: Document) -> str:
    found_marker = False
    for p in doc.paragraphs:
        txt = " ".join(p.text.split()).strip().lower()
        if "рабочая программа" in txt or "программа дисциплины" in txt:
            found_marker = True
            continue
        if found_marker:
            title = " ".join(p.text.split()).strip()
            if title:
                return title
    for p in doc.paragraphs[:12]:
        txt = " ".join(p.text.split()).strip()
        if txt and len(txt) > 5 and not txt.isupper() and "министерство" not in txt.lower():
            return txt
    return ""


def parse_authors(doc: Document) -> str:
    authors = []
    found_compiler_section = False
    parsed_paragraphs_count = 0

    stop_phrases = [
        "рабочая программа составлена", "протокол", "заведующий", "согласовано",
        "председатель", "учебно-методической", "количество часов", "декан", "директор"
    ]

    for p in doc.paragraphs[:100]:
        text = p.text.strip()
        if not text:
            continue

        text_lower = text.lower()

        if found_compiler_section:
            if any(phrase in text_lower for phrase in stop_phrases) or "содержание" in text_lower or "1." in text_lower:
                break
            parsed_paragraphs_count += 1
            if parsed_paragraphs_count > 4:
                break

        if "составитель" in text_lower or "разработчик" in text_lower or "составители" in text_lower:
            found_compiler_section = True
            parsed_paragraphs_count = 0

        if found_compiler_section:
            clean_line = re.sub(
                r'^(составитель|составители|разработчик|разработчики|составитель\s+-\s+|составители\s+-\s+)',
                '', text, flags=re.IGNORECASE
            ).strip()

            words_clean = []
            for w in clean_line.split():
                w_stripped = w.strip(".,()\"';:-")
                if w_stripped:
                    words_clean.append(w_stripped)

            words_filtered = [w for w in words_clean if w[0].isupper() and w.isalpha()]

            if len(words_filtered) >= 2:
                last_name = words_filtered[0].capitalize()
                first_init = words_filtered[1][0].upper() if words_filtered[1] else ""
                middle_init = ""
                if len(words_filtered) >= 3:
                    middle_init = words_filtered[2][0].upper() if words_filtered[2] else ""

                initials = f"{last_name}{first_init}{middle_init}"
                if initials not in authors:
                    authors.append(initials)

    if authors:
        return " ".join(authors)
    return ""


class FilenameStandardizer:

    def __init__(self, root_dir: str, excel_path: str):
        self.root_dir = Path(root_dir)
        self.excel_path = Path(excel_path)

    def run(self):
        if not self.root_dir.exists() or not self.excel_path.exists():
            logger.error("Проверьте правильность введенных путей к папке или файлу Excel.")
            return

        logger.info("Загрузка данных из учебного плана...")
        try:
            plan_map, specialty_code = load_excel_mapping(self.excel_path)
        except Exception as e:
            logger.error(f"Не удалось прочитать Excel-файл: {e}")
            return

        logger.info(f"Шифр направления: {specialty_code}")
        logger.info(f"Индексировано дисциплин из плана: {len(plan_map)}")

        updated_count = 0

        for docx_path in self.root_dir.glob("*.docx"):
            if docx_path.name.startswith("~$"):
                continue

            try:
                doc = Document(docx_path)
                title = get_rp_title(doc)
                if not title:
                    continue

                cleaned_title = clean_name_for_match(title)
                code = plan_map.get(cleaned_title)

                if not code:
                    logger.warning(f"  [-] Дисциплина '{title}' не найдена в учебном плане. Файл не переименован.")
                    continue

                authors = parse_authors(doc)
                abbr_discipline = abbreviate_discipline(title)

                new_base_name = f"{code} РП {abbr_discipline} {specialty_code} {authors}".strip()
                new_docx_name = f"{new_base_name}.docx"

                if new_docx_name == docx_path.name:
                    continue

                new_docx_path = docx_path.with_name(new_docx_name)
                counter = 1
                while new_docx_path.exists():
                    new_docx_path = docx_path.with_name(f"{new_base_name}_{counter}.docx")
                    counter += 1

                docx_path.rename(new_docx_path)
                logger.info(f"  [+] Переименован DOCX: '{docx_path.name}' -> '{new_docx_path.name}'")
                updated_count += 1

                pdf_path = docx_path.with_suffix(".pdf")
                if pdf_path.exists():
                    new_pdf_name = new_docx_path.with_suffix(".pdf").name
                    new_pdf_path = pdf_path.with_name(new_pdf_name)
                    pdf_path.rename(new_pdf_path)
                    logger.info(f"  [+] Переименован парный PDF: '{pdf_path.name}' -> '{new_pdf_path.name}'")

            except Exception as e:
                logger.error(f"Не удалось обработать файл {docx_path.name}: {e}")

        print(f"\n[Готово] Процесс стандартизации завершен. Успешно переименовано дисциплин: {updated_count}")


def main():
    print("=== Массовое приведение имен файлов РП к стандарту ===")
    try:
        user_root_dir = input("Шаг 1. Введите путь к папке с РП: ").strip().strip('"')
        user_excel_path = input("Шаг 2. Введите путь к файлу учебного плана Excel (plan.xlsx): ").strip().strip('"')
    except UnicodeDecodeError:
        print("Ошибка кодировки консоли! Используются пути по умолчанию.")
        user_root_dir = "."
        user_excel_path = "plan.xlsx"

    if not user_root_dir:
        user_root_dir = "."
    if not user_excel_path:
        user_excel_path = "plan.xlsx"

    standardizer = FilenameStandardizer(user_root_dir, user_excel_path)
    standardizer.run()


if __name__ == "__main__":
    main()