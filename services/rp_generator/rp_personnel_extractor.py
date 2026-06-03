import os
import re
import json
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from docx import Document
from openpyxl import load_workbook

# Настройка логирования
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s | %(levelname)s | %(message)s"
)
logger = logging.getLogger(__name__)

# === ГИБКИЕ РЕГУЛЯРНЫЕ ВЫРАЖЕНИЯ ===
ROLE_DEAN = re.compile(
    r"(?:(?:и\.?\s*о\.?|врио|временно\s+исполняющий\s+обязанности)\s+)?(декан|директор|руководитель\s+института)\b",
    re.IGNORECASE
)
ROLE_HEAD = re.compile(
    r"(?:(?:и\.?\s*о\.?|врио)\s+)?(?:заведующий|зав\.?)\s*(?:кафедрой|каф\.?)\b",
    re.IGNORECASE
)
ROLE_DIRECTOR = re.compile(
    r"(?:руководитель\s+(?:образовательной\s+)?программы|руководитель\s+оп|руководитель\s+направления)\b",
    re.IGNORECASE
)
ROLE_COMPILER = re.compile(
    r"(?:составител[ьяи]|разработчик[и]?)\b",
    re.IGNORECASE
)

DEGREE_PATTERN = re.compile(
    r"(?:преподаватель|доцент|профессор|ассистент|к\.?\s*[тфмэ]\.?\s*н\.?|д\.?\s*[тфмэ]\.?\s*н\.?|ст\.?\s*преп)",
    re.IGNORECASE
)

STOP_COMPILER_PHRASES = [
    "рабочая программа составлена",
    "в соответствии с требованиями",
    "рассмотрена на заседании",
    "протокол от",
    "согласовано",
    "кафедра",
    "заведующий",
    "декан"
]

DIRECTION_ROLE = re.compile(r"(?:направление|специальность|подготовк\w*)\b", re.IGNORECASE)
CODE_PATTERN = re.compile(r"(\d{2}\.\d{2}\.\d{2})")
PROFILE_ROLE = re.compile(r"(?:направленность|профиль|программа/специализация|специализация)\b", re.IGNORECASE)
DISCIPLINE_ROLE = re.compile(
    r"(?:рабочая\s+программа\s+дисциплины|оценочные\s+средства\s+по\s+дисциплине)",
    re.IGNORECASE
)

INITIALS_PATTERN = re.compile(
    r'(?:[А-Я]\s*\.\s*[А-Я]\s*\.\s*[А-Я][а-я]+|[А-Я][а-я]+\s+[А-Я]\s*\.\s*[А-Я]\s*\.)'
)

# Скорректированный паттерн для поиска ФИО без жесткой привязки \b к символу точки в конце инициалов
FIO_REGEX = re.compile(
    r'(?:'
    r'\b[А-Я][а-я]+(?:-[А-Я][а-я]+)?\s+[А-Я]\s*\.\s*[А-Я]\s*\.?|'  # Фамилия И.О.
    r'\b[А-Я]\s*\.\s*[А-Я]\s*\.?\s*[А-Я][а-я]+(?:-[А-Я][а-я]+)?\b|'  # И.О. Фамилия
    r'\b[А-Я][а-я]+(?:-[А-Я][а-я]+)?\s+[А-Я][а-я]+\s+[А-Я][а-я]+\b|'  # ... Имя Отчество
    r'\b[А-Я][а-я]+(?:-[А-Я][а-я]+)?\s+[А-Я][а-я]+\b'  # Фамилия Имя
    r')'
)


def clean_text(text: str) -> str:
    if not text:
        return ""
    s = re.sub(r'[_/\\|]+', ' ', text)
    s = " ".join(s.split()).strip()
    return s


def is_placeholder_name(name_str: str) -> bool:
    """Определяет, является ли переданная строка шаблонной заглушкой (например, 'И.О. Фамилия')."""
    norm = name_str.lower().strip()
    if "фамилия" in norm:
        return True
    if "имя" in norm or "отчество" in norm:
        return True
    if re.match(r'^[ио\s\._/\\]+$', norm):
        return True
    return False


def standardize_initials_name(name_str: str) -> str:
    """Приводит инициалы к единообразному формату 'И.О. Фамилия'."""
    name_str = clean_text(name_str)

    # Сценарий 1: Фамилия И.О. -> И.О. Фамилия
    pattern_surname_first = re.compile(
        r'^([А-Я][а-я]+(?:-[А-Я][а-я]+)?)\s+([А-Я]\s*\.\s*[А-Я]\s*\.?)$'
    )
    m1 = pattern_surname_first.match(name_str)
    if m1:
        surname = m1.group(1)
        initials = m1.group(2).replace(" ", "")
        if not initials.endswith("."):
            initials += "."
        return f"{initials} {surname}"

    # Сценарий 2: И.О. Фамилия
    pattern_initials_first = re.compile(
        r'^([А-Я]\s*\.\s*[А-Я]\s*\.?)\s+([А-Я][а-я]+(?:-[А-Я][а-я]+)?)$'
    )
    m2 = pattern_initials_first.match(name_str)
    if m2:
        initials = m2.group(1).replace(" ", "")
        if not initials.endswith("."):
            initials += "."
        surname = m2.group(2)
        return f"{initials} {surname}"

    return name_str


def clean_fio_spaces(name_str: str) -> str:
    s = clean_text(name_str)
    s = re.sub(r'([А-Я]\.)\s+([А-Я]\.)', r'\1\2', s)
    s = re.sub(r'([А-Я][а-я]+)([А-Я]\.[А-Я]\.)', r'\1 \2', s)
    s = re.sub(r'([А-Я]\.[А-Я]\.)([А-Я][а-я]+)', r'\1 \2', s)
    s = standardize_initials_name(s)
    return s


def normalize_profile_text(text: str) -> str:
    """Нормализует строку профиля для сопоставления, удаляя кавычки и знаки препинания."""
    s = text.lower()
    s = re.sub(r'[«»""\'\'„“]', '', s)
    s = re.sub(r'^[\s,.;:-]+|[\s,.;:-]+$', '', s)
    s = " ".join(s.split())
    return s


def has_name_pattern(text: str) -> bool:
    if INITIALS_PATTERN.search(text):
        return True
    capitalized = re.findall(r'\b[А-Я][а-я]+\b', text)
    if len(capitalized) >= 2:
        return True
    return False


def split_multi_compiler_string(text: str) -> List[str]:
    """Разбивает сложную строку с несколькими авторами на отдельные сегменты по границам ФИО."""
    matches = list(FIO_REGEX.finditer(text))
    if not matches:
        return [text]

    segments = []
    for i, match in enumerate(matches):
        start = match.start()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        segment_text = text[start:end].strip()
        segment_text = re.sub(r'^[\s,.;:-]+|[\s,.;:-]+$', '', segment_text)
        if segment_text:
            segments.append(segment_text)
    return segments


def split_compiler_name_and_degree(compiler_str: str) -> Tuple[str, str]:
    """Выделяет каноническое ФИО и ученую степень/должность составителя."""
    compiler_str = clean_text(compiler_str)

    match = FIO_REGEX.search(compiler_str)
    if not match:
        parts = compiler_str.split(",", 1)
        name = clean_text(parts[0])
        degree = clean_text(parts[1]) if len(parts) > 1 else ""
        return name, degree

    name = match.group(0).strip()

    start, end = match.span()
    prefix = compiler_str[:start].strip()
    suffix = compiler_str[end:].strip()

    degree_parts = []
    if prefix:
        degree_parts.append(prefix)
    if suffix:
        degree_parts.append(suffix)

    degree = ", ".join(degree_parts)
    degree = re.sub(r'^[\s,.;:-]+|[\s,.;:-]+$', '', degree).strip()

    return name, degree


def levenshtein_distance(s1: str, s2: str) -> int:
    if len(s1) < len(s2):
        return levenshtein_distance(s2, s1)
    if len(s2) == 0:
        return len(s1)

    previous_row = range(len(s2) + 1)
    for i, c1 in enumerate(s1):
        current_row = [i + 1]
        for j, c2 in enumerate(s2):
            insertions = previous_row[j + 1] + 1
            deletions = current_row[j] + 1
            substitutions = previous_row[j] + (c1 != c2)
            current_row.append(min(insertions, deletions, substitutions))
        previous_row = current_row

    return previous_row[-1]


def is_similar_name(name1: str, name2: str) -> bool:
    n1 = name1.lower().replace("ё", "е").strip()
    n2 = name2.lower().replace("ё", "е").strip()

    if n1 == n2:
        return True

    w1 = [w for w in re.split(r'[^а-яa-z]', n1) if w]
    w2 = [w for w in re.split(r'[^а-яa-z]', n2) if w]

    if not w1 or not w2:
        return False

    surname1, surname2 = w1[0], w2[0]

    if levenshtein_distance(surname1, surname2) <= 1:
        init1 = "".join([w[0] for w in w1[1:] if len(w) > 0])
        init2 = "".join([w[0] for w in w2[1:] if len(w) > 0])

        if init1 and init2:
            min_len = min(len(init1), len(init2))
            if init1[:min_len] == init2[:min_len]:
                return True
    return False


def find_person_by_role(lines: List[str], role_regex: re.Pattern, name_regex: re.Pattern, scan_offset: int = 4) -> str:
    """Ищет должностное лицо по роли, игнорируя шаблонные заглушки."""
    for idx, line in enumerate(lines):
        line_clean = clean_text(line)
        if role_regex.search(line_clean):
            for offset in range(0, scan_offset):
                if idx + offset < len(lines):
                    pot_line = clean_text(lines[idx + offset])
                    m = name_regex.search(pot_line)
                    if m:
                        potential_name = clean_fio_spaces(m.group(0))
                        if not is_placeholder_name(potential_name):
                            return potential_name
    return ""


def extract_default_personnel_from_excel(sheet) -> Dict[str, str]:
    default_staff = {
        "dean": "",
        "head_of_department": "",
        "program_director": "",
        "umk_chairman": "",
        "oamr_head": ""
    }

    role_dean_pattern = re.compile(r"декан", re.IGNORECASE)
    role_head_pattern = re.compile(r"(?:зав\.\s*кафедрой|заведующий\s*кафедрой|зав\. кафедра)", re.IGNORECASE)
    role_director_pattern = re.compile(r"руководитель\s+(?:ооп|направления|ооп)\b", re.IGNORECASE)
    role_umk_pattern = re.compile(r"председатель\s+умк", re.IGNORECASE)
    role_oamr_pattern = re.compile(r"начальник\s+оамр", re.IGNORECASE)

    fio_excel_pattern = re.compile(
        r"/\s*([А-Яа-я]\s*\.\s*[А-Яа-я]\s*\.\s*[А-Яа-я]+|[А-Яа-я]+\s+[А-Яа-я]\s*\.\s*[А-Яа-я]\s*\.)\s*/"
    )
    fio_fallback_pattern = re.compile(
        r"([А-Яа-я]\s*\.\s*[А-Яа-я]\s*\.\s*[А-Яа-я]+|[А-Яа-я]+\s+[А-Яа-я]\s*\.\s*[А-Яа-я]\s*\.)"
    )

    for row in range(1, sheet.max_row + 1):
        for col in range(1, sheet.max_column + 1):
            cell_val = str(sheet.cell(row=row, column=col).value or "").strip()
            if not cell_val:
                continue

            target_key = None
            if role_dean_pattern.search(cell_val):
                target_key = "dean"
            elif role_head_pattern.search(cell_val):
                target_key = "head_of_department"
            elif role_director_pattern.search(cell_val):
                target_key = "program_director"
            elif role_umk_pattern.search(cell_val):
                target_key = "umk_chairman"
            elif role_oamr_pattern.search(cell_val):
                target_key = "oamr_head"

            if target_key:
                for offset in range(1, 15):
                    if col + offset <= sheet.max_column:
                        test_cell = str(sheet.cell(row=row, column=col + offset).value or "").strip()
                        if test_cell:
                            m = fio_excel_pattern.search(test_cell)
                            if m:
                                potential_name = clean_fio_spaces(m.group(1))
                                if not is_placeholder_name(potential_name):
                                    default_staff[target_key] = potential_name
                                    break
                            m_fb = fio_fallback_pattern.search(test_cell)
                            if m_fb:
                                potential_name = clean_fio_spaces(m_fb.group(1))
                                if not is_placeholder_name(potential_name):
                                    default_staff[target_key] = potential_name
                                    break
    return default_staff


def parse_rp_file(file_path: Path) -> Optional[dict]:
    try:
        doc = Document(file_path)
    except Exception as e:
        logger.error(f"Не удалось открыть файл {file_path.name}: {e}")
        return None

    lines: List[str] = []
    for p in doc.paragraphs:
        val = p.text.strip()
        if val:
            lines.append(val)

    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                val = cell.text.strip()
                if val:
                    for line in val.split("\n"):
                        line_strip = line.strip()
                        if line_strip and line_strip not in lines:
                            lines.append(line_strip)

    direction_code = ""
    direction_name = ""
    for idx, line in enumerate(lines):
        line_clean = clean_text(line)
        if DIRECTION_ROLE.search(line_clean):
            m_code = CODE_PATTERN.search(line_clean)
            if m_code:
                direction_code = m_code.group(1)
                parts = line_clean.split(direction_code, 1)
                if len(parts) > 1:
                    pot_name = parts[1].strip()
                    pot_name = re.sub(r'(?:код|наименование|полностью).*', '', pot_name, flags=re.IGNORECASE).strip()
                    pot_name = re.sub(r'^[-\s,.:\)]+', '', pot_name).strip()
                    if pot_name:
                        direction_name = pot_name
                break

    profile = ""
    for idx, line in enumerate(lines):
        line_clean = clean_text(line)
        if PROFILE_ROLE.search(line_clean):
            # Попытка разделения по двоеточию или закрывающей скобке
            parts = re.split(r'[:\)]', line_clean, 1)
            if len(parts) > 1:
                pot_profile = parts[-1].strip()
            else:
                # Если разделителей нет (ваш случай), вырезаем само ключевое слово роли
                pot_profile = PROFILE_ROLE.sub("", line_clean, count=1).strip()

            # Если строка профиля осталась пустой/короткой, ищем в следующих строках (с фильтрацией)
            if not pot_profile or len(pot_profile) < 5 or any(
                    x in pot_profile.lower() for x in ["направленность", "профиль", "наименование"]):
                for offset in range(1, 3):
                    if idx + offset < len(lines):
                        pot = clean_text(lines[idx + offset])
                        if pot and len(pot) > 5 and not any(
                                x in pot.lower() for x in
                                ["направленность", "профиль", "наименование", "уровень образования", "форма обучения"]):
                            pot_profile = pot
                            break

            if pot_profile:
                pot_profile = re.sub(r"^.*?\)\s*", "", pot_profile)
                pot_profile = re.sub(r'(?:наименование|полностью).*', '', pot_profile, flags=re.IGNORECASE).strip()
                pot_profile = re.sub(r'^[-\s,.:\)]+', '', pot_profile).strip()
                if pot_profile:
                    profile = pot_profile
                    break

    discipline = ""
    for idx, line in enumerate(lines):
        line_clean = clean_text(line)
        if DISCIPLINE_ROLE.search(line_clean):
            for offset in range(1, 4):
                if idx + offset < len(lines):
                    pot = clean_text(lines[idx + offset])
                    if pot and not any(x in pot.lower() for x in ["наименование", "направление", "код"]):
                        discipline = pot
                        break
            if discipline:
                break

    dean = find_person_by_role(lines, ROLE_DEAN, INITIALS_PATTERN, scan_offset=4)
    head_of_department = find_person_by_role(lines, ROLE_HEAD, INITIALS_PATTERN, scan_offset=4)
    program_director = find_person_by_role(lines, ROLE_DIRECTOR, INITIALS_PATTERN, scan_offset=4)

    raw_compilers: List[str] = []
    for idx, line in enumerate(lines):
        line_clean = clean_text(line)
        if ROLE_COMPILER.search(line_clean):
            for offset in range(0, 7):
                if idx + offset < len(lines):
                    pot = clean_text(lines[idx + offset])

                    if offset > 0 and (ROLE_DEAN.search(pot) or ROLE_HEAD.search(pot) or ROLE_DIRECTOR.search(pot) or
                                       any(p in pot.lower() for p in STOP_COMPILER_PHRASES)):
                        break

                    if pot:
                        cleaned = ROLE_COMPILER.sub("", pot).strip()
                        cleaned = re.sub(r'^[-\s,.:]+', '', cleaned).strip()
                        cleaned = re.sub(r'\s*Ф\.И\.О\..*', '', cleaned, flags=re.IGNORECASE).strip()
                        if cleaned and not cleaned.startswith("_") and len(cleaned) > 4:
                            raw_compilers.append(cleaned)
            break

    compilers: List[str] = []
    for item in raw_compilers:
        if has_name_pattern(item):
            compilers.append(item)
        else:
            if compilers:
                compilers[-1] = (compilers[-1].rstrip(",") + ", " + item).strip()
                compilers[-1] = re.sub(r'\s*,\s*,', ',', compilers[-1])
                compilers[-1] = " ".join(compilers[-1].split())
            else:
                compilers.append(item)

    # Постобработка: разбиваем склеенных в одну строку составителей на отдельные элементы
    final_compilers: List[str] = []
    for c in compilers:
        split_items = split_multi_compiler_string(c)
        final_compilers.extend(split_items)
    compilers = final_compilers

    if not discipline:
        logger.warning(f"В файле {file_path.name} не удалось определить название дисциплины.")
        return None

    direction_full = f"{direction_code} {direction_name}".strip() if direction_code else "Неизвестное направление"

    return {
        "direction": direction_full,
        "profile": profile,
        "discipline": discipline,
        "data": {
            "dean": dean,
            "compilers": compilers,
            "head_of_department": head_of_department,
            "program_director": program_director,
            "profile": profile
        }
    }


class PersonnelExtractor:
    """Класс-экстрактор для сбора и нормализации должностных лиц по реляционной схеме."""

    def __init__(self, excel_path: Optional[Path] = None, rp_folder: Optional[Path] = None):
        self.excel_path = excel_path
        self.rp_folder = rp_folder

    def extract(self) -> Dict[str, Any]:
        """Запускает конвейер извлечения и нормализации данных."""
        default_staff = {}

        # 1. Сбор должностных лиц кафедры по умолчанию из Excel
        if self.excel_path and self.excel_path.exists():
            logger.info(f"Парсинг листа 'Титул' из Excel плана: {self.excel_path.name}...")
            try:
                wb = load_workbook(str(self.excel_path.absolute()), data_only=True)
                if "Титул" in wb.sheetnames:
                    default_staff = extract_default_personnel_from_excel(wb["Титул"])
                    logger.info(f"  [+] Извлечены должностные лица по умолчанию: {default_staff}")
                else:
                    logger.warning("Лист 'Титул' не обнаружен в Excel.")
            except Exception as e:
                logger.error(f"Не удалось распарсить Excel: {e}")

        # 2. Сбор сырых данных из Word-файлов
        raw_subjects: Dict[str, Dict[str, Any]] = {}
        if self.rp_folder and self.rp_folder.exists():
            docx_files = list(self.rp_folder.glob("*.docx"))
            logger.info(f"Найдено файлов для анализа: {len(docx_files)}")

            for docx_path in docx_files:
                if docx_path.name.startswith("~$"):
                    continue

                logger.info(f"Анализ файла РП: {docx_path.name}")
                result = parse_rp_file(docx_path)
                if result:
                    dir_key = result["direction"]
                    disc_key = result["discipline"]

                    if dir_key not in raw_subjects:
                        raw_subjects[dir_key] = {}

                    extracted_data = result["data"]

                    # Наложение значений по умолчанию, если поля пустые
                    for key in ["dean", "head_of_department", "program_director"]:
                        if not extracted_data.get(key) and default_staff.get(key):
                            extracted_data[key] = default_staff[key]

                    raw_subjects[dir_key][disc_key] = extracted_data
                    logger.info(f"  [+] Собран профиль для: {disc_key} ({dir_key})")
        else:
            logger.warning("Папка с Word РП не найдена или не указана.")

        # 3. Нормализация сущностей по реляционному принципу (БД)
        return self._normalize_all_entities(raw_subjects, default_staff)

    def _cluster_names(self, raw_names: List[str], prefix: str) -> Tuple[Dict[str, Any], Dict[str, str]]:
        """Группирует вариации имен людей и генерирует словарь сущностей и lookup карту."""
        clusters = []
        for r_name in raw_names:
            c_name = clean_fio_spaces(r_name)
            if not c_name or is_placeholder_name(c_name):
                continue
            matched_cluster = None
            for cl in clusters:
                canonical_name = max(cl["names_freq"], key=cl["names_freq"].get)
                if is_similar_name(c_name, canonical_name):
                    matched_cluster = cl
                    break
            if matched_cluster is None:
                matched_cluster = {"names_freq": {}}
                clusters.append(matched_cluster)
            matched_cluster["names_freq"][c_name] = matched_cluster["names_freq"].get(c_name, 0) + 1

        entities = {}
        lookup = {}
        for idx, cl in enumerate(clusters, start=1):
            entity_id = f"{prefix}{idx}"
            canonical_name = max(cl["names_freq"], key=cl["names_freq"].get)
            entities[entity_id] = {"name": canonical_name}

            for var_name in cl["names_freq"]:
                lookup[var_name.lower()] = entity_id

        return entities, lookup

    def _cluster_profiles(self, raw_profiles: List[str]) -> Tuple[Dict[str, Any], Dict[str, str]]:
        """Нормализует текстовые названия профилей, устраняя мелкие различия и кавычки."""
        entities = {}
        lookup = {}
        seen_profiles = {}

        idx = 1
        for prof in raw_profiles:
            clean_prof = clean_text(prof)
            if not clean_prof:
                continue

            norm_prof = normalize_profile_text(clean_prof)

            # Пропуск технических/информационных строк, не являющихся названиями профилей
            if any(x in norm_prof for x in ["уровень образования", "бакалавриат", "магистратура", "форма обучения"]):
                continue

            matched_id = None
            for existing_norm, existing_id in seen_profiles.items():
                if existing_norm == norm_prof:
                    matched_id = existing_id
                    break

                # Сопоставление с помощью расстояния Левенштейна для фильтрации опечаток/вариаций
                max_len = max(len(existing_norm), len(norm_prof))
                if max_len > 10 and levenshtein_distance(existing_norm, norm_prof) / max_len < 0.1:
                    matched_id = existing_id
                    break

            if matched_id is None:
                entity_id = f"profile_{idx}"
                seen_profiles[norm_prof] = entity_id

                # Формируем каноническое имя (удаляя внешние кавычки для корректного отображения)
                canonical_name = re.sub(r'^[\s«"\'„]+|[\s»"\'“]+$', '', clean_prof).strip()
                entities[entity_id] = {"name": canonical_name}
                idx += 1
                matched_id = entity_id

            lookup[clean_prof.lower()] = matched_id

        return entities, lookup

    def _normalize_teachers(self, raw_subjects: Dict[str, Dict[str, Any]]) -> Tuple[Dict[str, Any], Dict[str, str]]:
        """Выделяет канонические записи преподавателей и строит lookup карту."""
        raw_compilers_seen = []
        for dir_key, subjects in raw_subjects.items():
            for disc_key, data in subjects.items():
                for comp_str in data.get("compilers", []):
                    name, degree = split_compiler_name_and_degree(comp_str)
                    # Фильтрация шаблонных имен-заглушек
                    if name and not is_placeholder_name(name):
                        raw_compilers_seen.append((comp_str, name, degree))

        clusters = []
        for raw_str, name, degree in raw_compilers_seen:
            matched_cluster = None
            for cl in clusters:
                canonical_name = max(cl["names_freq"], key=cl["names_freq"].get)
                if is_similar_name(name, canonical_name):
                    matched_cluster = cl
                    break

            if matched_cluster is None:
                matched_cluster = {
                    "names_freq": {},
                    "degrees_freq": {},
                    "raw_strings": set()
                }
                clusters.append(matched_cluster)

            matched_cluster["names_freq"][name] = matched_cluster["names_freq"].get(name, 0) + 1
            if degree:
                matched_cluster["degrees_freq"][degree] = matched_cluster["degrees_freq"].get(degree, 0) + 1
            matched_cluster["raw_strings"].add(raw_str)

        teachers_dict = {}
        raw_to_teacher_id = {}

        for idx, cl in enumerate(clusters, start=1):
            teacher_id = f"teacher_{idx}"
            canonical_name = max(cl["names_freq"], key=cl["names_freq"].get)

            canonical_degree = ""
            if cl["degrees_freq"]:
                canonical_degree = max(cl["degrees_freq"], key=lambda k: (cl["degrees_freq"][k], len(k)))

            teachers_dict[teacher_id] = {
                "name": canonical_name,
                "degree_and_title": canonical_degree
            }

            for raw_str in cl["raw_strings"]:
                raw_to_teacher_id[raw_str.lower()] = teacher_id

        return teachers_dict, raw_to_teacher_id

    def _lookup_id(self, val: str, lookup_dict: Dict[str, str]) -> Optional[str]:
        if not val:
            return None
        return lookup_dict.get(clean_text(val).lower(), None)

    def _normalize_all_entities(self, raw_subjects: Dict[str, Dict[str, Any]], default_staff: Dict[str, str]) -> Dict[
        str, Any]:
        """Группирует все упоминания в единую структуру реляционных данных."""

        # Сбор сырых значений для каждого типа сущностей
        raw_deans = []
        raw_heads = []
        raw_pds = []
        raw_umk = []
        raw_oamr = []
        raw_profiles = []

        # Сбор из дефолтных настроек Excel
        if default_staff.get("dean"):
            raw_deans.append(default_staff["dean"])
        if default_staff.get("head_of_department"):
            raw_heads.append(default_staff["head_of_department"])
        if default_staff.get("program_director"):
            raw_pds.append(default_staff["program_director"])
        if default_staff.get("umk_chairman"):
            raw_umk.append(default_staff["umk_chairman"])
        if default_staff.get("oamr_head"):
            raw_oamr.append(default_staff["oamr_head"])

        # Сбор из распарсенных Word РП
        for dir_key, subjects in raw_subjects.items():
            for disc_key, data in subjects.items():
                if data.get("dean"):
                    raw_deans.append(data["dean"])
                if data.get("head_of_department"):
                    raw_heads.append(data["head_of_department"])
                if data.get("program_director"):
                    raw_pds.append(data["program_director"])
                if data.get("profile"):
                    raw_profiles.append(data["profile"])

        # Вызов кластеризации для всех справочников
        deans_dict, deans_lookup = self._cluster_names(raw_deans, "dean_")
        heads_dict, heads_lookup = self._cluster_names(raw_heads, "hod_")
        pds_dict, pds_lookup = self._cluster_names(raw_pds, "pd_")
        umk_dict, umk_lookup = self._cluster_names(raw_umk, "umk_")
        oamr_dict, oamr_lookup = self._cluster_names(raw_oamr, "oamr_")
        profiles_dict, profiles_lookup = self._cluster_profiles(raw_profiles)
        teachers_dict, teachers_lookup = self._normalize_teachers(raw_subjects)

        # Нормализация значений по умолчанию
        mapped_default_staff = {
            "dean": self._lookup_id(default_staff.get("dean"), deans_lookup),
            "head_of_department": self._lookup_id(default_staff.get("head_of_department"), heads_lookup),
            "program_director": self._lookup_id(default_staff.get("program_director"), pds_lookup),
            "umk_chairman": self._lookup_id(default_staff.get("umk_chairman"), umk_lookup),
            "oamr_head": self._lookup_id(default_staff.get("oamr_head"), oamr_lookup)
        }

        # Связывание в subjects_mapping по идентификаторам
        subjects_mapping = {}
        for dir_key, subjects in raw_subjects.items():
            subjects_mapping[dir_key] = {}
            for disc_key, data in subjects.items():

                # Маппинг преподавателей
                mapped_teacher_ids = []
                for comp_str in data.get("compilers", []):
                    tid = teachers_lookup.get(comp_str.lower())
                    if tid and tid not in mapped_teacher_ids:
                        mapped_teacher_ids.append(tid)

                subjects_mapping[dir_key][disc_key] = {
                    "profile": self._lookup_id(data.get("profile"), profiles_lookup),
                    "dean": self._lookup_id(data.get("dean"), deans_lookup),
                    "head_of_department": self._lookup_id(data.get("head_of_department"), heads_lookup),
                    "program_director": self._lookup_id(data.get("program_director"), pds_lookup),
                    "teachers": mapped_teacher_ids
                }

        return {
            "default_department_personnel": mapped_default_staff,
            "deans": deans_dict,
            "heads_of_department": heads_dict,
            "program_directors": pds_dict,
            "umk_chairmen": umk_dict,
            "oamr_heads": oamr_dict,
            "profiles": profiles_dict,
            "teachers": teachers_dict,
            "subjects_mapping": subjects_mapping
        }


def main():
    print("=== Комплексный экстрактор и маппер должностных лиц РП ===")

    user_excel = input("Шаг 1. Введите путь к файлу Excel учебного плана (например, plan.xlsx): ").strip()
    if not user_excel:
        user_excel = "plan.xlsx"
    excel_path = Path(user_excel)

    user_rp_folder = input("\nШаг 2. Введите путь к папке с архивными файлами РП (docx): ").strip()
    if not user_rp_folder:
        user_rp_folder = "."
    rp_folder = Path(user_rp_folder)

    user_output_dir = input(
        "\nШаг 3. Введите путь к папке для сохранения итогового JSON (по умолчанию 'services/rp_generator'): ").strip()
    if not user_output_dir:
        user_output_dir = "services/rp_generator"

    extractor = PersonnelExtractor(excel_path=excel_path, rp_folder=rp_folder)
    try:
        final_output = extractor.extract()

        output_dir = Path(user_output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
        output_mapping_path = output_dir / "rp_personnel_mapping.json"

        with open(output_mapping_path, "w", encoding="utf-8") as f:
            json.dump(final_output, f, ensure_ascii=False, indent=2)
        print(f"\n[Успешно] Данные по кадровому составу РП сохранены в:\n{output_mapping_path.absolute()}")
    except Exception as e:
        logger.error(f"Не удалось завершить извлечение данных: {e}", exc_info=True)


if __name__ == "__main__":
    main()