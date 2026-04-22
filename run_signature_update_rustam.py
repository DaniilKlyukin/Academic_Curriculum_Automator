import os
import re
from pathlib import Path
from typing import Optional, Tuple, List
from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run


def replace_text_in_run(run: Run, old_text: str, new_text: str) -> bool:
    """
    Заменяет текст в run, сохраняя форматирование.
    
    Args:
        run: Run для замены
        old_text: Текст для поиска
        new_text: Новый текст
        
    Returns:
        True если замена выполнена, False иначе
    """
    if old_text not in run.text:
        return False
    
    # Сохраняем форматирование
    old_format = {
        'bold': run.bold,
        'italic': run.italic,
        'underline': run.underline,
        'color': run.font.color.rgb,
        'size': run.font.size,
        'name': run.font.name,
    }
    
    # Заменяем текст
    run.text = run.text.replace(old_text, new_text)
    
    # Восстанавливаем форматирование
    run.bold = old_format['bold']
    run.italic = old_format['italic']
    run.underline = old_format['underline']
    run.font.color.rgb = old_format['color']
    run.font.size = old_format['size']
    run.font.name = old_format['name']
    
    return True


def process_docx_signatures(
    file_path: str,
    old_fio: Optional[str],
    new_fio: Optional[str],
    old_title: Optional[str],
    new_title: Optional[str]
) -> Tuple[bool, str]:
    """
    Замена ФИО и должности в тексте документа.
    
    Args:
        file_path: Путь к файлу .docx
        old_fio: Старое ФИО для замены
        new_fio: Новое ФИО
        old_title: Старая должность для замены
        new_title: Новая должность
        
    Returns:
        Кортеж (успех, сообщение)
    """
    # Проверка входных данных
    if not old_fio and not old_title:
        return False, "Нечего заменять"
    
    try:
        doc = Document(file_path)
        replaced = False
        
        # Проходим по всем параграфам документа
        for para in doc.paragraphs:
            # Замена ФИО
            if old_fio and old_fio in para.text:
                for run in para.runs:
                    if old_fio in run.text:
                        if replace_text_in_run(run, old_fio, new_fio):
                            replaced = True
            
            # Замена должности
            if old_title and old_title in para.text:
                for run in para.runs:
                    if old_title in run.text:
                        if replace_text_in_run(run, old_title, new_title):
                            replaced = True
        
        # Сохраняем изменения
        if replaced:
            doc.save(file_path)
            return True, "OK"
        
        return False, "не найдены"
        
    except FileNotFoundError:
        return False, f"Файл не найден: {file_path}"
    except Exception as e:
        return False, str(e)


def collect_all_paragraphs(doc: Document) -> List[Paragraph]:
    """
    Собирает все параграфы документа включая таблицы, заголовки и подвалы.
    
    Args:
        doc: Объект Document
        
    Returns:
        Список всех параграфов
    """
    all_paras: List[Paragraph] = list(doc.paragraphs)
    
    # Добавляем параграфы из таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paras.extend(cell.paragraphs)
    
    # Добавляем параграфы из заголовков и подвалов
    for section in doc.sections:
        if section.header.paragraphs:
            all_paras.extend(section.header.paragraphs)
        if section.footer.paragraphs:
            all_paras.extend(section.footer.paragraphs)
    
    return all_paras


def replace_word_in_docx(
    file_path: str,
    old_word: str,
    new_word: str
) -> Tuple[bool, str]:
    """
    Замена целого слова с сохранением форматирования.
    
    Args:
        file_path: Путь к файлу .docx
        old_word: Слово для поиска
        new_word: Новое слово
        
    Returns:
        Кортеж (успех, сообщение)
    """
    if not old_word:
        return False, "Не указано слово для замены"
    
    try:
        doc = Document(file_path)
        pattern = re.compile(rf'\b{re.escape(old_word)}\b', re.IGNORECASE)
        replaced = 0
        
        # Собираем все параграфы
        all_paras = collect_all_paragraphs(doc)
        
        # Проходим по всем параграфам
        for para in all_paras:
            if not pattern.search(para.text):
                continue
            
            # Сохраняем форматирование первого run
            fmt = para.runs[0] if para.runs else None
            
            # Заменяем текст
            new_text = pattern.sub(new_word, para.text)
            
            if new_text != para.text:
                para.text = new_text
                
                # Восстанавливаем форматирование
                if fmt:
                    for run in para.runs:
                        run.bold = fmt.bold
                        run.italic = fmt.italic
                        run.underline = fmt.underline
                        run.font.color.rgb = fmt.font.color.rgb
                        run.font.size = fmt.font.size
                        run.font.name = fmt.font.name
                
                replaced += 1
        
        # Сохраняем изменения
        if replaced:
            doc.save(file_path)
            return True, f"Замен: {replaced}"
        
        return False, "не найдено"
        
    except FileNotFoundError:
        return False, f"Файл не найден: {file_path}"
    except Exception as e:
        return False, str(e)


def main() -> None:
    """Основная функция программы."""
    print("=== Замена в .docx (рекурсивно) ===")
    
    # Получаем путь к папке
    dir_path = input("Путь к папке: ").strip().strip('"')
    
    if not os.path.isdir(dir_path):
        print("Ошибка: папка не найдена")
        return
    
    # Получаем параметры замены
    old_fio = input("Старое ФИО (Enter - пропустить): ").strip()
    new_fio = input("Новое ФИО: ").strip() if old_fio else ""
    
    old_title = input("Старая должность (Enter - пропустить): ").strip()
    new_title = input("Новая должность: ").strip() if old_title else ""
    
    old_word = input("Слово для замены (Enter - пропустить): ").strip()
    new_word = input("Новое слово: ").strip() if old_word else ""
    
    # Проверка, есть ли что-то для замены
    if not (old_fio or old_title or old_word):
        print("Ошибка: ничего не указано для замены")
        return
    
    # Статистика обработки
    stats = {"ok": 0, "skip": 0, "err": 0}
    
    # Проходим по всем файлам в директории
    for root, _, files in os.walk(dir_path):
        for f in files:
            # Фильтруем только .docx файлы, игнорируем системные файлы
            if not f.lower().endswith('.docx') or f.startswith('~$'):
                continue
            
            path = os.path.join(root, f)
            rel = os.path.relpath(path, dir_path)
            modified = False
            
            # Обработка замены ФИО и должности
            if old_fio or old_title:
                ok, msg = process_docx_signatures(
                    path, old_fio, new_fio, old_title, new_title
                )
                if ok:
                    modified = True
                elif "не найдены" not in msg:
                    print(f"[ERR] {rel}: {msg}")
                    stats["err"] += 1
                    continue
            
            # Обработка замены слова
            if old_word:
                ok, msg = replace_word_in_docx(path, old_word, new_word)
                if ok:
                    modified = True
                elif "не найдено" not in msg:
                    print(f"[ERR] {rel}: {msg}")
                    stats["err"] += 1
                    continue
            
            # Вывод результата обработки файла
            if modified:
                print(f"[OK] {rel}")
                stats["ok"] += 1
            else:
                print(f"[SKIP] {rel}")
                stats["skip"] += 1
    
    # Вывод итоговой статистики
    print(f"\n=== Готово ===")
    print(f"Обработано: {stats['ok']}, Пропущено: {stats['skip']}, Ошибок: {stats['err']}")


if __name__ == "__main__":
    main()
