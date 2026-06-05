# -*- coding: utf-8 -*-
"""
Модуль rp_doc_sections.py
Содержит шаблоны разделов РПД и ФОС с динамическим вычислением дат и реляционным выводом ЗУН.
"""

import re
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

from services.rp_generator.rp_doc_styles import (
    add_paragraph_with_spacing, set_cell_text, set_row_cant_split,
    set_repeat_table_header, merge_cells_vertically, set_cell_background,
    set_font, set_cell_width, get_department_acronym
)


def get_ugsn_info(direction_code: str, faculty_fallback: str = "") -> tuple[str, str]:
    """Динамически определяет код и наименование укрупненной группы специальностей."""
    if not direction_code or "." not in direction_code:
        return "01.00.00", faculty_fallback or "Математика и механика"
    prefix = direction_code.split(".")[0]
    ugsn_code = f"{prefix}.00.00"

    ugsn_names = {
        "01": "Математика и механика",
        "02": "Компьютерные и информационные науки",
        "03": "Физика и астрономия",
        "04": "Химия",
        "05": "Науки о земле",
        "08": "Техника и технологии строительства",
        "09": "Информатика и вычислительная техника",
        "10": "Информационная безопасность",
        "11": "Электроника, радиотехника и системы связи",
        "12": "Фотоника, приборостроение, оптические и биотехнические системы и технологии",
        "13": "Электро- и теплоэнергетика",
        "15": "Машиностроение",
        "20": "Техносферная безопасность и природообустройство",
        "27": "Управление в технических системах",
        "38": "Экономика и управление",
        "45": "Языкознание и литературоведение"
    }
    ugsn_name = ugsn_names.get(prefix, faculty_fallback or "Естественные и технические науки")
    return ugsn_code, ugsn_name


def generate_title_page(doc, metadata: dict, subj_info: dict, staff: dict):
    """Страница 1: Титульный лист."""
    start_year = metadata.get("start_year") or "2026"
    add_paragraph_with_spacing(doc, "МИНОБРНАУКИ РОССИИ", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph_with_spacing(doc,
                               "Федеральное государственное бюджетное образовательное учреждение высшего образования",
                               align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_paragraph_with_spacing(doc,
                               "«Ижевский государственный технический университет имени М.Т. Калашникова»",
                               bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

    add_paragraph_with_spacing(doc, "УТВЕРЖДАЮ", bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_paragraph_with_spacing(doc, f"Декан/Директор\n_____________/ {staff['dean']}",
                               align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_paragraph_with_spacing(doc, f"_________________ {start_year} г.",
                               align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=48)

    add_paragraph_with_spacing(doc, "РАБОЧАЯ ПРОГРАММА ДИСЦИПЛИНЫ", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
                               space_after=6)
    add_paragraph_with_spacing(doc, subj_info["name"], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

    add_paragraph_with_spacing(doc,
                               f"направление (специальность): {metadata.get('direction_code')} {metadata.get('direction_name')}",
                               align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph_with_spacing(doc, f"направленность (профиль/программа/специализация): {metadata.get('profile')}",
                               align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph_with_spacing(doc, f"уровень образования: {metadata.get('qualification')}",
                               align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph_with_spacing(doc, f"форма обучения: {metadata.get('education_form')}",
                               align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph_with_spacing(doc,
                               f"общая трудоемкость дисциплины составляет: {subj_info.get('credit_units')} зачетных единиц(ы)",
                               align=WD_ALIGN_PARAGRAPH.CENTER, space_after=144)

    add_paragraph_with_spacing(doc, f"Ижевск {start_year}", align=WD_ALIGN_PARAGRAPH.CENTER)


def generate_compilers_page(doc, metadata: dict, subj_info: dict, staff: dict):
    """Страница 2: Составители и согласование."""
    doc.add_page_break()
    start_year = metadata.get("start_year") or "2026"
    add_paragraph_with_spacing(doc, f"Кафедра «{subj_info.get('department_name', '')}»", bold=True, space_after=18)

    if len(staff["compilers"]) > 1:
        add_paragraph_with_spacing(doc, "Составители:", bold=True, space_after=4)
        for compiler in staff["compilers"]:
            add_paragraph_with_spacing(doc, f"\t{compiler}", space_before=2, space_after=2)
    else:
        comp_name = staff["compilers"][0] if staff["compilers"] else "Преподаватель кафедры"
        add_paragraph_with_spacing(doc, f"Составитель: {comp_name}", space_after=12)

    add_paragraph_with_spacing(doc, space_after=12)
    add_paragraph_with_spacing(doc,
                               f"Рабочая программа составлена в соответствии с требованиями образовательного стандарта ФГОС {metadata.get('fgos_standard')}, "
                               f"рассмотрена и одобрена на заседании кафедры.")

    add_paragraph_with_spacing(doc, f"Протокол от «____» ________________ {start_year} г. №_______", space_after=24)
    hod_role = staff.get("head_of_department_role") or "Заведующий кафедрой"
    hod_role_cap = hod_role[0].upper() + hod_role[1:] if hod_role else "Заведующий кафедрой"

    add_paragraph_with_spacing(doc, f"{hod_role_cap} __________________ {staff['head_of_department']}",
                               space_after=24)

    add_paragraph_with_spacing(doc, "СОГЛАСОВАНО", bold=True, space_after=12)
    add_paragraph_with_spacing(doc,
                               f"Количество часов рабочей программы и формируемые компетенции соответствуют учебному плану "
                               f"{metadata.get('direction_code')} «{metadata.get('direction_name')}» "
                               f"(профиль «{metadata.get('profile')}»)", space_after=18)

    # Динамическое определение УГСН
    dir_code = metadata.get("direction_code") or "01.03.04"
    faculty_name = metadata.get("faculty") or ""
    ugsn_code, ugsn_name = get_ugsn_info(dir_code, faculty_name)

    add_paragraph_with_spacing(doc, f"Протокол заседания учебно-методической комиссии по УГСН\n\n"
                                    f"{ugsn_code} «{ugsn_name}» от «____» _______________ {start_year} г. №_______",
                               space_after=24)

    add_paragraph_with_spacing(doc, f"Председатель учебно-методической комиссии по УГСН\n\n"
                                    f"{ugsn_code} «{ugsn_name}» _________________ {staff['umk_chairman']}",
                               space_after=24)

    add_paragraph_with_spacing(doc,
                               f"Руководитель образовательной программы _________________ {staff['program_director']}")


def generate_annotation_page(doc, metadata: dict, subj_info: dict, subj_ai: dict, mapped_comp: dict,
                             comp_registry: dict):
    """Страница 3: Аннотация дисциплины."""
    doc.add_page_break()
    add_paragraph_with_spacing(doc, "Аннотация к дисциплине", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
                               space_after=12)

    table_annot = doc.add_table(rows=9, cols=2)
    table_annot.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_annot.style = "Table Grid"

    total_hours_sum = subj_info.get("total_hours", {}).get("total", 0)

    # Сбор компетенций
    comp_texts_list = []
    for c_code in mapped_comp.keys():
        c_desc = comp_registry.get(c_code, {}).get('competency_text', '')
        comp_texts_list.append(f"{c_code}. {c_desc}")
    competencies_val = "\n".join(comp_texts_list)

    # Сбор разделов
    sections_list = []
    for s in subj_ai["thematic_plan"]["sections"]:
        sections_list.append(f"Раздел {s['number']}. {s['name']}")
    sections_val = "\n".join(sections_list)

    # Сбор форм промежуточной аттестации
    control_forms_str_list = []
    if subj_info["control_forms"]["exams"]:
        control_forms_str_list.append(f"Экзамен ({', '.join(map(str, subj_info['control_forms']['exams']))} семестр)")
    if subj_info["control_forms"]["graded_credits"]:
        control_forms_str_list.append(
            f"Зачет с оценкой ({', '.join(map(str, subj_info['control_forms']['graded_credits']))} семестр)")
    if subj_info["control_forms"]["credits"]:
        control_forms_str_list.append(f"Зачет ({', '.join(map(str, subj_info['control_forms']['credits']))} семестр)")
    control_val = ", ".join(control_forms_str_list)

    annot_data = [
        ("Название дисциплины", subj_info["name"]),
        ("Направление подготовки (специальность)",
         f"{metadata.get('direction_code')} {metadata.get('direction_name')}"),
        ("Направленность (профиль/программа/специализация)", metadata.get('profile', '')),
        ("Место дисциплины",
         f"{subj_info.get('structure', {}).get('part', 'Обязательная часть')} Блока 1. Дисциплины (модули)"),
        ("Трудоемкость (з.е. / часы)", f"{subj_info.get('credit_units')} з.е. / {total_hours_sum} часов"),
        ("Цель изучения дисциплины", subj_ai["pedagogical_frame"]["goals"]),
        ("Компетенции, формируемые в результате освоения дисциплины", competencies_val),
        ("Содержание дисциплины (основные разделы и темы)", sections_val),
        ("Форма промежуточной аттестации", control_val)
    ]

    for i, (label, val) in enumerate(annot_data):
        row = table_annot.rows[i]
        set_row_cant_split(row)
        set_cell_text(row.cells[0], label, bold=True, size_pt=11)
        set_cell_text(row.cells[1], val, size_pt=11)
        set_cell_width(row.cells[0], 5.0)
        set_cell_width(row.cells[1], 11.5)


def generate_sections_1_2(doc, subj_ai: dict, mapped_comp: dict, comp_registry: dict):
    """Страница 4: Разделы 1 и 2 РПД (Цели, Задачи, Компетенции) со слиянием ячеек."""
    doc.add_page_break()
    add_paragraph_with_spacing(doc, "1.  Цели и задачи дисциплины:", bold=True, space_after=12)
    add_paragraph_with_spacing(doc, f"Целью освоения дисциплины является: {subj_ai['pedagogical_frame']['goals']}")
    add_paragraph_with_spacing(doc, "Задачи дисциплины:", bold=True)
    for task in subj_ai['pedagogical_frame']['tasks']:
        add_paragraph_with_spacing(doc, f"— {task}")

    add_paragraph_with_spacing(doc, "2. Планируемые результаты обучения", bold=True, space_after=12)
    add_paragraph_with_spacing(doc, "В результате освоения дисциплины у студента должны быть сформированы:")

    # Сбор уникальных глобальных списков ЗУН произвольной длины из структуры ИИ
    knowledge_list = subj_ai["pedagogical_frame"].get("knowledge_list", [])
    skills_list = subj_ai["pedagogical_frame"].get("skills_list", [])
    abilities_list = subj_ai["pedagogical_frame"].get("abilities_list", [])

    # Построение маппингов для точного определения индексов
    knowledge_map = {text: idx for idx, text in enumerate(knowledge_list, start=1)}
    skills_map = {text: idx for idx, text in enumerate(skills_list, start=1)}
    abilities_map = {text: idx for idx, text in enumerate(abilities_list, start=1)}

    # Таблица Знаний
    add_paragraph_with_spacing(doc, "Знания, приобретаемые в ходе освоения дисциплины:", bold=True, space_before=12,
                               space_after=6)
    table_k = doc.add_table(rows=1, cols=2)
    table_k.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_k.style = "Table Grid"
    set_cell_text(table_k.rows[0].cells[0], "№ п/п", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, fill_hex="F2F2F2")
    set_cell_text(table_k.rows[0].cells[1], "Знания", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, fill_hex="F2F2F2")
    set_cell_width(table_k.rows[0].cells[0], 1.5)
    set_cell_width(table_k.rows[0].cells[1], 14.5)
    for idx_k, k_text in enumerate(knowledge_list, start=1):
        row_cells = table_k.add_row().cells
        set_row_cant_split(table_k.rows[-1])
        set_cell_text(row_cells[0], str(idx_k), align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)
        set_cell_text(row_cells[1], k_text, size_pt=10)
        set_cell_width(row_cells[0], 1.5)
        set_cell_width(row_cells[1], 14.5)

    # Таблица Умений
    add_paragraph_with_spacing(doc, "Умения, приобретаемые в ходе освоения дисциплины:", bold=True, space_before=12,
                               space_after=6)
    table_s = doc.add_table(rows=1, cols=2)
    table_s.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_s.style = "Table Grid"
    set_cell_text(table_s.rows[0].cells[0], "№ п/п", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, fill_hex="F2F2F2")
    set_cell_text(table_s.rows[0].cells[1], "Умения", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, fill_hex="F2F2F2")
    set_cell_width(table_s.rows[0].cells[0], 1.5)
    set_cell_width(table_s.rows[0].cells[1], 14.5)
    for idx_s, s_text in enumerate(skills_list, start=1):
        row_cells = table_s.add_row().cells
        set_row_cant_split(table_s.rows[-1])
        set_cell_text(row_cells[0], str(idx_s), align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)
        set_cell_text(row_cells[1], s_text, size_pt=10)
        set_cell_width(row_cells[0], 1.5)
        set_cell_width(row_cells[1], 14.5)

    # Таблица Навыков
    add_paragraph_with_spacing(doc, "Навыки, приобретаемые в ходе освоения дисциплины:", bold=True, space_before=12,
                               space_after=6)
    table_a = doc.add_table(rows=1, cols=2)
    table_a.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_a.style = "Table Grid"
    set_cell_text(table_a.rows[0].cells[0], "№ п/п", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, fill_hex="F2F2F2")
    set_cell_text(table_a.rows[0].cells[1], "Навыки", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, fill_hex="F2F2F2")
    set_cell_width(table_a.rows[0].cells[0], 1.5)
    set_cell_width(table_a.rows[0].cells[1], 14.5)
    for idx_a, a_text in enumerate(abilities_list, start=1):
        row_cells = table_a.add_row().cells
        set_row_cant_split(table_a.rows[-1])
        set_cell_text(row_cells[0], str(idx_a), align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)
        set_cell_text(row_cells[1], a_text, size_pt=10)
        set_cell_width(row_cells[0], 1.5)
        set_cell_width(row_cells[1], 14.5)

    # Таблица распределения ЗУН по индикаторам (диагональная структура)
    add_paragraph_with_spacing(doc, "Компетенции, приобретаемые в ходе освоения дисциплины:", bold=True,
                               space_before=18)
    table_comp = doc.add_table(rows=1, cols=5)
    table_comp.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_comp.style = "Table Grid"
    set_repeat_table_header(table_comp.rows[0])
    set_row_cant_split(table_comp.rows[0])

    comp_headers = ["Компетенции", "Индикаторы", "Знания", "Умения", "Навыки"]
    for c_idx, h_text in enumerate(comp_headers):
        set_cell_text(table_comp.rows[0].cells[c_idx], h_text, bold=True, size_pt=11, align=WD_ALIGN_PARAGRAPH.CENTER,
                      fill_hex="F2F2F2")

    for c_code, ind_list in mapped_comp.items():
        comp_text = f"{c_code}. {comp_registry.get(c_code, {}).get('competency_text', '')}"
        for i_code in ind_list:
            row_cells = table_comp.add_row().cells
            set_row_cant_split(table_comp.rows[-1])
            ind_desc = comp_registry.get(c_code, {}).get('indicators', {}).get(i_code, {}).get('indicator_text', '')
            ind_text = f"{i_code}. {ind_desc}"

            mapping_entry = next((m for m in subj_ai["pedagogical_frame"].get("indicator_mappings", []) if
                                  m["indicator_code"] == i_code), None)

            cell_k = ""
            cell_s = ""
            cell_a = ""

            if mapping_entry:
                k_idx = mapping_entry.get("knowledge_indices", [])
                if k_idx:
                    cell_k = f"{min(k_idx)}-{max(k_idx)}" if len(k_idx) > 1 else f"{k_idx[0]}-{k_idx[0]}"

                s_idx = mapping_entry.get("skills_indices", [])
                if s_idx:
                    cell_s = f"{min(s_idx)}-{max(s_idx)}" if len(s_idx) > 1 else f"{s_idx[0]}-{s_idx[0]}"

                a_idx = mapping_entry.get("abilities_indices", [])
                if a_idx:
                    cell_a = f"{min(a_idx)}-{max(a_idx)}" if len(a_idx) > 1 else f"{a_idx[0]}-{a_idx[0]}"

            set_cell_text(row_cells[0], comp_text, size_pt=10)
            set_cell_text(row_cells[1], ind_text, size_pt=10)
            set_cell_text(row_cells[2], cell_k, size_pt=10, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(row_cells[3], cell_s, size_pt=10, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(row_cells[4], cell_a, size_pt=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Применение вертикального объединения дубликатов компетенций в первом столбце
    merge_cells_vertically(table_comp, 0)


def generate_section_3(doc, subj_info: dict, subj_ai: dict):
    """Страница 5: Раздел 3 Место дисциплины в структуре ООП с расчетом курсов."""
    add_paragraph_with_spacing(doc)
    add_paragraph_with_spacing(doc, "3. Место дисциплины в структуре ООП", bold=True, space_after=12)
    add_paragraph_with_spacing(doc,
                               f"Дисциплина относится к обязательной части Блока 1 «Дисциплины (модули)» ООП.")

    # Динамический расчет курсов обучения на основе семестров
    sems = sorted([int(s) for s in subj_info.get("load_by_semester", {}).keys()])
    courses = sorted(list(set([(s + 1) // 2 for s in sems])))
    sems_str = ", ".join(map(str, sems))
    courses_str = ", ".join(map(str, courses))

    add_paragraph_with_spacing(doc, f"Дисциплина изучается на {courses_str} курсе(ах) в {sems_str} семестре(ах).")

    add_paragraph_with_spacing(doc,
                               f"Изучение дисциплины базируется на знаниях, умениях и навыках, полученных при изучении школьной программы.")
    add_paragraph_with_spacing(doc, f"Последующие дисциплины: {subj_ai['pedagogical_frame']['postrequisites_text']}")


def get_questions_for_competency(c_code: str, discipline_name: str, subj_ai: dict) -> list:
    """Возвращает набор из 5 тестовых вопросов по компетенции из JSON-файла ИИ-генератора."""
    competency_tests = subj_ai.get("resources_and_evaluation", {}).get("competency_tests", [])
    for test_block in competency_tests:
        if test_block.get("competency_code") == c_code:
            questions = test_block.get("questions", [])
            if len(questions) >= 5:
                return questions[:5]
    return []


def generate_fos_appendix(doc, metadata: dict, subj_info: dict, subj_ai: dict, mapped_comp: dict, comp_registry: dict,
                          sems_active: list):
    """Генерация приложения оценочных средств (ФОС) из структурированных JSON-данных."""
    doc.add_page_break()
    start_year = metadata.get("start_year") or "2026"
    add_paragraph_with_spacing(doc, "Приложение к рабочей программе дисциплины", italic=True,
                               align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=24)

    add_paragraph_with_spacing(doc, "МИНОБРНАУКИ РОССИИ", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph_with_spacing(doc,
                               "Федеральное государственное бюджетное образовательное учреждение высшего образования",
                               align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph_with_spacing(doc, "«Ижевский государственный технический университет имени М.Т. Калашникова»",
                               bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=36)

    add_paragraph_with_spacing(doc, "Оценочные средства по дисциплине", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
                               space_after=6)
    add_paragraph_with_spacing(doc, subj_info["name"], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

    add_paragraph_with_spacing(doc,
                               f"направление (специальность) {metadata.get('direction_code')} {metadata.get('direction_name')}",
                               align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph_with_spacing(doc, f"направленность (профиль) {metadata.get('profile')}",
                               align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph_with_spacing(doc, f"уровень образования: {metadata.get('qualification')}",
                               align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph_with_spacing(doc, f"форма обучения: {metadata.get('education_form')}",
                               align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    add_paragraph_with_spacing(doc,
                               f"общая трудоемкость дисциплины составляет: {subj_info.get('credit_units')} зачетных единиц(ы)",
                               align=WD_ALIGN_PARAGRAPH.CENTER)

    # Перенос таблицы на новую страницу
    doc.add_page_break()
    add_paragraph_with_spacing(doc, "1. Оценочные средства", bold=True, space_after=12)
    add_paragraph_with_spacing(doc,
                               """Оценивание формирования компетенций производится на основе результатов обучения, приведенных в п. 2 рабочей программы и ФОС. Связь разделов компетенций, индикаторов и форм контроля (текущего и промежуточного) указаны в таблице 4.2 рабочей программы дисциплины.
Оценочные средства соотнесены с результатами обучения по дисциплине и индикаторами достижения компетенций, представлены ниже.""",
                               space_after=12)

    table_fos_map = doc.add_table(rows=1, cols=4)
    table_fos_map.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_fos_map.style = "Table Grid"
    set_repeat_table_header(table_fos_map.rows[0])

    fos_map_headers = ["№ п/п", "Коды компетенций и индикаторов", "Результат обучения (знания, умения и навыки)",
                       "Формы текущего и промежуточного контроля"]
    for idx_f, f_text in enumerate(fos_map_headers):
        set_cell_text(table_fos_map.rows[0].cells[idx_f], f_text, bold=True, size_pt=10,
                      align=WD_ALIGN_PARAGRAPH.CENTER, fill_hex="F2F2F2")

    # Считывание пулов для сопоставления индексов в таблице 1 ФОС
    knowledge_list = subj_ai["pedagogical_frame"].get("knowledge_list", [])
    skills_list = subj_ai["pedagogical_frame"].get("skills_list", [])
    abilities_list = subj_ai["pedagogical_frame"].get("abilities_list", [])

    idx_row = 1
    for c_code, ind_list in mapped_comp.items():
        for i_code in ind_list:
            r_cells = table_fos_map.add_row().cells
            set_row_cant_split(table_fos_map.rows[-1])

            set_cell_text(r_cells[0], str(idx_row), align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)

            # Объединение кода и описания индикатора в ячейке
            ind_desc = comp_registry.get(c_code, {}).get('indicators', {}).get(i_code, {}).get('indicator_text', '')
            full_indicator_text = f"{i_code}. {ind_desc}" if ind_desc else i_code
            set_cell_text(r_cells[1], full_indicator_text, size_pt=10, align=WD_ALIGN_PARAGRAPH.LEFT)

            # Группировка ЗУН по конкретному индикатору со ссылкой на глобальные индексы
            ksa_strings = []
            mapping_entry = next((m for m in subj_ai["pedagogical_frame"].get("indicator_mappings", []) if
                                  m["indicator_code"] == i_code), None)

            if mapping_entry:
                for k_idx in mapping_entry.get("knowledge_indices", []):
                    if 0 < k_idx <= len(knowledge_list):
                        ksa_strings.append(f"З{k_idx}: {knowledge_list[k_idx - 1]}")
                for s_idx in mapping_entry.get("skills_indices", []):
                    if 0 < s_idx <= len(skills_list):
                        ksa_strings.append(f"У{s_idx}: {skills_list[s_idx - 1]}")
                for a_idx in mapping_entry.get("abilities_indices", []):
                    if 0 < a_idx <= len(abilities_list):
                        ksa_strings.append(f"Н{a_idx}: {abilities_list[a_idx - 1]}")

            total_hours_dict = subj_info.get("total_hours", {})
            labs_h = total_hours_dict.get("laboratory_works", 0)
            practicals_h = total_hours_dict.get("practical_classes", 0)

            controls = []
            if labs_h > 0:
                controls.append("Защита лабораторных работ")
            if practicals_h > 0:
                controls.append("Защита практических работ")
            control_text = "\n".join(controls) if controls else "Устный опрос"

            set_cell_text(r_cells[2], "\n".join(ksa_strings), size_pt=9)
            set_cell_text(r_cells[3], control_text, size_pt=10)
            idx_row += 1

    # Вертикальное объединение
    merge_cells_vertically(table_fos_map, 1)

    # Вывод вопросов по каждой активной форме промежуточной аттестации
    add_paragraph_with_spacing(doc, space_before=18)
    control_forms = subj_info.get("control_forms", {})
    form_types = [
        ("exams", "экзамен", "экзамена"),
        ("graded_credits", "зачет с оценкой", "зачета с оценкой"),
        ("credits", "зачет", "зачета")
    ]

    for key, form_name_nom, form_name_gen in form_types:
        sems = control_forms.get(key, [])
        if sems:
            add_paragraph_with_spacing(doc, f"Наименование: {form_name_nom}", bold=True, space_before=12)
            add_paragraph_with_spacing(doc, "Представление в ФОС: перечень вопросов", italic=True)
            add_paragraph_with_spacing(doc, f"Перечень вопросов для проведения {form_name_gen}:", bold=True,
                                       space_after=6)

            for sem in sems:
                add_paragraph_with_spacing(doc, f"{sem} семестр:", italic=True, bold=True, space_before=6)

                # Чтение контрольных вопросов из JSON
                questions = subj_ai.get("resources_and_evaluation", {}).get("control_questions", [])
                for q_idx, q_text in enumerate(questions, start=1):
                    add_paragraph_with_spacing(doc, f"{q_idx}. {q_text}", space_before=2, space_after=2)

            # Обязательная надпись без ответов
            add_paragraph_with_spacing(doc, "Критерии оценки:", bold=True, space_before=6)
            add_paragraph_with_spacing(doc, "Приведены в разделе 2.", space_after=12)

    # Тестирование с ключами к тестам ДЛЯ КАЖДОЙ КОМПЕТЕНЦИИ ИЗ JSON
    add_paragraph_with_spacing(doc, space_before=18)
    add_paragraph_with_spacing(doc, "Наименование: проверочный тест", bold=True)
    add_paragraph_with_spacing(doc, "Представление в ФОС: набор вопросов для проведения тестирования", italic=True)

    competency_tests = subj_ai.get("resources_and_evaluation", {}).get("competency_tests", [])

    for c_test in competency_tests:
        c_code = c_test.get("competency_code", "")
        comp_text = comp_registry.get(c_code, {}).get('competency_text', '')
        add_paragraph_with_spacing(doc, f"Компетенция {c_code}. {comp_text}", bold=True, space_before=12,
                                   space_after=6)

        questions = c_test.get("questions", [])
        for idx_q, q_item in enumerate(questions, start=1):
            add_paragraph_with_spacing(doc, f"{idx_q}. {q_item['question']}", space_before=4)
            for opt_idx, option in enumerate(q_item["options"], start=1):
                add_paragraph_with_spacing(doc, f"   {opt_idx}) {option}", space_before=1, space_after=1)

        # Вывод таблицы Ключей для теста данной компетенции
        add_paragraph_with_spacing(doc, "Ключи теста:", bold=True, space_before=6, space_after=6)
        num_cols = len(questions) + 1
        table_keys = doc.add_table(rows=2, cols=num_cols)
        table_keys.alignment = WD_TABLE_ALIGNMENT.CENTER
        table_keys.style = "Table Grid"

        set_cell_text(table_keys.rows[0].cells[0], "Вопрос", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
                      fill_hex="F2F2F2")
        set_cell_text(table_keys.rows[1].cells[0], "Ответ", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
                      fill_hex="F2F2F2")

        for idx_tk, q_item in enumerate(questions, start=1):
            set_cell_text(table_keys.rows[0].cells[idx_tk], str(idx_tk), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(table_keys.rows[1].cells[idx_tk], str(q_item["correct_answer"]),
                          align=WD_ALIGN_PARAGRAPH.CENTER)

    # Генерация примера билета (при наличии экзамена или дифференцированного зачета)
    has_exam = len(control_forms.get("exams", [])) > 0
    has_graded = len(control_forms.get("graded_credits", [])) > 0

    if has_exam or has_graded:
        add_paragraph_with_spacing(doc, space_before=18)
        ticket_label = "Пример экзаменационного билета:" if has_exam else "Пример билета для зачета с оценкой:"
        add_paragraph_with_spacing(doc, ticket_label, bold=True, space_after=12)

        table_ticket = doc.add_table(rows=1, cols=1)
        table_ticket.alignment = WD_TABLE_ALIGNMENT.CENTER
        table_ticket.style = "Table Grid"
        set_row_cant_split(table_ticket.rows[0])

        cell_ticket = table_ticket.rows[0].cells[0]
        cell_ticket.width = Cm(15.0)

        p_ticket = cell_ticket.paragraphs[0]
        p_ticket.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run_h1 = p_ticket.add_run(
            "\nФГБОУ ВО «Ижевский государственный технический университет имени М.Т. Калашникова»\n\n")
        set_font(run_h1, size_pt=9)

        ticket_num_text = "ЭКЗАМЕНАЦИОННЫЙ БИЛЕТ № 1\n" if has_exam else "ОЦЕНОЧНЫЙ БИЛЕТ ДЛЯ ЗАЧЕТА С ОЦЕНКОЙ № 1\n"
        run_title = p_ticket.add_run(ticket_num_text)
        set_font(run_title, size_pt=12, bold=True)

        run_sub = p_ticket.add_run(
            f"по дисциплине «{subj_info['name']}»\nдля направления {metadata.get('direction_code')} «{metadata.get('direction_name')}»\n\n")
        set_font(run_sub, size_pt=10, italic=True)

        # Чтение вопросов для билета из ФОС JSON
        questions_ticket = subj_ai.get("resources_and_evaluation", {}).get("control_questions", [])[:3]
        if len(questions_ticket) < 3:
            questions_ticket = [
                "Теоретический вопрос по первому разделу дисциплины.",
                "Теоретический вопрос по второму разделу дисциплины.",
                "Практическое задание на применение изученных алгоритмов и методов."
            ]

        for idx_q, q_text in enumerate(questions_ticket, start=1):
            run_q = p_ticket.add_run(f"{idx_q}. {q_text}\n")
            set_font(run_q, size_pt=11)

        dept_name = subj_info.get("department_name") or metadata.get("department") or "Прикладная математика"
        dept_abbr = get_department_acronym(dept_name)

        run_footer = p_ticket.add_run(
            f"\nБилет рассмотрен на заседании кафедры {dept_abbr} от «____» _______________ {start_year} г.\n")
        set_font(run_footer, size_pt=9, italic=True)

    add_paragraph_with_spacing(doc)
    generate_evaluation_criteria_section(doc, subj_info, subj_ai)


def generate_evaluation_criteria_section(doc, subj_info: dict, subj_ai: dict):
    """Генерирует раздел '2. Критерии и шкалы оценивания' с динамическим БРС-распределением."""
    add_paragraph_with_spacing(doc, space_before=18)
    add_paragraph_with_spacing(doc, "2. Критерии и шкалы оценивания", bold=True, space_after=12)
    add_paragraph_with_spacing(doc,
                               "Результат обучения по дисциплине считается достигнутым при успешном прохождении "
                               "обучающимся всех контрольных мероприятий, относящихся к данному результату обучения.",
                               space_after=12
                               )

    # 1. Сбор количества лабораторных и практических работ
    num_labs = len(subj_ai.get("thematic_plan", {}).get("labs", []))
    num_pracs = len(subj_ai.get("thematic_plan", {}).get("practicals", []))

    # Динамическое распределение баллов БРС (Всего: max = 100, min = 80)
    events = []
    test_max = 20
    remaining_points = 80

    if num_labs > 0 and num_pracs > 0:
        labs_share = 40
        pracs_share = 40
    elif num_labs > 0:
        labs_share = 80
        pracs_share = 0
    elif num_pracs > 0:
        labs_share = 0
        pracs_share = 80
    else:
        labs_share = 0
        pracs_share = 0
        # Если работы отсутствуют, создаем 4 дефолтных этапа контроля успеваемости
        default_events = 4
        each_max = remaining_points // default_events
        remainder = remaining_points % default_events
        for i in range(1, default_events + 1):
            m_val = each_max + (1 if i <= remainder else 0)
            events.append({
                "name": f"Защита устного опроса / собеседования №{i}",
                "max": m_val,
                "min": int(round(0.8 * m_val))
            })

    # Расчет баллов для лабораторных
    if labs_share > 0:
        each_max = labs_share // num_labs
        remainder = labs_share % num_labs
        for i in range(1, num_labs + 1):
            m_val = each_max + (1 if i <= remainder else 0)
            events.append({
                "name": f"Защита лабораторной работы №{i}",
                "max": m_val,
                "min": int(round(0.8 * m_val))
            })

    # Расчет баллов для практических
    if pracs_share > 0:
        each_max = pracs_share // num_pracs
        remainder = pracs_share % num_pracs
        for i in range(1, num_pracs + 1):
            m_val = each_max + (1 if i <= remainder else 0)
            events.append({
                "name": f"Защита практической работы №{i}",
                "max": m_val,
                "min": int(round(0.8 * m_val))
            })

    # Добавление финального проверочного теста
    events.append({
        "name": "Проверочное тестирование",
        "max": test_max,
        "min": int(round(0.8 * test_max))
    })

    # Точная корректировка суммы min до 80 во избежание погрешностей округления
    total_min = sum(e["min"] for e in events)
    diff_min = 80 - total_min
    if diff_min != 0 and len(events) > 0:
        events[-1]["min"] += diff_min

    # Создание таблицы БРС
    table_brs = doc.add_table(rows=1, cols=4)
    table_brs.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_brs.style = "Table Grid"

    hdr_cells = table_brs.rows[0].cells
    set_cell_text(hdr_cells[0], "Разделы дисциплины", bold=True, size_pt=10, align=WD_ALIGN_PARAGRAPH.CENTER,
                  fill_hex="F2F2F2")
    set_cell_text(hdr_cells[1], "Форма контроля", bold=True, size_pt=10, align=WD_ALIGN_PARAGRAPH.CENTER,
                  fill_hex="F2F2F2")
    set_cell_text(hdr_cells[2], "Количество баллов (min)", bold=True, size_pt=10, align=WD_ALIGN_PARAGRAPH.CENTER,
                  fill_hex="F2F2F2")
    set_cell_text(hdr_cells[3], "Количество баллов (max)", bold=True, size_pt=10, align=WD_ALIGN_PARAGRAPH.CENTER,
                  fill_hex="F2F2F2")

    num_sections = len(subj_ai.get("thematic_plan", {}).get("sections", []))

    for idx_ev, ev in enumerate(events, start=1):
        row_cells = table_brs.add_row().cells
        set_row_cant_split(table_brs.rows[-1])

        # Равномерное сопоставление контрольных точек с разделами плана
        if idx_ev <= num_sections:
            section_lbl = str(idx_ev)
        elif idx_ev == len(events):
            section_lbl = "Все" if num_sections == 0 else f"1-{num_sections}"
        else:
            section_lbl = str(num_sections) if num_sections > 0 else "1"

        set_cell_text(row_cells[0], section_lbl, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)
        set_cell_text(row_cells[1], ev["name"], size_pt=10)
        set_cell_text(row_cells[2], str(ev["min"]), align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)
        set_cell_text(row_cells[3], str(ev["max"]), align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)

    # Строка Итого в БРС
    total_row = table_brs.add_row().cells
    set_row_cant_split(table_brs.rows[-1])
    set_cell_text(total_row[0], "", size_pt=10)
    set_cell_text(total_row[1], "Итого", bold=True, size_pt=10)
    set_cell_text(total_row[2], "80", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)
    set_cell_text(total_row[3], "100", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)

    # Пояснительный текст под таблицей БРС
    add_paragraph_with_spacing(doc, space_before=12)
    add_paragraph_with_spacing(doc,
                               "При оценивании результатов обучения по дисциплине в ходе текущего контроля "
                               "успеваемости используются следующие критерии. Минимальное количество баллов выставляется "
                               "обучающемуся при выполнении всех показателей, допускаются несущественные неточности "
                               "в изложении и оформлении материала.",
                               space_after=12
                               )

    # Таблица показателей выполнения работ
    table_indicators = doc.add_table(rows=1, cols=2)
    table_indicators.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_indicators.style = "Table Grid"

    ind_hdr = table_indicators.rows[0].cells
    set_cell_text(ind_hdr[0], "Наименование, обозначение", bold=True, size_pt=10, align=WD_ALIGN_PARAGRAPH.CENTER,
                  fill_hex="F2F2F2")
    set_cell_text(ind_hdr[1], "Показатели выставления минимального количества баллов", bold=True, size_pt=10,
                  align=WD_ALIGN_PARAGRAPH.CENTER, fill_hex="F2F2F2")
    set_cell_width(ind_hdr[0], 4.0)
    set_cell_width(ind_hdr[1], 12.0)

    if num_labs > 0:
        row_l = table_indicators.add_row().cells
        set_row_cant_split(table_indicators.rows[-1])
        set_cell_text(row_l[0], "Лабораторная работа", bold=True, size_pt=10)
        set_cell_text(row_l[1],
                      "Лабораторная работа выполнена в полном объеме;\n"
                      "Представлен отчет, содержащий необходимые расчеты, выводы, оформленный в соответствии с установленными требованиями;\n"
                      "Продемонстрирован удовлетворительный уровень владения материалом при защите лабораторной работы, даны правильные ответы не менее чем на 50% заданных вопросов.",
                      size_pt=10
                      )
        set_cell_width(row_l[0], 4.0)
        set_cell_width(row_l[1], 12.0)

    if num_pracs > 0:
        row_p = table_indicators.add_row().cells
        set_row_cant_split(table_indicators.rows[-1])
        set_cell_text(row_p[0], "Практическая работа", bold=True, size_pt=10)
        set_cell_text(row_p[1],
                      "Продемонстрирован удовлетворительный уровень владения материалом.\n"
                      "Правильно решено не менее 50% заданий.",
                      size_pt=10
                      )
        set_cell_width(row_p[0], 4.0)
        set_cell_width(row_p[1], 12.0)

    if num_labs == 0 and num_pracs == 0:
        row_o = table_indicators.add_row().cells
        set_row_cant_split(table_indicators.rows[-1])
        set_cell_text(row_o[0], "Собеседование", bold=True, size_pt=10)
        set_cell_text(row_o[1],
                      "Продемонстрирован достаточный уровень усвоения теоретических понятий.\n"
                      "Даны правильные и структурированные ответы на контрольные вопросы.",
                      size_pt=10
                      )
        set_cell_width(row_o[0], 4.0)
        set_cell_width(row_o[1], 12.0)

    # 3. Формирование динамических описаний промежуточной аттестации
    control_forms = subj_info.get("control_forms", {})
    exams_sems = control_forms.get("exams", [])
    gc_sems = control_forms.get("graded_credits", [])
    c_sems = control_forms.get("credits", [])

    control_parts = []
    if exams_sems:
        sem_str = ", ".join(map(str, exams_sems))
        suffix = "семестр" if len(exams_sems) == 1 else "семестры"
        control_parts.append(f"экзамена ({sem_str} {suffix})")
    if gc_sems:
        sem_str = ", ".join(map(str, gc_sems))
        suffix = "семестр" if len(gc_sems) == 1 else "семестры"
        control_parts.append(f"зачета с оценкой ({sem_str} {suffix})")
    if c_sems:
        sem_str = ", ".join(map(str, c_sems))
        suffix = "семестр" if len(c_sems) == 1 else "семестры"
        control_parts.append(f"зачета ({sem_str} {suffix})")

    if len(control_parts) > 1:
        control_str = ", ".join(control_parts[:-1]) + " и " + control_parts[-1]
    elif len(control_parts) == 1:
        control_str = control_parts[0]
    else:
        control_str = "экзамена"

    add_paragraph_with_spacing(doc, space_before=12)
    add_paragraph_with_spacing(doc,
                               f"Промежуточная аттестация по дисциплине проводится в форме {control_str}.\n"
                               "Итоговая оценка по дисциплине может быть выставлена на основе результатов текущего "
                               "контроля с использованием следующей шкалы:",
                               space_after=6
                               )

    # Вывод шкалы БРС (дифференцированная vs недифференцированная)
    table_scale = doc.add_table(rows=1, cols=2)
    table_scale.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_scale.style = "Table Grid"

    scale_hdr = table_scale.rows[0].cells
    set_cell_text(scale_hdr[0], "Оценка", bold=True, size_pt=10, align=WD_ALIGN_PARAGRAPH.CENTER, fill_hex="F2F2F2")
    set_cell_text(scale_hdr[1], "Набрано баллов", bold=True, size_pt=10, align=WD_ALIGN_PARAGRAPH.CENTER,
                  fill_hex="F2F2F2")

    scales = []
    if exams_sems or gc_sems:
        scales.append(("«отлично»", "90-100"))
        scales.append(("«хорошо»", "75-89"))
        scales.append(("«удовлетворительно»", "50-74"))
        scales.append(("«неудовлетворительно»", "0-49"))
    else:
        scales.append(("«зачтено»", "50-100"))
        scales.append(("«не зачтено»", "0-49"))

    for grade, points in scales:
        row_s = table_scale.add_row().cells
        set_row_cant_split(table_scale.rows[-1])
        set_cell_text(row_s[0], grade, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)
        set_cell_text(row_s[1], points, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)

    add_paragraph_with_spacing(doc, space_before=12)
    add_paragraph_with_spacing(doc,
                               "Если сумма набранных баллов менее 50 – обучающийся не допускается до промежуточной аттестации.")
    add_paragraph_with_spacing(doc,
                               "Если сумма баллов составляет от 50 до 100 баллов, обучающийся допускается до промежуточной аттестации.",
                               space_after=6)

    ticket_types = []
    if gc_sems:
        ticket_types.append("зачету с оценкой")
    if exams_sems:
        ticket_types.append("экзамену")
    if c_sems and not (gc_sems or exams_sems):
        ticket_types.append("зачету")

    ticket_types_str = ", ".join(ticket_types)
    if len(ticket_types) > 1:
        ticket_types_str = ", ".join(ticket_types[:-1]) + " и " + ticket_types[-1]

    ticket_desc = f"Билет к {ticket_types_str} включает 1 теоретический и 2 практических задания."

    add_paragraph_with_spacing(doc,
                               f"{ticket_desc} Промежуточная аттестация проводится в письменной форме. "
                               "Время на подготовку: 60-90 минут. При оценивании результатов обучения по дисциплине в ходе "
                               "промежуточной аттестации используются следующие критерии и шкала оценки:",
                               space_after=6
                               )

    # Таблица детальных критериев оценивания ответов на билеты
    table_desc = doc.add_table(rows=1, cols=2)
    table_desc.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_desc.style = "Table Grid"

    desc_hdr = table_desc.rows[0].cells
    set_cell_text(desc_hdr[0], "Оценка", bold=True, size_pt=10, align=WD_ALIGN_PARAGRAPH.CENTER, fill_hex="F2F2F2")
    set_cell_text(desc_hdr[1], "Критерии оценки", bold=True, size_pt=10, align=WD_ALIGN_PARAGRAPH.CENTER,
                  fill_hex="F2F2F2")
    set_cell_width(desc_hdr[0], 4.0)
    set_cell_width(desc_hdr[1], 12.0)

    criteria_entries = []
    if exams_sems or gc_sems:
        criteria_entries.append((
            "«отлично»",
            "Обучающийся показал всестороннее, систематическое и глубокое знание учебного материала, "
            "предусмотренного программой, умение уверенно применять на их практике при решении задач (выполнении заданий), "
            "способность полно, правильно и аргументировано отвечать на вопросы и делать необходимые выводы. "
            "Свободно использует основную литературу и знаком с дополнительной литературой, рекомендованной программой."
        ))
        criteria_entries.append((
            "«хорошо»",
            "Обучающийся показал полное знание теоретического материала, владение основной литературой, "
            "рекомендованной в программе, умение самостоятельно решать задачи (выполнять задания), "
            "способность аргументировано отвечать на вопросы и делать необходимые выводы, допускает единичные ошибки, "
            "исправляемые после замечания преподавателя. Способен к самостоятельному пополнению и обновлению "
            "знаний в ходе дальнейшей учебной работы и профессиональной деятельности."
        ))
        criteria_entries.append((
            "«удовлетворительно»",
            "Обучающийся демонстрирует неполное или фрагментарное знание основного учебного материала, "
            "допускает существенные ошибки в его изложении, испытывает затруднения и допускает ошибки при выполнении "
            "заданий (решении задач), выполняет задание при подсказке преподавателя, затрудняется в формулировке выводов. "
            "Владеет знанием основных разделов, необходимых для дальнейшего обучения, знаком с основной и "
            "дополнительной литературой, рекомендованной программой."
        ))
        criteria_entries.append((
            "«неудовлетворительно»",
            "Обучающийся при ответе демонстрирует существенные пробелы в знаниях основного учебного материала, "
            "допускает грубые ошибки в формулировании основных понятий и при решении типовых задач (при выполнении "
            "типовых заданий), не способен ответить на наводящие вопросы преподавателя. Оценка ставится обучающимся, "
            "которые не могут продолжить обучение или приступить к профессиональной деятельности по окончании "
            "образовательного учреждения без дополнительных занятий по рассматриваемой дисциплине."
        ))
    else:
        criteria_entries.append((
            "«зачтено»",
            "Обучающийся продемонстрировал знание теоретического материала в объеме, достаточном для "
            "понимания сути предмета, умеет решать простейшие практические задачи и отвечать на основные вопросы. "
            "Допускаются несущественные неточности в изложении материала."
        ))
        criteria_entries.append((
            "«не зачтено»",
            "Обучающийся демонстрирует отсутствие базовых знаний по ключевым разделам дисциплины, "
            "не может ответить на наводящие вопросы преподавателя и решить простейшие типовые задачи."
        ))

    for grade, crit_text in criteria_entries:
        row_d = table_desc.add_row().cells
        set_row_cant_split(table_desc.rows[-1])
        set_cell_text(row_d[0], grade, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)
        set_cell_text(row_d[1], crit_text, size_pt=10)
        set_cell_width(row_d[0], 4.0)
        set_cell_width(row_d[1], 12.0)