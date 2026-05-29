import os
import re
import logging
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Настройка логирования
logger = logging.getLogger(__name__)


class CompetencyData:
    """Класс для хранения структуры данных одной компетенции."""

    def __init__(self, code_and_name: str):
        self.code_and_name = code_and_name
        self.indicators = []  # Список индикаторов
        self.disciplines = []  # Список дисциплин
        self.practices = []  # Список практик


def set_font(run, font_name="Times New Roman", size_pt=14, bold=False, italic=False, color_rgb=(0, 0, 0)):
    """Устанавливает шрифт, размер и начертание для текстового прогона (run),

    обеспечивая поддержку кириллицы в MS Word.
    """
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(*color_rgb)

    try:
        rPr = run._r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:cs'), font_name)
        rFonts.set(qn('w:eastAsia'), font_name)
        rPr.append(rFonts)
    except Exception as e:
        logger.debug(f"Не удалось применить XML-стили шрифтов: {e}")


def add_styled_paragraph(doc, text="", style_name="Normal", bold=False, italic=False, align=None,
                         space_before=0, space_after=0, line_spacing=1.0):
    """Добавляет абзац с заданным стилем MS Word и параметрами форматирования."""
    try:
        p = doc.add_paragraph(style=style_name)
    except Exception:
        p = doc.add_paragraph()

    if align is not None:
        p.alignment = align

    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing

    if text:
        run = p.add_run(text)
        # Если стиль является Заголовком, НЕ применяем ручное форматирование run-level,
        # чтобы Word не копировал эти шрифты в автооглавление.
        if "heading" in style_name.lower() or "заголовок" in style_name.lower():
            pass
        else:
            set_font(run, bold=bold, italic=italic)
    return p


def resolve_style_names(doc: Document):
    """Определяет доступные системные имена стилей в открытом документе."""
    existing_styles = {s.name for s in doc.styles}

    style_h1 = "Heading 1"
    for name in ["Заголовок 1", "Heading 1", "Heading1"]:
        if name in existing_styles:
            style_h1 = name
            break

    style_h2 = "Heading 2"
    for name in ["Заголовок 2", "Heading 2", "Heading2"]:
        if name in existing_styles:
            style_h2 = name
            break

    style_normal = "Normal"
    for name in ["Обычный", "Normal"]:
        if name in existing_styles:
            style_normal = name
            break

    return style_h1, style_h2, style_normal


def set_style_font_robust(style_obj, font_name="Times New Roman", size_pt=14, bold=False, italic=False,
                          color_rgb=(0, 0, 0)):
    """Устанавливает шрифт для стиля с гарантированной поддержкой кириллицы

    и полным отключением переопределения системных тем Word (Calibri/Arial).
    """
    font = style_obj.font
    font.name = font_name
    font.size = Pt(size_pt)
    font.bold = bold
    font.italic = italic
    font.color.rgb = RGBColor(*color_rgb)

    try:
        rPr = style_obj.element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()

        theme_attrs = ['asciiTheme', 'hAnsiTheme', 'eastAsiaTheme', 'cstheme']
        for attr in theme_attrs:
            attrib_key = qn(f'w:{attr}')
            if attrib_key in rFonts.attrib:
                del rFonts.attrib[attrib_key]

        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:cs'), font_name)
        rFonts.set(qn('w:eastAsia'), font_name)
    except Exception as e:
        logger.debug(f"Не удалось применить глубокую очистку стилей шрифтов: {e}")


def configure_toc_styles(doc: Document):
    """Настраивает форматирование строк оглавления (стили TOC 1 и TOC 2), сканируя базу стилей."""
    styles = doc.styles
    black_color = (0, 0, 0)

    def setup_toc_level(style_obj, size_pt, bold, italic):
        set_style_font_robust(style_obj, font_name="Times New Roman", size_pt=size_pt, bold=bold, italic=italic,
                              color_rgb=black_color)
        pf = style_obj.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(2)
        pf.line_spacing = 1.0

    for s in styles:
        name_lower = s.name.lower()
        if "toc" in name_lower or "оглавление" in name_lower:
            if "1" in name_lower:
                setup_toc_level(s, size_pt=14, bold=True, italic=True)
            elif "2" in name_lower:
                setup_toc_level(s, size_pt=13, bold=False, italic=True)


def configure_document_styles(doc: Document):
    """Программно настраивает глобальные стили документа согласно заданным параметрам."""
    style_h1, style_h2, style_normal = resolve_style_names(doc)
    black_color = (0, 0, 0)

    def setup_style(style_obj, size_pt=14, bold=False, italic=False, align=None):
        set_style_font_robust(style_obj, font_name="Times New Roman", size_pt=size_pt, bold=bold, italic=italic,
                              color_rgb=black_color)

        pf = style_obj.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0
        if align is not None:
            pf.alignment = align

    if style_normal in doc.styles:
        setup_style(doc.styles[style_normal], size_pt=14, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    if style_h1 in doc.styles:
        setup_style(doc.styles[style_h1], size_pt=14, bold=True, italic=False, align=WD_ALIGN_PARAGRAPH.CENTER)

    if style_h2 in doc.styles:
        setup_style(doc.styles[style_h2], size_pt=14, bold=False, italic=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    configure_toc_styles(doc)


def format_cell_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT):
    """Сбрасывает интервалы абзаца внутри ячеек таблицы."""
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0


def parse_indicators(text: str) -> list:
    """Разбивает текст ячейки индикаторов по кодам УК-Х.Y, ОПК-Х.Y и т.д."""
    pattern = r'(\b[А-Яа-яA-Za-z]+-\d+\.\d+(?:\.\d+)?\b)'
    parts = re.split(pattern, text)

    indicators = []
    current_code = None

    for part in parts:
        part_str = part.strip()
        if not part_str:
            continue
        if re.match(r'^[А-Яа-яA-Za-z]+-\d+\.\d+(?:\.\d+)?$', part_str):
            if current_code:
                indicators.append(current_code)
            current_code = part_str
        else:
            if current_code:
                indicators.append(f"{current_code} {part_str}")
                current_code = None
            else:
                if indicators:
                    indicators[-1] += " " + part_str
                else:
                    indicators.append(part_str)

    if current_code:
        indicators.append(current_code)

    return [re.sub(r'\s+', ' ', ind).strip() for ind in indicators if ind.strip()]


def clean_disciplines_or_practices(text: str) -> list:
    """Очищает список дисциплин или практик от лишних заголовков и пустых строк."""
    items = []
    text = text.replace('\xa0', ' ')
    for line in text.split('\n'):
        line_clean = " ".join(line.split()).strip()
        if not line_clean:
            continue
        line_lower = line_clean.lower()
        if line_lower in ('дисциплины', 'дисциплины (модули)', 'практики'):
            continue
        line_clean = re.sub(r'^[-\d\.\•\*\s*]+', '', line_clean).strip()
        if line_clean:
            items.append(line_clean)
    return items


def parse_word_table(doc: Document) -> list:
    """Находит в документе таблицу с компетенциями и парсит её структуру."""
    target_table = None

    for table in doc.tables:
        found = False
        for row in table.rows:
            row_texts = [cell.text.lower() for cell in row.cells]
            if any("код и наименование компетенции" in txt for txt in row_texts):
                target_table = table
                found = True
                break
        if found:
            break

    if not target_table:
        logger.warning("Таблица сопоставления компетенций не найдена в документе.")
        return []

    competencies = []
    current_comp = None

    start_row = 0
    for i, row in enumerate(target_table.rows):
        row_texts = [cell.text.lower() for cell in row.cells]
        if any("код и наименование компетенции" in txt for txt in row_texts):
            start_row = i + 1
            break

    for row in target_table.rows[start_row:]:
        if len(row.cells) < 4:
            continue

        comp_text = row.cells[0].text.strip()
        indicator_text = row.cells[1].text.strip()
        disc_text = row.cells[2].text.strip()
        prac_text = row.cells[3].text.strip()

        if not comp_text and not indicator_text and not disc_text and not prac_text:
            continue

        if comp_text:
            comp_text_clean = re.sub(r'\s+', ' ', comp_text).strip()
            if current_comp is None or current_comp.code_and_name != comp_text_clean:
                if current_comp:
                    competencies.append(current_comp)
                current_comp = CompetencyData(comp_text_clean)

        if current_comp is None:
            continue

        if indicator_text:
            inds = parse_indicators(indicator_text)
            for ind in inds:
                if ind not in current_comp.indicators:
                    current_comp.indicators.append(ind)

        if disc_text:
            discs = clean_disciplines_or_practices(disc_text)
            for d in discs:
                if d not in current_comp.disciplines:
                    current_comp.disciplines.append(d)

        if prac_text:
            pracs = clean_disciplines_or_practices(prac_text)
            for p in pracs:
                if p not in current_comp.practices:
                    current_comp.practices.append(p)

    if current_comp:
        competencies.append(current_comp)

    return competencies


def generate_section_3(doc: Document, competencies: list):
    """Добавляет Раздел 3 в конец текущего документа с использованием стилей."""
    style_h1, style_h2, style_normal = resolve_style_names(doc)
    logger.info(f"Раздел 3: Использование стилей H1='{style_h1}', H2='{style_h2}', Normal='{style_normal}'")

    # Переход к портретной (книжной) ориентации для Раздела 3
    new_section = doc.add_section()
    new_section.orientation = WD_ORIENT.PORTRAIT
    if new_section.page_width > new_section.page_height:
        w, h = new_section.page_width, new_section.page_height
        new_section.page_width = h
        new_section.page_height = w

    # Настройка полей книжной секции (Верх 2, Лево 3, Низ 2, Право 1.5)
    new_section.top_margin = Cm(2.0)
    new_section.bottom_margin = Cm(2.0)
    new_section.left_margin = Cm(3.0)
    new_section.right_margin = Cm(1.5)

    # Заголовок раздела (Стиль "Заголовок 1")
    add_styled_paragraph(
        doc,
        text="Раздел 3. Варианты диагностической работы",
        style_name=style_h1,
        bold=True
    )
    # Пустая строка после Раздела 3
    add_styled_paragraph(doc, style_name=style_normal)

    for comp_idx, comp in enumerate(competencies):
        # Каждая компетенция (начиная со второй) начинается со следующей страницы
        if comp_idx > 0:
            doc.add_page_break()

        # Подзаголовок "Компетенция" (Стиль "Обычный")
        add_styled_paragraph(doc, text="Компетенция", style_name=style_normal, bold=True)

        # Сама компетенция (Стиль "Заголовок 2", курсив наследуется)
        add_styled_paragraph(doc, text=comp.code_and_name, style_name=style_h2, italic=True)
        # Пустая строка после компетенции
        add_styled_paragraph(doc, style_name=style_normal)

        # Подзаголовок "Вариант" (Стиль "Обычный")
        add_styled_paragraph(doc, text="Вариант", style_name=style_normal, bold=True)

        # Шаблон первого вопроса
        p_q = add_styled_paragraph(doc, style_name=style_normal)
        run_q = p_q.add_run("1. ")
        set_font(run_q, bold=True)

        # 3 пустые строки под вопрос
        for _ in range(3):
            add_styled_paragraph(doc, style_name=style_normal)

        # Блок "Ключи теста" (Стиль "Обычный")
        add_styled_paragraph(doc, text="Ключи теста", style_name=style_normal)

        # Таблица ключей
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_ALIGN_PARAGRAPH.LEFT

        col_widths = [Inches(1.5), Inches(4.5)]
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = width

        # Шапка таблицы (Стиль ячеек)
        headers = ["Вопрос", "Ответ"]
        for idx, text in enumerate(headers):
            cell = table.cell(0, idx)
            cell.text = ""
            p = cell.paragraphs[0]
            try:
                p.style = style_normal
            except Exception:
                pass
            format_cell_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER)
            run = p.add_run(text)
            set_font(run, bold=True)

        # 5 пустых строк
        for row_idx in range(1, 6):
            cell_num = table.cell(row_idx, 0)
            cell_num.text = ""
            p_num = cell_num.paragraphs[0]
            try:
                p_num.style = style_normal
            except Exception:
                pass
            format_cell_paragraph(p_num, align=WD_ALIGN_PARAGRAPH.CENTER)
            run_num = p_num.add_run(str(row_idx))
            set_font(run_num, bold=False)

            cell_ans = table.cell(row_idx, 1)
            cell_ans.text = ""
            p_ans = cell_ans.paragraphs[0]
            try:
                p_ans.style = style_normal
            except Exception:
                pass
            format_cell_paragraph(p_ans, align=WD_ALIGN_PARAGRAPH.LEFT)
            run_ans = p_ans.add_run("")
            set_font(run_ans, bold=False)

        # Отступ после таблицы (Стиль "Обычный")
        add_styled_paragraph(doc, text="", style_name=style_normal)


class CompetencyReportGenerator:
    """Класс управления обновлением исходного документа."""

    def __init__(self, word_path: str):
        self.word_path = word_path

    def generate(self):
        logger.info(f"Загрузка документа: {self.word_path}")
        if not os.path.exists(self.word_path):
            logger.error(f"Файл не найден: {self.word_path}")
            print(f"Ошибка: Файл '{self.word_path}' не существует.")
            return

        try:
            doc = Document(self.word_path)
        except Exception as e:
            logger.error(f"Не удалось открыть документ: {e}")
            print("Ошибка чтения файла Word.")
            return

        # Программное переопределение настроек стилей на старте
        configure_document_styles(doc)

        logger.info("Анализ исходной таблицы компетенций...")
        competencies = parse_word_table(doc)

        if not competencies:
            logger.warning("Таблица компетенций не обнаружена.")
            print("В документе не найдена таблица с подходящей структурой.")
            return

        logger.info(f"Обнаружено компетенций для обработки: {len(competencies)}")

        logger.info("Добавление Раздела 3...")
        try:
            generate_section_3(doc, competencies)
        except Exception as e:
            logger.error(f"Ошибка при вставке Раздела 3: {e}")
            print("Произошла ошибка при генерации разметки шаблона.")
            return

        logger.info(f"Сохранение изменений в файл {self.word_path}...")
        try:
            doc.save(self.word_path)
            print(f"\n[Успешно] Шаблон Раздела 3 добавлен в исходный файл: {self.word_path}")
        except Exception as e:
            logger.error(f"Ошибка сохранения файла: {e}")
            print(
                "Не удалось сохранить файл. Убедитесь, что результирующий документ не открыт в другой программе (например, в MS Word).")


def main():
    print("=== Панель управления добавлением Раздела 3 ===")

    user_word_path: str = input("Введите путь к исходному файлу Word с таблицей (например, plan.docx): ").strip()

    if not user_word_path:
        user_word_path = "plan.docx"
        print(f"Используется файл по умолчанию: {user_word_path}")

    print("\nЗапуск процесса обновления документа...")
    generator = CompetencyReportGenerator(word_path=user_word_path)
    generator.generate()


if __name__ == "__main__":
    main()