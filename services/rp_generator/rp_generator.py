# -*- coding: utf-8 -*-
"""
Модуль rp_generator.py
Главный управляющий скрипт оркестрации процесса сборки РПД и ФОС из баз данных JSON.
"""

import os
import json
import logging
import re
import math
from datetime import date
from pathlib import Path
from typing import Dict, Any, List, Optional, Tuple

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# Прямой импорт утилит стилей, разметки и сокращений из базового модуля
from services.rp_generator.rp_doc_styles import (
    add_paragraph_with_spacing, set_cell_text, set_cell_background,
    set_cell_width, set_repeat_table_header, set_row_cant_split,
    abbreviate_discipline, format_author_initials
)

import services.rp_generator.rp_doc_sections as rsec

logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO, format="%(asctime)s | %(levelname)s | %(message)s")


def normalize_str(s: str) -> str:
    """Нормализует строку для нечеткого сравнения (удаляет пробелы, кавычки и знаки препинания)."""
    if not s:
        return ""
    return re.sub(r'[^a-zA-Zа-яА-Я0-9]', '', s.lower())


def parse_study_duration_to_years(duration_str: str) -> int:
    """
    Парсит строку срока обучения (например, '4 г.', '2 года 6 мес', '4.5')
    и возвращает округленное в большую сторону количество лет (целое число).
    """
    if not duration_str:
        return 4  # Дефолтное значение

    duration_str = duration_str.strip().lower()

    # Попытка парсинга как простого дробного числа
    try:
        val = float(duration_str.replace("г.", "").replace("г", "").strip())
        return int(math.ceil(val))
    except ValueError:
        pass

    years = 0
    months = 0

    # Поиск лет: "лет", "года", "год", "г", "г."
    year_match = re.search(r'(\d+)\s*(?:лет|года?|г\.?|y\.?)', duration_str)
    if year_match:
        years = int(year_match.group(1))

    # Поиск месяцев: "мес", "месяцев", "м", "м."
    month_match = re.search(r'(\d+)\s*(?:мес|месяцев|м\.?|m\.?)', duration_str)
    if month_match:
        months = int(month_match.group(1))

    # Резервный поиск первой цифры
    if years == 0 and months == 0:
        digit_match = re.search(r'(\d+)', duration_str)
        if digit_match:
            return int(digit_match.group(1))
        return 4

    total_years = years + (months / 12.0)
    return int(math.ceil(total_years))


class RPGenerator:
    """Класс-оркестратор для сборки рабочих программ дисциплин (РПД)."""

    def __init__(self, project_dir: Path, output_dir: Path):
        self.project_dir = project_dir
        self.output_dir = output_dir

        self.workload_path = project_dir / "rp_academic_workload.json"
        self.comp_map_path = project_dir / "rp_subject_competency_map.json"
        self.personnel_path = project_dir / "rp_personnel_mapping.json"
        self.ai_data_path = project_dir / "rp_ai_generated_data.json"

    def _load_json(self, path: Path) -> Dict[str, Any]:
        if not path.exists():
            raise FileNotFoundError(f"Файл не найден: {path.name}")
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)

    def _lookup_personnel(self, code: str, name: str, personnel_data: dict, metadata: dict) -> Dict[str, Any]:
        """
        Реляционный поиск кадрового состава с поддержкой нечеткого сопоставления дисциплин.
        Обеспечивает корректный подбор заведующего кафедрой (HOD) и составителей.
        """
        default_staff = personnel_data.get("default_department_personnel", {})

        deans = personnel_data.get("deans", {})
        heads = personnel_data.get("heads_of_department", {})
        pds = personnel_data.get("program_directors", {})
        umks = personnel_data.get("umk_chairmen", {})
        teachers = personnel_data.get("teachers", {})

        dir_key = f"{metadata.get('direction_code', '')} {metadata.get('direction_name', '')}".strip()
        subjects_mapping = personnel_data.get("subjects_mapping", {})

        # 1. Попытка точного поиска по направлению подготовки
        subjects_map = subjects_mapping.get(dir_key, {})
        subj_staff = subjects_map.get(name, {})

        # 2. Нечеткое резервное сопоставление (позволяет найти "Техносферную безопасность" и др.)
        if not subj_staff:
            norm_target = normalize_str(name)
            for other_dir_key, other_map in subjects_mapping.items():
                for other_sub_name, other_staff in other_map.items():
                    if normalize_str(other_sub_name) == norm_target:
                        subj_staff = other_staff
                        logger.info(
                            f"Дисциплина '{name}' сопоставлена по нечеткому совпадению с записью для '{other_sub_name}'")
                        break
                if subj_staff:
                    break

        def get_name(id_val, source_dict):
            return source_dict.get(id_val, {}).get("name", "") if id_val else ""

        hod_id = subj_staff.get("head_of_department") or default_staff.get("head_of_department")
        hod_info = heads.get(hod_id, {}) if hod_id else {}

        resolved = {
            "dean": get_name(subj_staff.get("dean") or default_staff.get("dean"), deans),
            "head_of_department": hod_info.get("name", ""),
            "head_of_department_role": hod_info.get("role", "заведующий кафедрой"),
            "program_director": get_name(subj_staff.get("program_director") or default_staff.get("program_director"),
                                         pds),
            "umk_chairman": get_name(default_staff.get("umk_chairman"), umks),
            "compilers": []
        }

        compiler_ids = subj_staff.get("teachers", [])
        for tid in compiler_ids:
            t_info = teachers.get(tid, {})
            if t_info:
                resolved["compilers"].append(f"{t_info.get('name')}, {t_info.get('degree_and_title')}")

        if not resolved["compilers"]:
            resolved["compilers"].append("Преподаватель кафедры")

        return resolved

    def generate_all(self):
        """Метод-оркестратор для сборки и сохранения документов РПД и ФОС."""
        # Загрузка баз данных
        workload = self._load_json(self.workload_path)
        comp_map = self._load_json(self.comp_map_path)
        personnel = self._load_json(self.personnel_path)
        ai_data = self._load_json(self.ai_data_path)

        metadata = workload.get("metadata", {})
        disciplines = workload.get("disciplines", {})
        comp_registry = comp_map.get("competencies_registry", {})
        subject_to_comp = comp_map.get("subject_to_competencies", {})

        # Динамическое считывание года начала подготовки учебного плана
        start_year = metadata.get("start_year") or date.today().year
        start_year = str(start_year).strip()

        # Определение срока обучения в годах с округлением в большую сторону
        duration_years = parse_study_duration_to_years(metadata.get("study_duration", "4 г."))

        # Автоматическая генерация списка учебных лет программы
        try:
            start_yr = int(start_year)
        except ValueError:
            start_yr = 2026
        years_list = [f"{start_yr + i}-{start_yr + i + 1}" for i in range(duration_years)]

        self.output_dir.mkdir(parents=True, exist_ok=True)

        for idx, (code, subj_info) in enumerate(disciplines.items(), start=1):
            subj_name = subj_info["name"]

            if code not in ai_data:
                logger.warning(f"Пропуск {code} {subj_name} — отсутствуют сгенерированные ИИ-данные.")
                continue

            logger.info(f"Генерация РПД [{idx}/{len(disciplines)}]: {code} {subj_name}...")

            subj_ai = ai_data[code]
            staff = self._lookup_personnel(code, subj_name, personnel, metadata)

            total_hours_dict = subj_info.get("total_hours", {})
            lectures_h = total_hours_dict.get("lectures", 0)
            practicals_h = total_hours_dict.get("practical_classes", 0)
            labs_h = total_hours_dict.get("laboratory_works", 0)

            doc = Document()

            # Установка полей согласно ГОСТ ИжГТУ
            section = doc.sections[0]
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(3.0)
            section.right_margin = Cm(1.5)

            # === СТРАНИЦА 1: ТИТУЛЬНЫЙ ЛИСТ ===
            rsec.generate_title_page(doc, metadata, subj_info, staff)

            # === СТРАНИЦА 2: ЛИСТ СОСТАВИТЕЛЕЙ И СОГЛАСОВАНИЯ ===
            rsec.generate_compilers_page(doc, metadata, subj_info, staff)

            # === СТРАНИЦА 3: АННОТАЦИЯ ДИСЦИПЛИНЫ ===
            mapped_comp = subject_to_comp.get(code, {}).get("competencies", {})
            rsec.generate_annotation_page(doc, metadata, subj_info, subj_ai, mapped_comp, comp_registry)

            # === СТРАНИЦА 4: РАЗДЕЛЫ 1, 2 ===
            rsec.generate_sections_1_2(doc, subj_ai, mapped_comp, comp_registry)

            # === СТРАНИЦА 5: МЕСТО В СТРУКТУРЕ ООП (РАЗДЕЛ 3) ===
            rsec.generate_section_3(doc, subj_info, subj_ai)

            # === РАЗДЕЛ 4: СТРУКТУРА И СОДЕРЖАНИЕ ДИСЦИПЛИНЫ ===
            add_paragraph_with_spacing(doc, "4. Структура и содержание дисциплины", bold=True, space_after=12)
            add_paragraph_with_spacing(doc, "4.1 Структура учебной нагрузки", bold=True, space_after=6)

            # Отрисовка сложной таблицы 4.1 со сбалансированным распределением остатков часов по разделам
            sems_active = sorted(list(subj_info.get("load_by_semester", {}).keys()), key=int)
            sections = subj_ai["thematic_plan"]["sections"]
            num_sections = len(sections)
            num_sems = len(sems_active)

            sem_to_sections = {sem: [] for sem in sems_active}
            for i_sec, s in enumerate(sections):
                target_sem = sems_active[min(i_sec * num_sems // num_sections, num_sems - 1)]
                sem_to_sections[target_sem].append(s)

            table_struct = doc.add_table(rows=0, cols=10)
            table_struct.alignment = WD_TABLE_ALIGNMENT.CENTER
            table_struct.style = "Table Grid"

            # Шаблон сложного заголовка
            row0 = table_struct.add_row().cells
            row1 = table_struct.add_row().cells
            row2 = table_struct.add_row().cells
            row3 = table_struct.add_row().cells

            for r in table_struct.rows[:4]:
                set_row_cant_split(r)

            set_cell_text(row0[0], "№", bold=True, size_pt=10, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(row0[1], "Раздел дисциплины. Форма промежуточной аттестации (по семестрам)", bold=True,
                          size_pt=10, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(row0[2], "Всего часов на раздел", bold=True, size_pt=10, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(row0[3], "Семестр", bold=True, size_pt=10, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(row0[4], "Распределение трудоемкости раздела (в часах) по видам учебной работы", bold=True,
                          size_pt=10, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(row0[9], "Содержание самостоятельной работы", bold=True, size_pt=10,
                          align=WD_ALIGN_PARAGRAPH.CENTER)

            set_cell_text(row1[4], "контактная", bold=True, size_pt=10, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(row1[8], "СРС", bold=True, size_pt=10, align=WD_ALIGN_PARAGRAPH.CENTER)

            set_cell_text(row2[4], "лек", bold=True, size_pt=9, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(row2[5], "пр", bold=True, size_pt=9, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(row2[6], "лаб", bold=True, size_pt=9, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(row2[7], "КЧА", bold=True, size_pt=9, align=WD_ALIGN_PARAGRAPH.CENTER)

            col_labels = ["1", "2", "3", "4", "5", "6", "7", "8", "10", "11"]
            for ci, label in enumerate(col_labels):
                set_cell_text(row3[ci], label, size_pt=9, align=WD_ALIGN_PARAGRAPH.CENTER, fill_hex="F2F2F2")

            row0[0].merge(row1[0]).merge(row2[0])
            row0[1].merge(row1[1]).merge(row2[1])
            row0[2].merge(row1[2]).merge(row2[2])
            row0[3].merge(row1[3]).merge(row2[3])
            row0[4].merge(row0[5]).merge(row0[6]).merge(row0[7]).merge(row0[8])
            row1[4].merge(row1[5]).merge(row1[6]).merge(row1[7])
            row1[8].merge(row2[8])
            row0[9].merge(row1[9]).merge(row2[9])

            for cell_to_paint in list(row0) + list(row1) + list(row2):
                set_cell_background(cell_to_paint, "F2F2F2")

            grand_total_hours = 0
            grand_total_lec = 0
            grand_total_prac = 0
            grand_total_lab = 0
            grand_total_kcha = 0.0
            grand_total_srs = 0.0

            for sem in sems_active:
                sem_load = subj_info["load_by_semester"][sem]
                sem_sections = sem_to_sections[sem]
                num_sem_sections = len(sem_sections)

                lec_h = sem_load.get("lectures", 0)
                prac_h = sem_load.get("practical_classes", 0)
                lab_h = sem_load.get("laboratory_works", 0)
                srs_h = sem_load.get("self_study", 0)

                # Накапливаем итоговые аудиторные часы по семестрам
                grand_total_lec += lec_h
                grand_total_prac += prac_h
                grand_total_lab += lab_h

                for idx_s, s in enumerate(sem_sections):
                    s_row = table_struct.add_row().cells
                    set_row_cant_split(table_struct.rows[-1])

                    s_lec = lec_h // num_sem_sections + (1 if idx_s < lec_h % num_sem_sections else 0)
                    s_prac = prac_h // num_sem_sections + (1 if idx_s < prac_h % num_sem_sections else 0)
                    s_lab = lab_h // num_sem_sections + (1 if idx_s < lab_h % num_sem_sections else 0)
                    s_srs = srs_h // num_sem_sections + (1 if idx_s < srs_h % num_sem_sections else 0)
                    s_total = s_lec + s_prac + s_lab + s_srs

                    # Динамическое определение характера СРС во избежание хардкодинга "лабораторных" там, где их нет
                    s_srs_desc = "Самостоятельное изучение теоретического материала и литературы по разделу"
                    if s_lab > 0 and s_prac > 0:
                        s_srs_desc = "Подготовка к защите лабораторных и практических работ по разделу"
                    elif s_lab > 0:
                        s_srs_desc = "Выполнение заданий и подготовка к защите лабораторных работ"
                    elif s_prac > 0:
                        s_srs_desc = "Решение практических задач и подготовка к защите практических работ"

                    set_cell_text(s_row[0], str(s["number"]), align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)
                    set_cell_text(s_row[1], s["name"], size_pt=10)
                    set_cell_text(s_row[2], str(s_total), align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)
                    set_cell_text(s_row[3], str(sem), align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)
                    set_cell_text(s_row[4], str(s_lec) if s_lec else "–", align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)
                    set_cell_text(s_row[5], str(s_prac) if s_prac else "–", align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)
                    set_cell_text(s_row[6], str(s_lab) if s_lab else "–", align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)
                    set_cell_text(s_row[7], "–", align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)
                    set_cell_text(s_row[8], str(s_srs) if s_srs else "–", align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)
                    set_cell_text(s_row[9], s_srs_desc, size_pt=9)

                control_info = sem_load.get("intermediate_control")
                ctrl_hours = 0
                if control_info:
                    ctrl_row = table_struct.add_row().cells
                    set_row_cant_split(table_struct.rows[-1])

                    ctrl_type_ru = "Экзамен" if control_info["type"] == "Exam" else "Зачет с оценкой" if control_info[
                                                                                                             "type"] == "GradedCredit" else "Зачет"
                    ctrl_hours = int(control_info["total"])
                    kcha_val = control_info["kcha"]
                    srs_val = control_info["self_study"]

                    set_cell_text(ctrl_row[0], "–", align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)
                    set_cell_text(ctrl_row[1], ctrl_type_ru, bold=True, size_pt=10)
                    set_cell_text(ctrl_row[2], str(ctrl_hours), align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)
                    set_cell_text(ctrl_row[3], str(sem), align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)
                    set_cell_text(ctrl_row[4], "–", align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)
                    set_cell_text(ctrl_row[5], "–", align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)
                    set_cell_text(ctrl_row[6], "–", align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)
                    set_cell_text(ctrl_row[7], f"{kcha_val:.1f}", align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)
                    set_cell_text(ctrl_row[8], f"{srs_val:.1f}".replace(".0", ""), align=WD_ALIGN_PARAGRAPH.CENTER,
                                  size_pt=10)
                    set_cell_text(ctrl_row[9],
                                  f"{ctrl_type_ru} выставляется по совокупности результатов текущего контроля успеваемости",
                                  size_pt=9)

                    grand_total_kcha += kcha_val

                # Суммируем часы напрямую по семестровым итогам (это на 100% исключает расхождения)
                sem_total_sum = lec_h + prac_h + lab_h + srs_h + ctrl_hours
                grand_total_hours += sem_total_sum

                sem_srs_sum = srs_h + (control_info['self_study'] if control_info else 0)
                grand_total_srs += sem_srs_sum

                sem_total_row = table_struct.add_row().cells
                set_row_cant_split(table_struct.rows[-1])

                set_cell_text(sem_total_row[1], f"Итого за {sem} семестр:", bold=True, size_pt=10)
                set_cell_text(sem_total_row[2], str(sem_total_sum), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
                              size_pt=10)
                set_cell_text(sem_total_row[3], str(sem), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)
                set_cell_text(sem_total_row[4], str(lec_h) if lec_h else "–", bold=True,
                              align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)
                set_cell_text(sem_total_row[5], str(prac_h) if prac_h else "–", bold=True,
                              align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)
                set_cell_text(sem_total_row[6], str(lab_h) if lab_h else "–", bold=True,
                              align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)
                set_cell_text(sem_total_row[7], f"{control_info['kcha']:.1f}" if control_info else "–", bold=True,
                              align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)
                set_cell_text(sem_total_row[8],
                              f"{srs_h + (control_info['self_study'] if control_info else 0):.1f}".replace(".0", ""),
                              bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)

            final_row = table_struct.add_row().cells
            set_row_cant_split(table_struct.rows[-1])
            set_cell_text(final_row[1], "Итого:", bold=True, size_pt=10)
            set_cell_text(final_row[2], str(grand_total_hours), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)
            set_cell_text(final_row[4], str(grand_total_lec) if grand_total_lec else "–", bold=True,
                          align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)
            set_cell_text(final_row[5], str(grand_total_prac) if grand_total_prac else "–", bold=True,
                          align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)
            set_cell_text(final_row[6], str(grand_total_lab) if grand_total_lab else "–", bold=True,
                          align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)
            set_cell_text(final_row[7], f"{grand_total_kcha:.1f}".replace(".0", ""), bold=True,
                          align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)
            set_cell_text(final_row[8], f"{grand_total_srs:.1f}".replace(".0", ""), bold=True,
                          align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)

            # === 4.2 Содержание разделов курса и формируемых в них компетенций ===
            add_paragraph_with_spacing(doc, "")
            add_paragraph_with_spacing(doc, "4.2 Содержание разделов курса и формируемых в них компетенций", bold=True,
                                       space_after=6)

            table_comp_map = doc.add_table(rows=1, cols=7)
            table_comp_map.alignment = WD_TABLE_ALIGNMENT.CENTER
            table_comp_map.style = "Table Grid"
            set_repeat_table_header(table_comp_map.rows[0])

            headers_map = ["№ п/п", "Раздел дисциплины", "Коды компетенций и индикаторов", "Знания", "Умения", "Навыки",
                           "Форма контроля"]
            for i_h, h_t in enumerate(headers_map):
                set_cell_text(table_comp_map.rows[0].cells[i_h], h_t, bold=True, size_pt=10,
                              align=WD_ALIGN_PARAGRAPH.CENTER, fill_hex="F2F2F2")

            for s in sections:
                row_cells = table_comp_map.add_row().cells
                set_row_cant_split(table_comp_map.rows[-1])

                set_cell_text(row_cells[0], str(s["number"]), align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)
                set_cell_text(row_cells[1], s["name"], size_pt=10)

                mapped_indicators = []
                for c_code, ind_list in mapped_comp.items():
                    mapped_indicators.extend(ind_list)
                indicators_str = "\n".join(mapped_indicators)

                # Динамический сбор уникальных меток ЗУН по привязанным к разделу индикаторам
                k_list, s_list, a_list = [], [], []
                for i_code in mapped_indicators:
                    ksa_entry = next((item for item in subj_ai["pedagogical_frame"].get("indicators_ksa", []) if
                                      item["indicator_code"] == i_code), None)
                    if ksa_entry:
                        k_list.extend(f"З{idx}" for idx in range(1, len(ksa_entry.get("knowledge", [])) + 1))
                        s_list.extend(f"У{idx}" for idx in range(1, len(ksa_entry.get("skills", [])) + 1))
                        a_list.extend(f"Н{idx}" for idx in range(1, len(ksa_entry.get("abilities", [])) + 1))

                k_labels = ", ".join(sorted(list(set(k_list))))
                s_labels = ", ".join(sorted(list(set(s_list))))
                a_labels = ", ".join(sorted(list(set(a_list))))

                set_cell_text(row_cells[2], indicators_str, size_pt=10)
                set_cell_text(row_cells[3], k_labels or "–", align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)
                set_cell_text(row_cells[4], s_labels or "–", align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)
                set_cell_text(row_cells[5], a_labels or "–", align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)

                controls = []
                if labs_h > 0:
                    controls.append("Защита лабораторных работ")
                if practicals_h > 0:
                    controls.append("Защита практических работ")
                set_cell_text(row_cells[6], "\n".join(controls) if controls else "Устный опрос", size_pt=9)

            # === 4.3 ТЕМЫ ЛЕКЦИЙ ===
            add_paragraph_with_spacing(doc, "")
            add_paragraph_with_spacing(doc, "4.3 Наименование тем лекций, их содержание и объем в часах", bold=True,
                                       space_after=6)

            table_lec = doc.add_table(rows=1, cols=4)
            table_lec.alignment = WD_TABLE_ALIGNMENT.CENTER
            table_lec.style = "Table Grid"
            set_repeat_table_header(table_lec.rows[0])

            set_cell_text(table_lec.rows[0].cells[0], "№ п/п", bold=True, size_pt=10, align=WD_ALIGN_PARAGRAPH.CENTER,
                          fill_hex="F2F2F2")
            set_cell_text(table_lec.rows[0].cells[1], "№ раздела", bold=True, size_pt=10,
                          align=WD_ALIGN_PARAGRAPH.CENTER, fill_hex="F2F2F2")
            set_cell_text(table_lec.rows[0].cells[2], "Наименование лекций", bold=True, size_pt=10,
                          align=WD_ALIGN_PARAGRAPH.CENTER, fill_hex="F2F2F2")
            set_cell_text(table_lec.rows[0].cells[3], "Трудоемкость (час)", bold=True, size_pt=10,
                          align=WD_ALIGN_PARAGRAPH.CENTER, fill_hex="F2F2F2")

            for i_l, l in enumerate(subj_ai["thematic_plan"]["lectures"], start=1):
                row_cells = table_lec.add_row().cells
                set_row_cant_split(table_lec.rows[-1])
                set_cell_text(row_cells[0], str(i_l), align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)
                set_cell_text(row_cells[1], str(l["section_number"]), align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)
                set_cell_text(row_cells[2], f"{l['theme']}\nСодержание: {l['content']}", size_pt=10)
                set_cell_text(row_cells[3], str(l["hours"]), align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)

            # === 4.4 ПРАКТИЧЕСКИЕ ЗАНЯТИЯ ===
            if practicals_h > 0:
                add_paragraph_with_spacing(doc, "")
                add_paragraph_with_spacing(doc,
                                           "4.4 Наименование тем практических занятий, их содержание и объем в часах",
                                           bold=True, space_after=6)

                table_prac = doc.add_table(rows=1, cols=4)
                table_prac.alignment = WD_TABLE_ALIGNMENT.CENTER
                table_prac.style = "Table Grid"
                set_repeat_table_header(table_prac.rows[0])

                set_cell_text(table_prac.rows[0].cells[0], "№ п/п", bold=True, size_pt=10,
                              align=WD_ALIGN_PARAGRAPH.CENTER, fill_hex="F2F2F2")
                set_cell_text(table_prac.rows[0].cells[1], "№ раздела", bold=True, size_pt=10,
                              align=WD_ALIGN_PARAGRAPH.CENTER, fill_hex="F2F2F2")
                set_cell_text(table_prac.rows[0].cells[2], "Наименование практических работ", bold=True, size_pt=10,
                              align=WD_ALIGN_PARAGRAPH.CENTER, fill_hex="F2F2F2")
                set_cell_text(table_prac.rows[0].cells[3], "Трудоемкость (час)", bold=True, size_pt=10,
                              align=WD_ALIGN_PARAGRAPH.CENTER, fill_hex="F2F2F2")

                for i_p, p in enumerate(subj_ai["thematic_plan"]["practicals"], start=1):
                    row_cells = table_prac.add_row().cells
                    set_row_cant_split(table_prac.rows[-1])
                    set_cell_text(row_cells[0], str(i_p), align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)
                    set_cell_text(row_cells[1], str(p["section_number"]), align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)
                    set_cell_text(row_cells[2], p["theme"], size_pt=10)
                    set_cell_width(row_cells[2], 10.0)
                    set_cell_text(row_cells[3], str(p["hours"]), align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)

            # === 4.5 ЛАБОРАТОРНЫЕ РАБОТЫ ===
            if labs_h > 0:
                add_paragraph_with_spacing(doc, "")
                add_paragraph_with_spacing(doc,
                                           "4.5 Наименование тем лабораторных работ, их содержание и объем в часах",
                                           bold=True, space_after=6)

                table_lab = doc.add_table(rows=1, cols=4)
                table_lab.alignment = WD_TABLE_ALIGNMENT.CENTER
                table_lab.style = "Table Grid"
                set_repeat_table_header(table_lab.rows[0])

                set_cell_text(table_lab.rows[0].cells[0], "№ п/п", bold=True, size_pt=10,
                              align=WD_ALIGN_PARAGRAPH.CENTER, fill_hex="F2F2F2")
                set_cell_text(table_lab.rows[0].cells[1], "№ раздела", bold=True, size_pt=10,
                              align=WD_ALIGN_PARAGRAPH.CENTER, fill_hex="F2F2F2")
                set_cell_text(table_lab.rows[0].cells[2], "Наименование лабораторных работ", bold=True, size_pt=10,
                              align=WD_ALIGN_PARAGRAPH.CENTER, fill_hex="F2F2F2")
                set_cell_text(table_lab.rows[0].cells[3], "Трудоемкость (час)", bold=True, size_pt=10,
                              align=WD_ALIGN_PARAGRAPH.CENTER, fill_hex="F2F2F2")

                for i_lb, lb in enumerate(subj_ai["thematic_plan"]["labs"], start=1):
                    row_cells = table_lab.add_row().cells
                    set_row_cant_split(table_lab.rows[-1])
                    set_cell_text(row_cells[0], str(i_lb), align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)
                    set_cell_text(row_cells[1], str(lb["section_number"]), align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)
                    set_cell_text(row_cells[2], lb["theme"], size_pt=10)
                    set_cell_width(row_cells[2], 10.0)
                    set_cell_text(row_cells[3], str(lb["hours"]), align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)

            # === РАЗДЕЛЫ 5, 6, 7 ===
            add_paragraph_with_spacing(doc, "")
            add_paragraph_with_spacing(doc,
                                       "5. Оценочные материалы для текущего контроля успеваемости и промежуточной аттестации по дисциплине",
                                       bold=True, space_after=12)
            add_paragraph_with_spacing(doc, "Для контроля результатов освоения дисциплины проводятся:")
            if labs_h > 0:
                add_paragraph_with_spacing(doc, "— защиты лабораторных работ;")
            if practicals_h > 0:
                add_paragraph_with_spacing(doc, "— защиты практических работ.")

            add_paragraph_with_spacing(doc,
                                       "Примечание: детальные оценочные материалы (варианты заданий, тесты, контрольные вопросы) "
                                       "приведены в Приложении к настоящей рабочей программе дисциплины.")

            add_paragraph_with_spacing(doc, "")
            add_paragraph_with_spacing(doc, "6.  Учебно-методическое и информационное обеспечение дисциплины:",
                                       bold=True, space_after=12)

            add_paragraph_with_spacing(doc, "а) основная литература:", bold=True)
            for idx_lit, lit in enumerate(subj_ai["resources_and_evaluation"].get("primary_literature", []), start=1):
                add_paragraph_with_spacing(doc, f"{idx_lit}. {lit}")

            add_paragraph_with_spacing(doc, "б) дополнительная литература:", bold=True)
            for idx_lit, lit in enumerate(subj_ai["resources_and_evaluation"].get("secondary_literature", []), start=1):
                add_paragraph_with_spacing(doc, f"{idx_lit}. {lit}")

            add_paragraph_with_spacing(doc, "в) методические указания:", bold=True)
            for idx_lit, lit in enumerate(subj_ai["resources_and_evaluation"].get("methodological_guidelines", []),
                                          start=1):
                add_paragraph_with_spacing(doc, f"{idx_lit}. {lit}")

            add_paragraph_with_spacing(doc, "г) перечень ресурсов информационно-коммуникационной сети Интернет:",
                                       bold=True)

            # Динамический перенос веб-ресурсов с надежной базовой заглушкой
            default_links = subj_ai["resources_and_evaluation"].get("internet_resources", [])
            if not default_links:
                default_links = [
                    "Система электронного обучения ИжГТУ имени М.Т. Калашникова — http://ee.istu.ru/",
                    "Электронно-библиотечная система IPRbooks / IPR SMART — http://istu.ru/material/elektronno-bibliotechnaya-sistema-iprbooks",
                    "Электронный каталог научной библиотеки ИжГТУ Web ИРБИС — http://94.181.117.43/cgi-bin/irbis64r_12/cgiirbis_64.exe",
                    "Национальная электронная библиотека (НЭБ) — http://нэб.рф",
                    "Научная электронная библиотека eLIBRARY.RU — https://elibrary.ru/defaultx.asp"
                ]
            for idx_l, link in enumerate(default_links, start=1):
                add_paragraph_with_spacing(doc, f"{idx_l}. {link}")

            add_paragraph_with_spacing(doc, "д) лицензионное и свободно распространяемое программное обеспечение:",
                                       bold=True)

            # ДИНАМИЧЕСКИЙ ВЫВОД ПО ИЗ JSON (Устранение хардкода Visual Studio и C# библиотек)
            default_software = subj_ai["resources_and_evaluation"].get("software", [])
            if not default_software:
                default_software = [
                    "Операционная система семейства Linux / MS Windows",
                    "Офисный пакет (Microsoft Office / LibreOffice)",
                    "Doctor Web Enterprise Suite (Лицензия № 116663324)"
                ]
            for idx_s, sw in enumerate(default_software, start=1):
                add_paragraph_with_spacing(doc, f"{idx_s}. {sw}")

            add_paragraph_with_spacing(doc, "")
            add_paragraph_with_spacing(doc, "7. Материально-техническое обеспечение дисциплины:", bold=True,
                                       space_after=12)
            add_paragraph_with_spacing(doc, "1. Лекционные занятия:")
            add_paragraph_with_spacing(doc,
                                       "Учебные аудитории для лекционных занятий укомплектованы специализированной мебелью "
                                       "и техническими средствами обучения (мультимедийный проектор, экран, персональный компьютер).")
            add_paragraph_with_spacing(doc, "2. Практические и лабораторные занятия:")
            add_paragraph_with_spacing(doc,
                                       "Лаборатории и компьютерные классы оснащены персональными компьютерами (ПЭВМ) "
                                       "с возможностью выхода в Интернет и доступом к локальной электронно-образовательной среде ИжГТУ.")

            add_paragraph_with_spacing(doc, "")
            add_paragraph_with_spacing(doc,
                                       "При необходимости рабочая программа дисциплины может быть адаптирована для обеспечения "
                                       "образовательного процесса инвалидов и лиц с ограниченными возможностями здоровья. Для этого "
                                       "требуется письменное заявление студента и заключение ПМПК.")

            # === ЛИСТЫ СОГЛАСОВАНИЙ ===
            doc.add_page_break()
            add_paragraph_with_spacing(doc, "Лист согласования рабочей программы дисциплины на учебный год", bold=True,
                                       align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
            add_paragraph_with_spacing(doc, f"Рабочая программа дисциплины «{subj_name}» по направлению подготовки "
                                            f"«{metadata.get('direction_code')} {metadata.get('direction_name')}» "
                                            f"согласована на ведение учебного процесса:")

            table_agree = doc.add_table(rows=len(years_list) + 1, cols=2)
            table_agree.alignment = WD_TABLE_ALIGNMENT.CENTER
            table_agree.style = "Table Grid"
            set_cell_text(table_agree.rows[0].cells[0], "Учебный год", bold=True, size_pt=11,
                          align=WD_ALIGN_PARAGRAPH.CENTER, fill_hex="F2F2F2")
            set_cell_text(table_agree.rows[0].cells[1],
                          "Согласовано\n(подпись и дата)", bold=True,
                          size_pt=11, align=WD_ALIGN_PARAGRAPH.CENTER, fill_hex="F2F2F2")

            for i_y, y_str in enumerate(years_list, start=1):
                set_row_cant_split(table_agree.rows[i_y])
                set_cell_text(table_agree.rows[i_y].cells[0], y_str, size_pt=11, align=WD_ALIGN_PARAGRAPH.CENTER)
                set_cell_text(table_agree.rows[i_y].cells[1], "", size_pt=11)

            # === ПРИЛОЖЕНИЕ: ОЦЕНОЧНЫЕ СРЕДСТВА (ФОС) ===
            rsec.generate_fos_appendix(doc, metadata, subj_info, subj_ai, mapped_comp, comp_registry, sems_active)

            # Сохранение готового файла DOCX
            abbr_discipline = abbreviate_discipline(subj_name)
            specialty_code = metadata.get("direction_code", "").strip()

            authors_abbr_list = [format_author_initials(c) for c in staff["compilers"] if format_author_initials(c)]
            authors_abbr = " ".join(authors_abbr_list)

            # Шаблон названия файла: {Код дисциплины} РП {АббревиатураПредмета} {КодСпециальности} {ИнициалыСоставителей}
            new_base_name = f"{code} РП {abbr_discipline} {specialty_code} {authors_abbr}".strip()
            safe_filename = re.sub(r'[\\/*?:"<>|]', "", new_base_name)
            output_file_path = self.output_dir / f"{safe_filename}.docx"

            try:
                doc.save(str(output_file_path.absolute()))
                logger.info(f"  [+] Успешно сгенерирована РПД: {output_file_path.name}")
            except Exception as e:
                logger.error(f"Не удалось сохранить РПД {subj_name}: {e}")

        print(
            f"\n[Успешно] Генерация всех рабочих программ завершена. Файлы сохранены в:\n{self.output_dir.absolute()}")


def main():
    print("=== Комплексный генератор рабочих программ дисциплин (РПД) ===")

    user_project_dir = input("Введите путь к папке с файлами JSON (по умолчанию 'services/rp_generator'): ").strip()
    if not user_project_dir:
        user_project_dir = "services/rp_generator"
    project_dir = Path(user_project_dir)

    user_output_dir = input(
        "Введите путь к папке для сохранения сгенерированных РПД (по умолчанию 'RPD_Output'): ").strip()
    if not user_output_dir:
        user_output_dir = "RPD_Output"
    output_dir = Path(user_output_dir)

    generator = RPGenerator(project_dir=project_dir, output_dir=output_dir)
    try:
        generator.generate_all()
    except Exception as e:
        logger.error(f"Ошибка при работе генератора РПД: {e}", exc_info=True)


if __name__ == "__main__":
    main()