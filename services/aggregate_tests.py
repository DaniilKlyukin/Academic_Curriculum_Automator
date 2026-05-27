import re
import copy
import io
import logging
from pathlib import Path
from typing import List, Dict, Tuple, Optional, Any
import openpyxl
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.shared import Pt


logger = logging.getLogger(__name__)


COMP_PATTERN = re.compile(
    r"\b(УК|ОПК|ПК|UK|OPK|PK|YK|OK|Yk|Pk|Opk|Uk)\s*[-–—\.\s]*(\d+)\b",
    re.IGNORECASE
)


def standardize_comp_code(code: str) -> str:
    """Нормализует код компетенции, переводя латинские буквы в кириллицу."""
    mapping = {
        'Y': 'У', 'y': 'у',
        'K': 'К', 'k': 'к',
        'O': 'О', 'o': 'о',
        'P': 'П', 'p': 'п',
        'C': 'С', 'c': 'с',
        'U': 'У', 'u': 'у',
        'B': 'В', 'b': 'в'
    }
    cleaned = re.sub(r"\s+", "", code).upper()
    return "".join(mapping.get(char, char) for char in cleaned)


def standardize_subject_code(code: str) -> str:
    """Приводит код дисциплины к единому стандарту с точками (например, 'Б1 О 11' -> 'Б1.О.11')."""
    mapping = {'B': 'Б', 'V': 'В', 'O': 'О', 'C': 'С'}
    cleaned = re.sub(r'[\s\._\-]+', '.', code).upper().strip('.')
    parts = cleaned.split('.')
    normalized_parts = []
    for part in parts:
        normalized_part = "".join(mapping.get(char, char) for char in part)
        normalized_parts.append(normalized_part)
    return ".".join(normalized_parts)


def clean_discipline_name(name: str) -> str:
    """Удаляет коды (шифры) из названия дисциплины, оставляя только текст."""
    cleaned = re.sub(r'^(Б\d+[\w\.\(\)]*[\s\.\-_]*)+', '', name)
    cleaned = re.sub(r'^(ФТД[\w\.\(\)]*[\s\.\-_]*)+', '', cleaned)
    return cleaned.strip().strip("«»\"")


def extract_code_from_filename(filename: str) -> str:
    """Извлекает код дисциплины из начала имени файла РП/ФОС."""
    name = filename.strip()
    match = re.match(r'^([БВОбвоБ1ФТДFTD\d\s\.\-_]+)', name)
    if match:
        raw_code = match.group(1).strip()
        raw_code = re.split(r'\s+(РП|ФОС)\s+', raw_code, flags=re.IGNORECASE)[0]
        return standardize_subject_code(raw_code)
    return ""


def sort_comp_key(code: str) -> Tuple[int, int]:
    """Сортировка компетенций в правильном порядке: УК -> ОПК -> ПК."""
    parts = code.split('-')
    prefix = parts[0].upper()
    order_map = {"УК": 1, "ОПК": 2, "ПК": 3}
    prefix_order = order_map.get(prefix, 4)

    num_part = 0
    if len(parts) > 1:
        num_str = re.sub(r'\D', '', parts[1])
        if num_str.isdigit():
            num_part = int(num_str)
    return (prefix_order, num_part)


def iter_block_items(parent: Any) -> Any:
    """Генератор, обходящий параграфы и таблицы в порядке их следования."""
    from docx.document import Document as _Document
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P

    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    else:
        parent_elm = parent._element

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def clear_document_content(doc: Document) -> None:
    """Полностью очищает контент документа, сохраняя стили."""
    body_elm = doc.element.body
    for child in list(body_elm):
        if child.tag.endswith('sectPr'):
            continue
        body_elm.remove(child)


def copy_media_relationships(src_doc: Document, dest_doc: Document, element_xml: Any) -> None:
    """Копирует связанные изображения и исправляет rId во избежание ошибок открытия файла."""
    embed_attr = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
    id_attr = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id'
    IMAGE_RELTYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"

    for el in list(element_xml.iter()):
        rId = el.get(embed_attr) or el.get(id_attr)
        if rId:
            copied_successfully = False
            if rId in src_doc.part.rels:
                src_rel = src_doc.part.rels[rId]
                if src_rel.reltype == IMAGE_RELTYPE:
                    try:
                        img_bytes = src_rel.target_part.blob

                        # Кросс-версионный способ доступа к созданию медиа-объектов
                        image_part = dest_doc.part.package.get_or_add_image_part(io.BytesIO(img_bytes))
                        new_rId = dest_doc.part.relate_to(image_part, IMAGE_RELTYPE)

                        if el.get(embed_attr):
                            el.set(embed_attr, new_rId)
                        else:
                            el.set(id_attr, new_rId)
                        copied_successfully = True
                    except Exception as e:
                        logger.warning(f"Пропущено копирование изображения rId {rId}: {e}")

            if not copied_successfully:
                # Если перенос не удался, удаляем весь сломанный XML-контейнер изображения,
                # чтобы Word не выдавал ошибку повреждения структуры при открытии.
                curr = el
                removed = False
                while curr is not None:
                    tag_local = curr.tag.split('}')[-1].lower()
                    if tag_local in ['drawing', 'shape', 'object', 'hyperlink']:
                        parent = curr.getparent()
                        if parent is not None:
                            parent.remove(curr)
                            removed = True
                            break
                    curr = curr.getparent()
                if not removed:
                    parent = el.getparent()
                    if parent is not None:
                        try:
                            parent.remove(el)
                        except Exception:
                            pass


def append_element_to_doc(dest_doc: Document, src_doc: Document, src_block: Any) -> None:
    """Безопасно добавляет XML-элемент в итоговый документ."""
    body = dest_doc.element.body
    sectPr = None
    for child in body:
        if child.tag.endswith('sectPr'):
            sectPr = child
            break

    copied_el = copy.deepcopy(src_block._element)
    copy_media_relationships(src_doc, dest_doc, copied_el)

    if sectPr is not None:
        sectPr.addprevious(copied_el)
    else:
        body.append(copied_el)


def parse_curriculum_excel(excel_path: Path) -> Tuple[Dict[str, str], Dict[str, Any]]:
    """Извлекает из Excel учебного плана все дисциплины, компетенции и индикаторы."""
    wb = openpyxl.load_workbook(str(excel_path.absolute()), data_only=True)

    # 1. Сбор названий дисциплин
    disciplines_map = {}
    plan_sheet = None
    for name in ["ПланСвод", "План"]:
        if name in wb.sheetnames:
            plan_sheet = wb[name]
            break
    if plan_sheet is None:
        plan_sheet = wb.active

    for row in range(1, plan_sheet.max_row + 1):
        col_b = str(plan_sheet.cell(row=row, column=2).value or "").strip()
        col_c = str(plan_sheet.cell(row=row, column=3).value or "").strip()
        if col_b and col_c:
            std_code = standardize_subject_code(col_b)
            if std_code.startswith("Б") or std_code.startswith("ФТД"):
                disciplines_map[std_code] = col_c

    # 2. Сбор компетенций и индикаторов
    competency_map = {}
    comp_sheet = None
    for name in wb.sheetnames:
        if "Компетенции" in name:
            comp_sheet = wb[name]
            break

    if comp_sheet:
        current_parent_code = None
        for row in range(1, comp_sheet.max_row + 1):
            c2 = str(comp_sheet.cell(row=row, column=2).value or "").strip()
            c3 = str(comp_sheet.cell(row=row, column=3).value or "").strip()
            c4 = str(comp_sheet.cell(row=row, column=4).value or "").strip()
            c5 = str(comp_sheet.cell(row=row, column=5).value or "").strip()

            # Определение базовой компетенции (например, УК-1)
            if c2 and not c3 and not c4:
                comp_code = standardize_comp_code(c2)
                current_parent_code = comp_code
                if comp_code not in competency_map:
                    competency_map[comp_code] = {
                        "description": c5,
                        "indicators": {},
                        "disciplines": []
                    }
            # Определение индикатора (например, УК-1.1)
            elif c3:
                ind_code = standardize_comp_code(c3)
                parent_comp = ind_code.split('.')[0] if '.' in ind_code else ind_code
                if parent_comp not in competency_map:
                    competency_map[parent_comp] = {
                        "description": "",
                        "indicators": {},
                        "disciplines": []
                    }
                competency_map[parent_comp]["indicators"][ind_code] = c5
                current_parent_code = parent_comp
            # Определение формирующей дисциплины
            elif c4 and current_parent_code:
                disp_code = standardize_subject_code(c4)
                raw_name = disciplines_map.get(disp_code, c5)
                disp_name = clean_discipline_name(raw_name)

                # Добавление в список без повторов
                if not any(d["code"] == disp_code for d in competency_map[current_parent_code]["disciplines"]):
                    competency_map[current_parent_code]["disciplines"].append({
                        "code": disp_code,
                        "name": disp_name
                    })
    return disciplines_map, competency_map


class SyllabusParser:
    """Парсер для интеллектуального извлечения тестов."""

    def parse_file(self, file_path: Path) -> List[Dict[str, Any]]:
        doc = Document(str(file_path.absolute()))
        extracted_tests = []

        in_diagnostic_section = False
        current_test_blocks = []
        current_test_competencies = set()
        test_started = False

        for block in iter_block_items(doc):
            if isinstance(block, Paragraph):
                text = block.text.strip()
                text_lower = text.lower()

                # Проверка входа во 2-й раздел диагностических материалов
                if any(t in text_lower for t in
                       ["раздел 2", "диагностические материалы", "оценочные средства", "оценочные материалы",
                        "фонд оценочных средств"]):
                    in_diagnostic_section = True

                if in_diagnostic_section:
                    # Поиск упоминания кодов компетенций
                    comp_matches = COMP_PATTERN.findall(text)
                    for m in comp_matches:
                        comp_code = standardize_comp_code(f"{m[0]}-{m[1]}")
                        current_test_competencies.add(comp_code)

                    # Триггер старта нового теста
                    is_test_start = False
                    if any(trig in text_lower for trig in
                           ["наименование: тест", "варианты тестов", "тест №", "проведение работы, заключающейся"]) or (
                            text.startswith("1. ") and not test_started):
                        is_test_start = True

                    if is_test_start:
                        if current_test_blocks:
                            extracted_tests.append({
                                "competencies": list(current_test_competencies),
                                "blocks": list(current_test_blocks)
                            })
                            current_test_blocks = []
                            current_test_competencies = set()
                            for m in comp_matches:
                                current_test_competencies.add(standardize_comp_code(f"{m[0]}-{m[1]}"))

                        test_started = True

                    if test_started:
                        # Проверка маркеров завершения тестов
                        if any(end in text_lower for end in ["критерии и шкалы", "раздел 3", "2. критерии"]):
                            if current_test_blocks:
                                extracted_tests.append({
                                    "competencies": list(current_test_competencies),
                                    "blocks": list(current_test_blocks)
                                })
                            current_test_blocks = []
                            current_test_competencies = set()
                            test_started = False
                        else:
                            current_test_blocks.append((block, doc))

            elif isinstance(block, Table):
                if in_diagnostic_section and test_started:
                    current_test_blocks.append((block, doc))

        # Сохранение последнего незавершенного блока
        if current_test_blocks:
            extracted_tests.append({
                "competencies": list(current_test_competencies),
                "blocks": list(current_test_blocks)
            })

        return extracted_tests


def main() -> None:
    print("=== Панель агрегации оценочных материалов (тестов) из РП ===")
    excel_path_str = input("Шаг 1. Введите путь к Excel-плану: ").strip()
    folder_path_str = input("Шаг 2. Введите путь к папке с РП (файлы .docx): ").strip()
    word_path_str = input("Шаг 3. Введите путь для сохранения файла или папки: ").strip()

    if not folder_path_str:
        folder_path_str = "."

    folder_path = Path(folder_path_str)

    if word_path_str.endswith(".docx"):
        word_path = Path(word_path_str)
    else:
        word_path = Path(word_path_str) / "Сводные_оценочные_материалы_ФОС.docx"
    word_path.parent.mkdir(parents=True, exist_ok=True)

    # 1. Загрузка данных плана
    print("\n[Excel] Чтение учебного плана...")
    disciplines_map, competency_map = parse_curriculum_excel(Path(excel_path_str))

    # 2. Анализ файлов РП
    files = [f for f in folder_path.glob("*.docx") if not f.name.startswith("~$")]
    if not files:
        print("В указанной папке файлы .docx не найдены.")
        return

    # Извлеченные тесты: {Код_Компетенции: {Код_Дисциплины: [Блоки_РП]}}
    extracted_tests_db: Dict[str, Dict[str, List[Tuple[Any, Document]]]] = {}
    parser = SyllabusParser()

    print("\n[Word] Извлечение оценочных материалов из файлов...")
    for i, file_path in enumerate(files, 1):
        try:
            subj_code = extract_code_from_filename(file_path.name)
            disc_name = clean_discipline_name(disciplines_map.get(subj_code, file_path.stem))

            # Парсинг теста из РП
            tests = parser.parse_file(file_path)
            print(f"[{i:03}/{len(files):03}] Анализ файла: {file_path.name} | Предмет: «{disc_name}»")

            for test in tests:
                comps = test["competencies"]
                blocks = test["blocks"]

                if comps:
                    # Тест привязан к конкретным компетенциям в файле
                    for comp_code in comps:
                        if comp_code not in extracted_tests_db:
                            extracted_tests_db[comp_code] = {}
                        extracted_tests_db[comp_code][subj_code] = blocks
                        print(f"      -> Найдено тестовое задание для {comp_code}")
                else:
                    # Тест не имеет явной разметки компетенций в РП
                    # Дублируем его во все компетенции, к которым привязана дисциплина по учебному плану!
                    matched_any = False
                    for comp_code, comp_data in competency_map.items():
                        if any(d["code"] == subj_code for d in comp_data["disciplines"]):
                            if comp_code not in extracted_tests_db:
                                extracted_tests_db[comp_code] = {}
                            extracted_tests_db[comp_code][subj_code] = blocks
                            matched_any = True
                    if matched_any:
                        print(
                            f"      -> (!) Неразмеченный тест распределен по всем компетенциям дисциплины «{disc_name}»")

        except Exception as e:
            logger.error(f"Не удалось обработать файл {file_path.name}: {e}")

    # 3. Сборка итогового документа
    print("\n[Сборка] Создание итогового структурированного документа...")
    output_doc = Document()
    clear_document_content(output_doc)

    # Титульный заголовок раздела
    p_title = output_doc.add_paragraph()
    r_title = p_title.add_run("Раздел 2. Диагностические материалы для оценки сформированности компетенций")
    r_title.bold = True
    r_title.font.size = Pt(14)
    p_title.alignment = 1  # По центру

    # Сортировка по типу компетенций (УК -> ОПК -> ПК)
    sorted_comp_codes = sorted(competency_map.keys(), key=sort_comp_key)

    for comp_code in sorted_comp_codes:
        comp_data = competency_map[comp_code]

        output_doc.add_paragraph()  # Пустая строка для отступа

        # Вывод названия компетенции
        p_comp_label = output_doc.add_paragraph()
        p_comp_label.add_run("Компетенция").bold = True

        p_comp_desc = output_doc.add_paragraph()
        p_comp_desc.add_run(f"{comp_code}. {comp_data['description']}").italic = True

        output_doc.add_paragraph()

        # Вывод индикаторов
        p_ind_label = output_doc.add_paragraph()
        p_ind_label.add_run("Индикаторы достижения компетенции:").bold = True
        for ind_code, ind_desc in comp_data["indicators"].items():
            p_ind = output_doc.add_paragraph()
            p_ind.add_run(f"{ind_code}. {ind_desc}")

        output_doc.add_paragraph()

        # Вывод списка формирующих дисциплин
        p_disc_label = output_doc.add_paragraph()
        p_disc_label.add_run("Дисциплины и практики, формирующие компетенцию:").bold = True
        for disc in comp_data["disciplines"]:
            p_disc = output_doc.add_paragraph()
            p_disc.add_run(f"{disc['name']}.")

        output_doc.add_paragraph()

        # Вывод оценочных материалов
        p_eval_label = output_doc.add_paragraph()
        p_eval_label.add_run("Оценочные материалы").bold = True

        has_any_tests = False
        for disc in comp_data["disciplines"]:
            code = disc["code"]
            name = disc["name"]

            # Если для данной дисциплины и данной компетенции есть тест
            if comp_code in extracted_tests_db and code in extracted_tests_db[comp_code]:
                has_any_tests = True
                p_disc_header = output_doc.add_paragraph()
                p_disc_header.add_run(f"Дисциплина «{name}»").bold = True

                # Копируем структуру теста
                for block, src_doc in extracted_tests_db[comp_code][code]:
                    append_element_to_doc(output_doc, src_doc, block)

                output_doc.add_paragraph()

        if not has_any_tests:
            p_empty = output_doc.add_paragraph()
            p_empty.add_run("[Тестовые материалы для данной компетенции отсутствуют или не найдены в РП]").italic = True

        output_doc.add_page_break()

    output_doc.save(str(word_path.absolute()))
    print(f"\nАгрегация завершена. Файл успешно сохранен в: {word_path.name}")


if __name__ == "__main__":
    main()