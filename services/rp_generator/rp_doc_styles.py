# -*- coding: utf-8 -*-
"""
Модуль rp_doc_styles.py
Содержит базовые манипуляции стилями Word, XML-разметкой ячеек и утилиты сокращения строк.
"""

import re
import logging
from typing import Any

from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


def set_font(run, font_name="Times New Roman", size_pt=14, bold=False, italic=False, color_rgb=(0, 0, 0)):
    """Устанавливает шрифт, размер и начертание для текстового прогона."""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(*color_rgb)

    try:
        rPr = run._r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:cs'), font_name)
        rFonts.set(qn('w:eastAsia'), font_name)
        rPr.append(rFonts)
    except Exception as e:
        logger.debug(f"Не удалось применить стили шрифтов: {e}")


def add_paragraph_with_spacing(doc, text="", style="Normal", bold=False, italic=False, align=None,
                               space_after=6, space_before=0, line_spacing=1.0) -> Any:
    """Добавляет абзац с фиксированными отступами и одинарным интервалом."""
    p = doc.add_paragraph()
    try:
        p.style = style
    except Exception:
        pass
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    if align:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_font(run, bold=bold, italic=italic)
    return p


def set_repeat_table_header(row):
    """Повторение шапки таблицы на каждой странице."""
    trPr = row._tr.get_or_add_trPr()
    trPr.append(OxmlElement('w:tblHeader'))


def set_row_cant_split(row):
    """Запрет разрыва строки таблицы на границе страниц."""
    trPr = row._tr.get_or_add_trPr()
    trPr.append(OxmlElement('w:cantSplit'))


def set_cell_background(cell, fill_hex):
    """Устанавливает фоновый цвет заливки для ячейки."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Задает внутренние поля ячейки в dxa (1 pt = 20 dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def set_cell_text(cell, text: str, bold=False, italic=False, size_pt=12, align=WD_ALIGN_PARAGRAPH.LEFT,
                  vertical_align=WD_ALIGN_VERTICAL.CENTER, fill_hex=None):
    """Форматирует ячейку таблицы и записывает в нее текст."""
    cell.text = ""
    cell.vertical_alignment = vertical_align
    set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
    if fill_hex:
        set_cell_background(cell, fill_hex)

    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    if text:
        run = p.add_run(text)
        set_font(run, size_pt=size_pt, bold=bold, italic=italic)


def set_cell_width(cell, width_cm: float):
    """Задает ширину ячейки в сантиметрах."""
    cell.width = Cm(width_cm)


def abbreviate_word(word: str) -> str:
    """Сокращает отдельное слово до первой согласной после первой гласной."""
    vowels = set("аеёиоуыэюя")
    word_clean = "".join([c for c in word if c.isalnum() or c in ('+', '#')])
    if not word_clean:
        return ""

    first_vowel_idx = -1
    for idx, char in enumerate(word_clean):
        if char.lower() in vowels:
            first_vowel_idx = idx
            break

    if first_vowel_idx == -1:
        return word_clean

    first_consonant_after_vowel_idx = -1
    for idx in range(first_vowel_idx + 1, len(word_clean)):
        if word_clean[idx].lower() not in vowels:
            first_consonant_after_vowel_idx = idx
            break

    if first_consonant_after_vowel_idx == -1:
        return word_clean

    return word_clean[:first_consonant_after_vowel_idx + 1]


def abbreviate_discipline(name: str) -> str:
    """Формирует сжатое название дисциплины (например, Информационные технологии -> ИнформТехнолог)."""
    normalized_name = name.replace("-", " ")
    words = normalized_name.split()
    stop_words = {"на", "по", "в", "с", "для", "под", "о", "об", "за", "из", "от", "до", "без"}

    abbr_words = []
    for w in words:
        if w.lower() in stop_words:
            continue
        abbr = abbreviate_word(w)
        if not abbr:
            continue
        if w.lower() == "и":
            abbr_words.append("и")
        else:
            abbr_words.append(abbr.capitalize())
    return "".join(abbr_words)


def format_author_initials(compiler_str: str) -> str:
    """Преобразует строку составителя в компактные инициалы (например, 'Севостьянов Б.В. -> БВСевостьянов')."""
    clean_line = compiler_str.split(",", 1)[0].strip()
    if any(word in clean_line.lower() for word in ["преподаватель", "кафедр", "составител", "разработчик"]):
        return ""

    words = clean_line.split()
    if not words:
        return ""

    last_name = words[0].capitalize()
    first_init = words[1][0].upper() if len(words) > 1 and words[1] else ""
    middle_init = words[2][0].upper() if len(words) > 2 and words[2] else ""
    return f"{first_init}{middle_init}{last_name}"


def merge_cells_vertically(table, col_idx: int):
    """Выполняет вертикальное объединение ячеек с одинаковым значением в столбце col_idx."""
    start_row = 1
    while start_row < len(table.rows):
        val = table.cell(start_row, col_idx).text.strip()
        if not val:
            start_row += 1
            continue
        end_row = start_row + 1
        while end_row < len(table.rows) and table.cell(end_row, col_idx).text.strip() == val:
            end_row += 1

        if end_row > start_row + 1:
            base_cell = table.cell(start_row, col_idx)
            for r in range(start_row + 1, end_row):
                base_cell.merge(table.cell(r, col_idx))
            set_cell_text(base_cell, val, size_pt=10, align=WD_ALIGN_PARAGRAPH.LEFT,
                          vertical_align=WD_ALIGN_VERTICAL.CENTER)
        start_row = end_row


def get_department_acronym(name: str) -> str:
    """Формирует аббревиатуру кафедры (например, 'Прикладная математика и информационные технологии' -> 'ПМиИТ')."""
    if not name:
        return "ПМиИТ"
    words = name.replace("-", " ").split()
    stop_words = {"и", "на", "по", "в", "с", "для", "под", "о", "об", "за", "из", "от", "до", "без"}
    parts = []
    for w in words:
        clean_w = "".join(c for c in w if c.isalpha())
        if not clean_w:
            continue
        if clean_w.lower() in stop_words:
            parts.append(clean_w.lower())
        else:
            parts.append(clean_w[0].upper())
    return "".join(parts)